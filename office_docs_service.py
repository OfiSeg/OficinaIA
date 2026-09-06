"""Persistencia y manipulación de los documentos internos de OficinaIA.

Extraído de app.py en V20 Etapa 15. No conoce Flask, request, session ni rutas.
"""
from pathlib import Path
import os
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter
from docx import Document


def fila_vacia(fila):
    return not any(str(valor or "").strip() for valor in fila)


def normalizar_matriz_excel(filas):
    if not isinstance(filas, list):
        raise ValueError("La matriz no es válida.")
    normalizadas = []
    for fila in filas[:500]:
        if not isinstance(fila, list):
            continue
        normalizadas.append(["" if valor is None else str(valor) for valor in fila[:30]])
    return normalizadas


def limpiar_filas_excel(filas, conservar_vacias=False):
    normalizadas = normalizar_matriz_excel(filas)
    if not normalizadas:
        return []
    if conservar_vacias:
        return normalizadas
    encabezado = normalizadas[0]
    cuerpo = [fila for fila in normalizadas[1:] if not fila_vacia(fila)]
    return [encabezado] + cuerpo


def limpiar_columnas_excel(filas):
    filas = normalizar_matriz_excel(filas)
    if not filas:
        return []
    max_cols = max((len(f) for f in filas), default=0)
    if not max_cols:
        return filas
    vivas = [c for c in range(max_cols) if any(str(f[c] if c < len(f) else "").strip() for f in filas)]
    if not vivas:
        return [[""]]
    return [[fila[c] if c < len(fila) else "" for c in vivas] for fila in filas]


class OfficeDocumentsService:
    def __init__(self, *, base_dir, libros_excel, word_file,
                 descargar_excel, subir_excel,
                 usar_pg_documento, pg_obtener_documento, pg_guardar_documento,
                 invalidar_cache_excel=None):
        self.base_dir = Path(base_dir)
        self.libros_excel = libros_excel
        self.word_file = Path(word_file)
        self.descargar_excel = descargar_excel
        self.subir_excel = subir_excel
        self.usar_pg_documento = usar_pg_documento
        self.pg_obtener_documento = pg_obtener_documento
        self.pg_guardar_documento = pg_guardar_documento
        self.invalidar_cache_excel = invalidar_cache_excel

    @staticmethod
    def r2_excel_configurado():
        return all(os.getenv(n) for n in (
            "R2_ENDPOINT_URL", "R2_ACCESS_KEY_ID", "R2_SECRET_ACCESS_KEY", "R2_BUCKET_NAME"
        ))

    def _libro(self, libro_id):
        libro_id = str(libro_id or "1")
        if libro_id not in self.libros_excel:
            raise ValueError("Libro de Excel no válido.")
        return libro_id, self.libros_excel[libro_id]

    def asegurar_excel(self, libro_id="1"):
        libro_id, libro = self._libro(libro_id)
        archivo = self.base_dir / libro["archivo"]
        if self.r2_excel_configurado():
            try:
                if self.descargar_excel(archivo, libro["r2_key"]):
                    return archivo
                if not archivo.exists():
                    self._crear_excel(archivo)
                self.subir_excel(archivo, libro["r2_key"])
                return archivo
            except Exception as error:
                if archivo.exists():
                    print("ADVERTENCIA EXCEL R2:", error)
                    print(f"Se utilizará temporalmente la copia local del libro {libro_id}.")
                    return archivo
                raise RuntimeError(
                    f"No se pudo recuperar el libro Excel {libro_id} desde Cloudflare R2 y tampoco existe una copia local."
                ) from error
        if not archivo.exists():
            self._crear_excel(archivo)
        return archivo

    @staticmethod
    def _crear_excel(archivo):
        wb = Workbook(); ws = wb.active; ws.title = "Datos"
        ws.append(["Dato", "Valor", "Observaciones"]); wb.save(archivo)

    def leer_excel(self, libro_id="1"):
        libro_id, libro = self._libro(libro_id)
        archivo = self.asegurar_excel(libro_id)
        wb = load_workbook(archivo, data_only=False); ws = wb.active
        filas = [["" if v is None else str(v) for v in row] for row in ws.iter_rows(values_only=True)]
        filas = limpiar_filas_excel(filas, conservar_vacias=True)
        columnas = max(1, min(max([len(f) for f in filas], default=1), 30))
        filas = [f[:columnas] + [""] * (columnas - len(f)) for f in filas]
        return {"hoja": ws.title, "filas": filas, "columnas": columnas}

    def guardar_excel(self, filas, nombre_hoja="Datos", libro_id="1"):
        libro_id, libro = self._libro(libro_id)
        archivo = self.base_dir / libro["archivo"]
        filas = limpiar_filas_excel(filas, conservar_vacias=True)
        max_cols = min(max([len(f) for f in filas], default=1), 30)
        wb = Workbook(); ws = wb.active; ws.title = (nombre_hoja or "Datos")[:31]
        for r, fila in enumerate(filas, start=1):
            for c in range(1, max_cols + 1):
                valor = fila[c - 1] if c - 1 < len(fila) else ""
                ws.cell(row=r, column=c, value="" if valor is None else str(valor))
        for c in range(1, max_cols + 1):
            letra = get_column_letter(c)
            valores = [str(ws.cell(r, c).value or "") for r in range(1, min(ws.max_row, 30) + 1)]
            ws.column_dimensions[letra].width = min(max([len(v) for v in valores] + [10]) + 2, 32)
        wb.save(archivo)
        if self.r2_excel_configurado():
            try:
                self.subir_excel(archivo, libro["r2_key"])
            except Exception as error:
                print(f"ERROR SINCRONIZANDO LIBRO {libro_id} CON R2:", error)
                raise
        if libro_id == "1" and self.invalidar_cache_excel:
            try:
                self.invalidar_cache_excel()
            except Exception as error:
                print("AVISO INVALIDANDO CACHE EXCEL IA:", error)
        return archivo

    def importar_excel(self, archivo_subido, libro_id="1"):
        libro_id, libro = self._libro(libro_id)
        destino = self.base_dir / libro["archivo"]
        temporal = destino.with_suffix(".upload.xlsx")
        try:
            archivo_subido.save(temporal)
            wb = load_workbook(temporal, data_only=False)
            if not wb.sheetnames:
                raise ValueError("El Excel no contiene hojas.")
            ws = wb[wb.sheetnames[0]]
            filas = [["" if v is None else str(v) for v in row] for row in ws.iter_rows(values_only=True)]
            self.guardar_excel(filas, ws.title, libro_id)
            return self.leer_excel(libro_id)
        finally:
            try: temporal.unlink(missing_ok=True)
            except Exception: pass

    def asegurar_word(self):
        if not self.word_file.exists():
            doc = Document(); doc.add_paragraph(""); doc.save(self.word_file)

    def leer_word(self):
        if self.usar_pg_documento():
            try:
                contenido = self.pg_obtener_documento()
                return contenido if contenido is not None else ""
            except Exception as error:
                print("ERROR leer_word_interno PG:", error)
                return ""
        self.asegurar_word()
        doc = Document(self.word_file)
        return "\n\n".join(p.text for p in doc.paragraphs)

    def guardar_word(self, contenido):
        if self.usar_pg_documento():
            self.pg_guardar_documento(str(contenido or "")); return
        doc = Document()
        for linea in str(contenido or "").splitlines(): doc.add_paragraph(linea)
        doc.save(self.word_file)

    def generar_docx(self):
        if self.usar_pg_documento():
            contenido = self.leer_word(); doc = Document()
            for linea in contenido.splitlines(): doc.add_paragraph(linea)
            doc.save(self.word_file)
        else:
            self.asegurar_word()
        return self.word_file
