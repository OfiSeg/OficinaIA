from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    session,
    jsonify,
    send_from_directory,
    Response,
    stream_with_context,
)

from pathlib import Path
from functools import wraps
import fitz  # PyMuPDF
import re
import unicodedata
import os
import json
import sqlite3
import logging
import traceback
from contextlib import closing
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

import pendientes_ops
from database_pg import (
    inicializar_postgres,
    listar_manuales,
    registrar_manual,
    actualizar_manual,
    eliminar_manual as eliminar_manual_pg,
    obtener_manual_por_r2_key,
    listar_polizas as pg_listar_polizas,
    obtener_poliza_por_r2_key,
    registrar_poliza,
    eliminar_poliza as eliminar_poliza_pg,
    listar_usuarios as pg_listar_usuarios,
    obtener_usuario as pg_obtener_usuario,
    obtener_usuario_por_id as pg_obtener_usuario_por_id,
    usuario_existe as pg_usuario_existe,
    crear_usuario as pg_crear_usuario,
    actualizar_usuario as pg_actualizar_usuario,
    eliminar_usuario as pg_eliminar_usuario,
    crear_conversacion as pg_crear_conversacion,
    listar_chats as pg_listar_chats,
    validar_chat as pg_validar_chat,
    obtener_chat_con_mensajes as pg_obtener_chat_con_mensajes,
    eliminar_chat as pg_eliminar_chat,
    agregar_mensaje as pg_agregar_mensaje,
    obtener_configuracion as pg_obtener_configuracion,
    guardar_configuracion as pg_guardar_configuracion,
    obtener_documento_interno as pg_obtener_documento_interno,
    guardar_documento_interno as pg_guardar_documento_interno,
    obtener_flota_activa as pg_obtener_flota_activa,
    guardar_flota_activa as pg_guardar_flota_activa,
    borrar_flota_activa as pg_borrar_flota_activa,
)
from storage_r2 import (
    subir_pdf as r2_subir_pdf,
    eliminar_pdf as r2_eliminar_pdf,
    descargar_pdf_temporal,
    obtener_objeto_stream,
    EXCEL_INTERNO_R2_KEY,
    subir_excel_interno,
    descargar_excel_interno,
)
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter
from docx import Document
from coti import procesar_comando_coti
from servicios_ia import buscar_en_metadatos
from companias import normalizar_compania, nombre_compania as _nombre_compania_canonico

# ==========================================================
# CONFIGURACIÓN
# ==========================================================

BASE_DIR = Path(__file__).resolve().parent
# load_dotenv ANTES de leer cualquier env (P0.7). En Render las vars del
# panel ya están en el proceso; en local el .env debe aplicar a secret_key.
load_dotenv(BASE_DIR / ".env")

app = Flask(__name__)

# P0.1 — Secret de sesión: en producción (Neon/Render) es obligatorio.
# El default solo se tolera en desarrollo local sin DATABASE_URL.
_secret = os.getenv("FLASK_SECRET_KEY")
_es_produccion = bool(os.getenv("DATABASE_URL") or os.getenv("RENDER"))
if _es_produccion:
    if not _secret or _secret == "OFICINA_SEGUROS_CAMBIAR_CLAVE":
        raise RuntimeError(
            "FLASK_SECRET_KEY debe estar definida en producción y no puede "
            "ser el valor por defecto. Configurala en el panel de Render."
        )
    app.secret_key = _secret
else:
    app.secret_key = _secret or "OFICINA_SEGUROS_CAMBIAR_CLAVE"

# Cookies de sesión endurecidas (P0.1 / P1.1).
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=_es_produccion,  # solo HTTPS en prod
    MAX_CONTENT_LENGTH=20 * 1024 * 1024,
)

# P-FIX500 — Logger técnico para excepciones no controladas. Va a stdout,
# que es lo que Render captura como logs del servicio.
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger("oficinaia")


# P-FIX500 — Red de seguridad global: ninguna excepción no controlada debe
# terminar en el HTML genérico de Flask/Werkzeug ("Internal Server Error").
# Esto NO reemplaza el manejo específico que ya tienen /api/chat y otras
# rutas (que arman mensajes más precisos): es la última línea de defensa
# para cualquier endpoint que no tenga su propio try/except, para que el
# frontend (leerJsonSeguro) siempre reciba JSON parseable.
@app.errorhandler(Exception)
def _manejar_excepcion_no_controlada(error):
    from werkzeug.exceptions import HTTPException

    if isinstance(error, HTTPException):
        # Errores HTTP legítimos (404, 405, etc.) mantienen su status,
        # pero devolvemos JSON en vez del HTML por defecto de Werkzeug.
        logger.warning(
            "HTTPException en %s %s: %s", request.method, request.path, error
        )
        return jsonify({
            "ok": False,
            "error": error.description or "No se pudo completar la operación.",
        }), error.code

    # Excepción real no controlada: log técnico completo (sin credenciales)
    # y respuesta JSON genérica y amigable para el usuario.
    logger.error(
        "EXCEPCIÓN NO CONTROLADA en %s %s: %s\n%s",
        request.method,
        request.path,
        error,
        traceback.format_exc(),
    )
    return jsonify({
        "ok": False,
        "error": "Ocurrió un problema al procesar la solicitud. Intentá nuevamente. "
                 "Si el problema continúa, avisá al administrador.",
    }), 500


# Request demasiado grande (ej. PDF > MAX_CONTENT_LENGTH): Werkzeug corta la
# petición en el parsing, antes de que el código de /api/chat pueda dar su
# propio mensaje. Sin este handler, esto también devuelve HTML.
@app.errorhandler(413)
def _manejar_request_demasiado_grande(error):
    return jsonify({
        "ok": False,
        "error": "El archivo adjunto es demasiado grande. El máximo permitido es 20 MB.",
    }), 413


DOCUMENTOS_DIR = BASE_DIR / "documentos"

NOTAS_FILE = BASE_DIR / "notas.json"
WORD_FILE = BASE_DIR / "documento_interno.docx"

# Planilla interna editable de Oficina IA.
EXCEL_FILE = BASE_DIR / "excel_interno.xlsx"

LIBROS_EXCEL = {
    "1": {
        "archivo": "excel_interno.xlsx",
        "r2_key": EXCEL_INTERNO_R2_KEY,
        "nombre": "Asegurados",
    },
    "2": {
        "archivo": "excel_flotas.xlsx",
        "r2_key": "excel/flotas.xlsx",
        "nombre": "Flotas",
    },
}

DOCUMENTOS_DIR.mkdir(
    exist_ok=True
)

CONFIG_FILE = BASE_DIR / "configuracion.json"
DB_FILE = BASE_DIR / "oficina.db"
ROLES_VALIDOS = {"admin", "usuario"}
USUARIO_ADMIN_PRINCIPAL = "admin"

# Accesos directos a plataformas de compañías. Se muestran solo los nombres.
CIAS_LINKS = [
    ("Self", "https://online.fedpat.com.ar/self/index.jsp"),
    ("ATM", "https://extranet.atmseguros.com.ar/ATM_COM_PROD/servlet/ar.com.glmsa.seguros.comercial.hlogin"),
    ("Rivadavia", "https://www.sistemas.segurosrivadavia.com/sistemas/login/login_intra_pas.php?u=P"),
    ("Triunfo", "https://www.triunfonet.com.ar/gauswebtriunfo/servlet/hlogon"),
    ("Prof", "https://pasnet.profseguros.seg.ar/Default.aspx"),
    ("Ags", "https://www.agsnet.com.ar/ingreprod.php"),
    ("San Cristobal", "https://productores.sancristobal.com.ar/"),
    ("Mercantil Andina", "https://servicios.mercantilandina.com.ar/sigmav3/"),
    ("EuroAmerica", "https://pas.euroamericaseguros.seg.ar/login"),
    ("Allianz", "https://auth.allianz.com.ar/login"),
]

MANUALES_DIR = BASE_DIR / "manuales_companias"
POLIZAS_DIR = BASE_DIR / "polizas"
MANUALES_DIR.mkdir(exist_ok=True)
POLIZAS_DIR.mkdir(exist_ok=True)

MANUALES_COMPANIAS = [
    "Mercantil Andina",
    "Federación Patronal",
    "ATM",
    "San Cristóbal",
    "Rivadavia",
    "EuroAmérica",
    "AgroSalta",
    "Triunfo",
    "PROF",
]

# Límites de recuperación de manuales. Una consulta con compañía identificada
# puede revisar todos los manuales de esa compañía; una consulta genérica se
# mantiene acotada para proteger memoria, tiempo y costo en Render.
# MANUALES_MAX_CANDIDATOS_CIA: 0 = sin tope (revisar todos los de esa compañía).
# Se baja el default de 0 a 10 para evitar quedarse sin memoria (SIGKILL) al
# descargar TODOS los manuales de una compañía en una sola consulta. Si hace
# falta cobertura completa y el plan de Render tiene RAM de sobra, se puede
# subir con la variable de entorno MANUALES_MAX_CANDIDATOS_CIA en Render.
MANUALES_MAX_CANDIDATOS_GENERAL = int(os.getenv("MANUALES_MAX_CANDIDATOS_GENERAL", "12"))
MANUALES_MAX_CANDIDATOS_CIA = int(os.getenv("MANUALES_MAX_CANDIDATOS_CIA", "10"))
MANUALES_MAX_ARCHIVOS_CON_CIA = int(os.getenv("MANUALES_MAX_ARCHIVOS_CON_CIA", "6"))
MANUALES_MAX_ARCHIVOS_GENERAL = int(os.getenv("MANUALES_MAX_ARCHIVOS_GENERAL", "3"))

def slug_manual_compania(nombre):
    equivalencias = {
        "Mercantil Andina": "mercantil_andina",
        "Federación Patronal": "federacion_patronal",
        "ATM": "atm",
        "San Cristóbal": "san_cristobal",
        "Rivadavia": "rivadavia",
        "EuroAmérica": "euroamerica",
        "AgroSalta": "agrosalta",
        "Triunfo": "triunfo",
        "PROF": "prof",
    }
    return equivalencias[nombre]

def manuales_companias():
    """
    Lista los manuales almacenados en Neon/R2 agrupados por compañía.
    Mantiene exactamente la estructura que espera la interfaz actual.
    """
    filas = listar_manuales()
    agrupados = {slug_manual_compania(nombre): [] for nombre in MANUALES_COMPANIAS}

    for fila in filas:
        r2_key = str(fila.get("r2_key") or "")
        partes = r2_key.split("/")
        if len(partes) < 3 or partes[0] != "manuales":
            continue
        slug = partes[1]
        if slug not in agrupados:
            continue

        fecha = fila.get("fecha_subida")
        if fecha:
            fecha_texto = fecha.strftime("%d/%m/%Y %H:%M")
        else:
            fecha_texto = ""

        agrupados[slug].append({
            "nombre": fila.get("nombre") or "manual.pdf",
            # La interfaz sigue llamando a este campo "archivo", pero ahora
            # contiene el r2_key privado, no una ruta local.
            "archivo": r2_key,
            "fecha": fecha_texto,
            "tamaño": round((fila.get("tamaño") or 0) / 1024, 1),
        })

    resultado = []
    for nombre in MANUALES_COMPANIAS:
        slug = slug_manual_compania(nombre)
        archivos = agrupados[slug]
        resultado.append({
            "nombre": nombre,
            "slug": slug,
            "cargado": bool(archivos),
            "cantidad": len(archivos),
            "archivos": archivos,
        })
    return resultado


# ==========================================================
# USUARIOS Y AUTENTICACIÓN
# ==========================================================

def conectar_db():
    conexion = sqlite3.connect(DB_FILE)
    conexion.row_factory = sqlite3.Row
    return conexion


def _usuarios_usar_pg():
    """True si hay DATABASE_URL (Neon): los usuarios se guardan ahí y
    sobreviven a los redeploys/reinicios de Render. Si no hay Neon
    configurada, se usa SQLite local como respaldo (solo para desarrollo)."""
    return bool(os.getenv("DATABASE_URL"))


# Los chats (conversaciones + mensajes) usan la misma regla que los usuarios:
# con Neon configurada viven en Postgres y sobreviven a los redeploys de
# Render; sin Neon, caen a SQLite local (solo development).
_chats_usar_pg = _usuarios_usar_pg

# La configuración global (nombre de oficina, colores, herramientas
# visibles) sigue la misma regla: Neon si está configurada, archivo JSON
# local como respaldo de desarrollo.
_config_usar_pg = _usuarios_usar_pg

# El documento interno (Word / "Notas") es texto plano: con Neon vive en
# Postgres; sin Neon cae al .docx local como respaldo de desarrollo.
_documento_interno_usar_pg = _usuarios_usar_pg


def inicializar_base_datos():
    with closing(conectar_db()) as db:
        db.execute("CREATE TABLE IF NOT EXISTS usuarios (id INTEGER PRIMARY KEY AUTOINCREMENT, usuario TEXT UNIQUE NOT NULL, password TEXT NOT NULL)")
        db.execute("""CREATE TABLE IF NOT EXISTS conversaciones (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario TEXT NOT NULL,
            titulo TEXT NOT NULL DEFAULT 'Nueva conversación',
            creado_en TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            actualizado_en TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
        )""")
        db.execute("""CREATE TABLE IF NOT EXISTS mensajes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversacion_id INTEGER NOT NULL,
            rol TEXT NOT NULL,
            contenido TEXT NOT NULL,
            creado_en TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(conversacion_id) REFERENCES conversaciones(id) ON DELETE CASCADE
        )""")
        db.execute("""CREATE TABLE IF NOT EXISTS flotas_activas (
            conversacion_id INTEGER PRIMARY KEY,
            estado TEXT NOT NULL DEFAULT 'nueva',
            libro_id TEXT NOT NULL DEFAULT '2',
            datos_generales TEXT NOT NULL DEFAULT '{}',
            vehiculos TEXT NOT NULL DEFAULT '[]',
            creado_en TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            actualizado_en TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(conversacion_id) REFERENCES conversaciones(id) ON DELETE CASCADE
        )""")
        db.execute("""CREATE TABLE IF NOT EXISTS metadatos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario TEXT NOT NULL,
            titulo TEXT NOT NULL,
            contenido TEXT NOT NULL,
            creado_en TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            actualizado_en TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )""")
        columnas = {fila[1] for fila in db.execute("PRAGMA table_info(usuarios)").fetchall()}
        if "email" not in columnas: db.execute("ALTER TABLE usuarios ADD COLUMN email TEXT NOT NULL DEFAULT ''")
        if "rol" not in columnas: db.execute("ALTER TABLE usuarios ADD COLUMN rol TEXT NOT NULL DEFAULT 'usuario'")
        if "protegido" not in columnas: db.execute("ALTER TABLE usuarios ADD COLUMN protegido INTEGER NOT NULL DEFAULT 0")
        # P2.6 / Tanda C — tipo de chat (flota|coti|alta|envios)
        cols_conv = {fila[1] for fila in db.execute("PRAGMA table_info(conversaciones)").fetchall()}
        if "tipo" not in cols_conv:
            db.execute("ALTER TABLE conversaciones ADD COLUMN tipo TEXT NOT NULL DEFAULT ''")
        # El bootstrap del admin en SQLite sólo importa cuando NO hay Neon
        # (desarrollo local). Con Neon configurada, el admin se crea/mantiene
        # en Postgres desde inicializar_postgres(), para no pisar contraseñas
        # cambiadas en la nube con el valor por defecto de este archivo local.
        if not _usuarios_usar_pg():
            admin = db.execute("SELECT id FROM usuarios WHERE usuario = ?", (USUARIO_ADMIN_PRINCIPAL,)).fetchone()
            if admin is None:
                db.execute("INSERT INTO usuarios (usuario,password,email,rol,protegido) VALUES (?,?,?,?,1)", ("admin", generate_password_hash("1234"), "", "admin"))
            else:
                db.execute("UPDATE usuarios SET rol='admin', protegido=1 WHERE usuario=?", (USUARIO_ADMIN_PRINCIPAL,))
        db.commit()
        pendientes_ops.asegurar_tabla(db)


def obtener_usuario(usuario):
    if _usuarios_usar_pg():
        try:
            return pg_obtener_usuario(usuario)
        except Exception as error:
            print("ERROR obtener_usuario PG:", error)
            return None
    with closing(conectar_db()) as db:
        return db.execute("SELECT id,usuario,password,email,rol,protegido FROM usuarios WHERE usuario=?", (usuario,)).fetchone()

def obtener_usuario_por_id(usuario_id):
    if _usuarios_usar_pg():
        try:
            return pg_obtener_usuario_por_id(usuario_id)
        except Exception as error:
            print("ERROR obtener_usuario_por_id PG:", error)
            return None
    with closing(conectar_db()) as db:
        return db.execute("SELECT id,usuario,password,email,rol,protegido FROM usuarios WHERE id=?", (usuario_id,)).fetchone()

def usuario_es_admin():
    u=obtener_usuario(session.get("usuario", "")) if session.get("usuario") else None
    return bool(u and u["rol"] == "admin")

def validar_email(email):
    return not email or bool(re.fullmatch(r"[^\s@]+@[^\s@]+\.[^\s@]+", email))

def requiere_admin(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        if "usuario" not in session: return redirect(url_for("login"))
        if not usuario_es_admin(): return ("Acceso no autorizado", 403)
        return func(*args, **kwargs)
    return wrapper

def _herramientas_legacy_a_lista(visibles=None, urls=None):
    """Migra la configuración vieja (diccionarios fijos) al modelo dinámico.

    Se conserva para que una oficina existente no pierda sus accesos al
    desplegar esta versión. A partir de ahora las herramientas se guardan como
    una lista libre: nombre + URL + visible.
    """
    visibles = visibles if isinstance(visibles, dict) else {}
    urls = urls if isinstance(urls, dict) else {}
    catalogo = [
        ("gmail", "Gmail", "https://mail.google.com/"),
        ("whatsapp", "WhatsApp", "https://web.whatsapp.com/"),
        ("datacar", "Datacar", "https://www.datacar.com.ar/"),
        ("nosis", "Nosis", "https://www.nosis.com/es"),
        ("chatgpt", "ChatGPT", "https://chatgpt.com/"),
        ("drive", "Drive", "https://drive.google.com/"),
        ("envios_ya", "Envíos Ya", ""),
    ]
    salida = []
    for clave, nombre, url_default in catalogo:
        url = str(urls.get(clave, url_default) or "").strip()
        if not url:
            continue
        salida.append({
            "id": clave,
            "nombre": nombre,
            "url": url,
            "visible": bool(visibles.get(clave, True)),
        })
    return salida


def cargar_configuracion():
    config = {
        "nombre_oficina": "Oficina Seguros",
        "notificaciones": True,
        "color_principal": "#122033",
        "color_acento": "#0d8b7c",
        "color_fondo": "#f7f9fb",
        "color_sidebar": "#ffffff",
        "color_botones": "#122033",
        "herramientas": [],
        "excel_visible": True,
    }
    try:
        datos = None
        if _config_usar_pg():
            try:
                datos = pg_obtener_configuracion()
            except Exception as error:
                print("ERROR cargar_configuracion PG:", error)
                datos = None
        elif CONFIG_FILE.exists():
            datos = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))

        if isinstance(datos, dict):
            config.update(datos)

            herramientas = config.get("herramientas")
            if not isinstance(herramientas, list):
                herramientas = _herramientas_legacy_a_lista(
                    config.get("herramientas_visibles"),
                    config.get("herramientas_urls"),
                )
            limpias = []
            vistos = set()
            for i, item in enumerate(herramientas):
                if not isinstance(item, dict):
                    continue
                nombre = str(item.get("nombre", "") or "").strip()
                url = str(item.get("url", "") or "").strip()
                ident = re.sub(r"[^a-z0-9_-]+", "-", str(item.get("id", "") or "").lower()).strip("-")
                if not ident:
                    ident = f"herramienta-{i+1}"
                base = ident
                n = 2
                while ident in vistos:
                    ident = f"{base}-{n}"; n += 1
                vistos.add(ident)
                if nombre and url:
                    limpias.append({"id": ident, "nombre": nombre, "url": url, "visible": bool(item.get("visible", True))})
            config["herramientas"] = limpias
            # Las claves viejas pueden seguir guardadas en Neon/JSON por
            # compatibilidad, pero la interfaz nueva ya no depende de ellas.
            config["excel_visible"] = bool(config.get("excel_visible", True))
    except Exception:
        pass
    return config

def contexto_usuario():
    u=obtener_usuario(session.get("usuario", "")) if session.get("usuario") else None
    config = cargar_configuracion()
    return {
        "usuario_rol": u["rol"] if u else None,
        "usuario_es_admin": bool(u and u["rol"] == "admin"),
        "config_global": config,
        "cias_links": CIAS_LINKS,
    }

app.context_processor(contexto_usuario)


# ==========================================================
# LOGIN
# ==========================================================

def requiere_login(func):

    @wraps(func)
    def wrapper(*args, **kwargs):

        if "usuario" not in session:

            return redirect(
                url_for("login")
            )

        return func(
            *args,
            **kwargs
        )

    return wrapper


# ==========================================================
# NOMBRES DE COMPAÑÍAS
# ==========================================================

def nombre_compania(nombre):
    return _nombre_compania_canonico(nombre)


app.jinja_env.globals["nombre_compania"] = nombre_compania


# ==========================================================
# OBTENER COMPAÑÍAS
# ==========================================================

def obtener_companias():
    """Devuelve únicamente las compañías que forman parte de la biblioteca de Manuales."""
    DOCUMENTOS_DIR.mkdir(parents=True, exist_ok=True)

    companias = [
        "atm",
        "mercantil_andina",
        "federacion_patronal",
        "san_cristobal",
        "rivadavia",
        "euroamerica",
        "agrosalta",
        "triunfo",
        "prof",
    ]

    for compania in companias:
        (DOCUMENTOS_DIR / compania).mkdir(parents=True, exist_ok=True)

    return sorted(
        companias,
        key=lambda x: nombre_compania(x).lower()
    )


# ==========================================================
# EXTRACCIÓN Y RETRIEVAL DE PDF
# ==========================================================

# Límites de memoria para Render.
MAX_PDF_PAGES_INDEX = 80
MAX_PDF_TEXT_CHARS_INDEX = 120_000
MAX_PDF_FILE_SIZE_BYTES = 15 * 1024 * 1024
MAX_PDF_PAGES_CHAT = 30
MAX_PDF_TEXT_CHARS_CHAT = 40_000
_PDF_CACHE = {}  # Se conserva por compatibilidad; no se utiliza para retener PDFs.

_STOPWORDS_ES = {
    "para", "como", "cual", "cuál", "que", "qué", "del", "las", "los",
    "una", "uno", "unos", "unas", "por", "con", "sin", "sobre", "entre",
    "desde", "hacia", "esta", "este", "estas", "estos", "tiene", "tienen",
    "debe", "deben", "puedo", "puede", "pueden", "quiero", "necesito",
    "donde", "dónde", "cuando", "cuándo", "hay", "son", "es", "el", "la",
    "y", "o", "a", "en", "un", "al", "se", "su", "sus", "mi", "mis",
    "me", "te", "lo", "le", "por", "del", "ya", "más", "mas"
}


def _normalizar_busqueda(texto):
    texto = str(texto or "").lower()
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", texto).strip()


def _tokens_busqueda(texto):
    normalizado = _normalizar_busqueda(texto)
    tokens = re.findall(r"[a-z0-9]+", normalizado)
    return [t for t in tokens if len(t) >= 3 and t not in _STOPWORDS_ES]


def _raiz_simple(token):
    """Pequeña normalización morfológica para español sin dependencias externas."""
    t = token.lower()
    for sufijo in (
        "amientos", "imientos", "aciones", "iciones", "amiento", "imiento",
        "mente", "ando", "iendo", "ados", "idas", "idos", "adas", "ados",
        "es", "os", "as", "o", "a", "e"
    ):
        if len(t) > len(sufijo) + 3 and t.endswith(sufijo):
            return t[:-len(sufijo)]
    return t


def extraer_paginas_pdf(ruta):
    """
    Extrae texto de un PDF de forma controlada para Render.
    Usa PyMuPDF en lugar de pypdf porque consume menos memoria en PDFs
    complejos. No conserva todos los PDFs procesados en una caché global.
    """
    ruta = Path(ruta)
    try:
        if not ruta.exists() or not ruta.is_file():
            return []
        if ruta.stat().st_size > MAX_PDF_FILE_SIZE_BYTES:
            print(f"PDF OMITIDO POR TAMAÑO: {ruta}")
            return []
    except OSError:
        return []

    paginas = []
    total_chars = 0

    try:
        documento = fitz.open(str(ruta))
        try:
            total_paginas = min(documento.page_count, MAX_PDF_PAGES_INDEX)

            for indice in range(total_paginas):
                if total_chars >= MAX_PDF_TEXT_CHARS_INDEX:
                    break

                try:
                    pagina = documento.load_page(indice)
                    contenido = pagina.get_text("text", sort=True) or ""
                    # Liberamos la referencia de página inmediatamente.
                    del pagina
                except Exception as error:
                    print(
                        f"ERROR EXTRAYENDO PÁGINA {indice + 1} DE PDF {ruta}: {error}"
                    )
                    continue

                contenido = re.sub(r"[ \t]+", " ", contenido)
                contenido = re.sub(r"\n{3,}", "\n\n", contenido).strip()

                if not contenido:
                    continue

                restante = MAX_PDF_TEXT_CHARS_INDEX - total_chars
                if len(contenido) > restante:
                    contenido = contenido[:restante]

                if contenido:
                    paginas.append({
                        "pagina": indice + 1,
                        "texto": contenido
                    })
                    total_chars += len(contenido)

        finally:
            documento.close()

        if not paginas:
            print(
                f"PDF SIN TEXTO EXTRAÍBLE: {ruta}. "
                "Si es un PDF escaneado, necesita OCR para poder consultarse."
            )

    except Exception as error:
        print("ERROR LEYENDO PDF CON PYMUPDF:", ruta, error)
        paginas = []

    return paginas

def extraer_texto_pdf(ruta):
    """Compatibilidad con las funciones existentes que necesitan texto completo."""
    return "\n\n".join(p["texto"] for p in extraer_paginas_pdf(ruta))


def extraer_texto_pdf_bytes(datos, max_paginas=25, max_chars=25_000):
    """Extrae texto desde bytes de un PDF (upload en memoria). P1.7 / Tanda B."""
    if not datos:
        return ""
    paginas = []
    total = 0
    try:
        documento = fitz.open(stream=datos, filetype="pdf")
        try:
            n = min(documento.page_count, max_paginas)
            for i in range(n):
                if total >= max_chars:
                    break
                try:
                    pagina = documento.load_page(i)
                    contenido = pagina.get_text("text", sort=True) or ""
                    del pagina
                except Exception:
                    continue
                contenido = re.sub(r"[ \t]+", " ", contenido)
                contenido = re.sub(r"\n{3,}", "\n\n", contenido).strip()
                if not contenido:
                    continue
                restante = max_chars - total
                if len(contenido) > restante:
                    contenido = contenido[:restante]
                paginas.append(contenido)
                total += len(contenido)
        finally:
            documento.close()
    except Exception as error:
        print("ERROR extraer_texto_pdf_bytes:", error)
        return ""
    return "\n\n".join(paginas)


def _proponer_ficha_desde_manual(texto, compania="", nombre_archivo=""):
    """
    Arma una ficha sugerida a partir del texto de un manual/póliza (P1.7).
    Heurística léxica: prioriza remolque/asistencia/cobertura/límites.
    No usa LLM (latencia/costo). El humano confirma antes de guardar.
    """
    texto = (texto or "").strip()
    if not texto or len(texto) < 40:
        return None

    compania = (compania or "").strip()
    nombre = (nombre_archivo or "").strip()
    if nombre.lower().endswith(".pdf"):
        nombre = nombre[:-4]

    lineas = [ln.strip() for ln in texto.splitlines() if ln.strip()]
    keywords = (
        "remolque", "grúa", "grua", "asistencia", "auxilio", "cobertura",
        "límite", "limite", "franquicia", "terceros", "all risk", "casco",
        "km", "kilómetro", "kilometro", "servicio", "exclusión", "exclusion",
        "suma asegurada", "deducible", "responsabilidad civil",
    )
    bullets = []
    vistos = set()
    for ln in lineas:
        low = ln.lower()
        if any(k in low for k in keywords):
            if len(ln) < 12 or len(ln) > 400:
                continue
            clave = re.sub(r"\s+", " ", low)[:120]
            if clave in vistos:
                continue
            vistos.add(clave)
            bullets.append("• " + ln)
            if len(bullets) >= 18:
                break

    if not bullets:
        for ln in lineas:
            if len(ln) < 20 or len(ln) > 300:
                continue
            if re.match(r"^(página|page|confidential|www\.|http)", ln, re.I):
                continue
            bullets.append("• " + ln)
            if len(bullets) >= 10:
                break

    if not bullets:
        return None

    if compania:
        titulo = f"{compania} — {nombre}" if nombre else f"Manual {compania}"
    else:
        titulo = nombre or "Ficha desde manual"
    titulo = titulo[:200]

    cuerpo_parts = []
    if compania:
        cuerpo_parts.append(f"Compañía: {compania}")
    if nombre:
        cuerpo_parts.append(f"Fuente: {nombre}.pdf")
    cuerpo_parts.append("")
    cuerpo_parts.append("Extracto operativo (revisar y completar):")
    cuerpo_parts.extend(bullets)
    cuerpo = "\n".join(cuerpo_parts)
    if len(cuerpo) > 12000:
        cuerpo = cuerpo[:12000] + "\n\n[Texto recortado por longitud]"

    return {"titulo": titulo, "contenido": cuerpo}


def _crear_chunks_paginas(paginas, chunk_chars=1400, overlap=220):
    """
    Divide el texto en fragmentos pequeños conservando página.
    Evita enviar un PDF completo al modelo cuando sólo una parte es relevante.
    """
    chunks = []

    for pagina in paginas:
        texto = pagina["texto"]
        if len(texto) <= chunk_chars:
            chunks.append({
                "pagina": pagina["pagina"],
                "texto": texto
            })
            continue

        inicio = 0
        while inicio < len(texto):
            fin = min(len(texto), inicio + chunk_chars)

            # Preferimos cortar cerca de un salto de párrafo o frase.
            if fin < len(texto):
                corte = max(
                    texto.rfind("\n", inicio + 700, fin),
                    texto.rfind(". ", inicio + 700, fin),
                    texto.rfind("; ", inicio + 700, fin)
                )
                if corte > inicio + 700:
                    fin = corte + 1

            fragmento = texto[inicio:fin].strip()
            if fragmento:
                chunks.append({
                    "pagina": pagina["pagina"],
                    "texto": fragmento
                })

            if fin >= len(texto):
                break

            inicio = max(inicio + 1, fin - overlap)

    return chunks


def _puntuar_chunk(consulta, chunk):
    """
    Ranking híbrido local:
    - coincidencia exacta de términos;
    - frecuencia;
    - frases de varias palabras;
    - coincidencias morfológicas simples.
    """
    consulta_norm = _normalizar_busqueda(consulta)
    texto_norm = _normalizar_busqueda(chunk["texto"])
    tokens = _tokens_busqueda(consulta)

    if not tokens or not texto_norm:
        return 0

    puntuacion = 0

    # Frase completa: una señal muy fuerte.
    if len(consulta_norm) >= 8 and consulta_norm in texto_norm:
        puntuacion += 30

    # Pares consecutivos de términos.
    for i in range(len(tokens) - 1):
        frase = f"{tokens[i]} {tokens[i+1]}"
        if frase in texto_norm:
            puntuacion += 10

    palabras_texto = set(re.findall(r"[a-z0-9]+", texto_norm))
    raices_texto = {_raiz_simple(x) for x in palabras_texto}

    for token in tokens:
        if token in palabras_texto:
            # La primera aparición es más útil que repetir una palabra 30 veces.
            puntuacion += min(8, 2 + texto_norm.count(token))
        elif _raiz_simple(token) in raices_texto:
            puntuacion += 3

    return puntuacion


def _manuales_r2_por_ruta(consulta="", max_manuales=None):
    """
    Prepara una cantidad acotada de manuales R2 por consulta.

    Si la consulta menciona una compañía, primero se restringe el universo a
    los manuales de esa compañía usando el mismo detector de aliases que usa
    servicios_ia._companias_mencionadas(). Recién después se aplica el orden
    por nombre. Si no hay compañía explícita, se conserva el filtro por nombre
    para no descargar todo R2 en cada request.

    Con compañía identificada se prioriza cobertura completa de ese universo;
    el límite opcional sólo se usa si se configura explícitamente y es mayor
    que cero. Esto aumenta la lectura de PDFs de una compañía, pero evita que
    un manual correcto quede fuera por el nombre de archivo y mantiene el
    universo de trabajo controlado frente a un barrido global.
    """
    mapa = {}
    try:
        manuales = listar_manuales()

        # Reutilizamos exactamente el criterio de aliases de servicios_ia.py.
        try:
            from servicios_ia import _companias_mencionadas
            companias_detectadas = _companias_mencionadas(consulta, [])
        except Exception as error:
            print("ERROR DETECTANDO COMPAÑIA PARA MANUALES:", error)
            companias_detectadas = set()

        slug_por_canon = {
            "mercantil andina": "mercantil_andina",
            "mercantilandina": "mercantil_andina",
            "federacion patronal": "federacion_patronal",
            "federacion": "federacion_patronal",
            "atm": "atm",
            "san cristobal": "san_cristobal",
            "sancristobal": "san_cristobal",
            "rivadavia": "rivadavia",
            "euroamerica": "euroamerica",
            "euro america": "euroamerica",
            "agrosalta": "agrosalta",
            "ags": "agrosalta",
            "triunfo": "triunfo",
            "prof": "prof",
        }

        # Algunos aliases del detector tienen una forma compacta distinta
        # (ej. "mercantilandina"). Se resuelven contra la misma compañía.
        slug_companias = set()
        for canon in companias_detectadas:
            canon_norm = _normalizar_busqueda(canon)
            slug = slug_por_canon.get(canon_norm)
            if slug:
                slug_companias.add(slug)
                continue
            compact = re.sub(r"[^a-z0-9]+", "", canon_norm)
            for nombre, slug_candidato in slug_por_canon.items():
                if compact == re.sub(r"[^a-z0-9]+", "", nombre):
                    slug_companias.add(slug_candidato)
                    break

        candidatos = []
        for fila in manuales:
            nombre = str(fila.get("nombre") or "")
            r2_key = str(fila.get("r2_key") or "")
            if not r2_key:
                continue

            partes = r2_key.split("/")
            slug = partes[1] if len(partes) > 1 else ""

            if slug_companias:
                # Con compañía explícita, los manuales de otras compañías no
                # compiten en absoluto por el contexto final.
                if slug not in slug_companias:
                    continue
                score = 0
            else:
                texto_nombre = _normalizar_busqueda(f"{nombre} {r2_key}")
                tokens = set(_tokens_busqueda(consulta))
                score = sum(1 for token in tokens if token in texto_nombre)

            candidatos.append((score, nombre, fila))

        candidatos.sort(key=lambda x: (x[0], x[1].lower()), reverse=True)

        if slug_companias:
            # Si la compañía está identificada, por defecto se revisan todos
            # sus manuales. Sólo un límite > 0 impuesto por configuración
            # reduce ese universo de forma explícita. 0 o None = sin tope.
            limite = max_manuales
            if limite is None:
                limite = MANUALES_MAX_CANDIDATOS_CIA
            seleccion = candidatos if not limite else candidatos[:limite]
        else:
            limite = max_manuales
            if limite is None:
                limite = MANUALES_MAX_CANDIDATOS_GENERAL
            seleccion = candidatos[:max(0, limite)]

        for score, _, fila in seleccion:
            r2_key = str(fila.get("r2_key") or "")
            try:
                path = descargar_pdf_temporal(r2_key)
            except Exception as error:
                print(f"ERROR PREPARANDO MANUAL R2 {r2_key}: {error}")
                continue

            mapa[str(path.resolve())] = fila

    except Exception as error:
        print("ERROR CONSULTANDO MANUALES R2:", error)

    return mapa


def _polizas_r2_por_ruta(consulta="", max_polizas=8):
    """
    Descarga a caché temporal una cantidad acotada de pólizas de R2,
    priorizadas por coincidencia de nombre con la consulta, para que
    buscar_en_documentos() pueda indexarlas igual que a los manuales.
    """
    mapa = {}
    try:
        polizas = pg_listar_polizas()
        tokens = set(_tokens_busqueda(consulta))

        candidatos = []
        for fila in polizas:
            nombre = str(fila.get("nombre") or "")
            r2_key = str(fila.get("r2_key") or "")
            if not r2_key:
                continue
            texto_nombre = _normalizar_busqueda(nombre)
            score = sum(1 for token in tokens if token in texto_nombre)
            candidatos.append((score, nombre, fila))

        candidatos.sort(key=lambda x: (x[0], x[1].lower()), reverse=True)
        limite = max_polizas if max_polizas else len(candidatos)
        seleccion = candidatos[:max(0, limite)]

        for score, _, fila in seleccion:
            r2_key = str(fila.get("r2_key") or "")
            try:
                path = descargar_pdf_temporal(r2_key)
            except Exception as error:
                print(f"ERROR PREPARANDO POLIZA R2 {r2_key}: {error}")
                continue
            mapa[str(path.resolve())] = fila

    except Exception as error:
        print("ERROR CONSULTANDO POLIZAS R2:", error)

    return mapa


def buscar_en_documentos(consulta, limite=16):
    """
    Recuperación por relevancia de PDFs.

    - Con compañía identificada, primero se acota R2 a esa compañía y se
      permite que compitan hasta MANUALES_MAX_ARCHIVOS_CON_CIA archivos.
    - Sin compañía identificada, se conserva un tope menor para no disparar
      descargas, memoria y costo en Render.
    - La cantidad total de fragmentos sigue limitada por ``limite`` y cada
      archivo conserva su máximo de 4/8 fragmentos según complejidad.
    """
    resultados = []
    tokens = _tokens_busqueda(consulta)

    if not tokens:
        return resultados

    try:
        from servicios_ia import _companias_mencionadas
        companias_detectadas = _companias_mencionadas(consulta, [])
    except Exception:
        companias_detectadas = set()

    r2_por_ruta = _manuales_r2_por_ruta(consulta)
    r2_por_ruta.update(_polizas_r2_por_ruta(consulta))

    archivos_locales = []
    if DOCUMENTOS_DIR.exists():
        archivos_locales.extend(
            p for p in DOCUMENTOS_DIR.rglob("*.pdf") if p.is_file()
        )

    archivos = archivos_locales + [Path(ruta) for ruta in r2_por_ruta]
    cantidad_archivos = len(archivos)

    for archivo in archivos:
        paginas = extraer_paginas_pdf(archivo)
        if not paginas:
            continue

        chunks = _crear_chunks_paginas(paginas)

        for chunk in chunks:
            puntuacion = _puntuar_chunk(consulta, chunk)
            if puntuacion <= 0:
                continue

            try:
                ruta_clave = str(archivo.resolve())

                if ruta_clave in r2_por_ruta:
                    fila = r2_por_ruta[ruta_clave]
                    r2_key = str(fila.get("r2_key") or "")

                    if r2_key.startswith("polizas/"):
                        nombre_archivo = fila.get("nombre") or archivo.name
                        compania = "Biblioteca de pólizas"
                        tipo = "poliza"
                        ruta_relativa = r2_key
                    else:
                        partes = r2_key.split("/")
                        slug = partes[1] if len(partes) > 1 else ""
                        compania = next(
                            (c for c in MANUALES_COMPANIAS
                             if slug_manual_compania(c) == slug),
                            "",
                        )
                        nombre_archivo = fila.get("nombre") or archivo.name
                        tipo = "manual"
                        ruta_relativa = r2_key

                else:
                    relativa = archivo.relative_to(DOCUMENTOS_DIR)
                    partes = relativa.parts
                    compania_slug = partes[0] if partes else ""
                    compania = nombre_compania(compania_slug)
                    nombre_archivo = archivo.name
                    tipo = "documento"
                    ruta_relativa = str(relativa)

            except Exception:
                compania = ""
                nombre_archivo = archivo.name
                tipo = "documento"
                ruta_relativa = archivo.name

            resultados.append({
                "archivo": nombre_archivo,
                "compania": compania,
                "ruta": ruta_relativa,
                "coincidencias": puntuacion,
                "texto": chunk["texto"],
                "pagina": chunk["pagina"],
                "tipo": tipo,
            })

    resultados.sort(
        key=lambda x: (
            x["coincidencias"],
            len(x.get("texto", ""))
        ),
        reverse=True
    )

    # La detección de compañía define cuántos archivos pueden competir.
    # El límite total de fragmentos sigue siendo ``limite``.
    max_archivos = (
        MANUALES_MAX_ARCHIVOS_CON_CIA
        if companias_detectadas
        else MANUALES_MAX_ARCHIVOS_GENERAL
    )

    seleccionados = []
    por_archivo = {}
    archivos_permitidos = []

    tokens_consulta = _tokens_busqueda(consulta)
    es_compleja = len(tokens_consulta) >= 8 or any(
        palabra in _normalizar_busqueda(consulta)
        for palabra in (
            "como", "cómo", "procedimiento", "documentacion",
            "documentación", "requisitos", "condiciones", "pasos",
            "explicame", "detalle", "completo",
        )
    )
    max_por_archivo = 8 if es_compleja else 4

    for resultado in resultados:
        clave = resultado["ruta"]
        if clave not in archivos_permitidos:
            if len(archivos_permitidos) >= max_archivos:
                continue
            archivos_permitidos.append(clave)

        cantidad = por_archivo.get(clave, 0)
        if cantidad >= max_por_archivo:
            continue

        seleccionados.append(resultado)
        por_archivo[clave] = cantidad + 1
        if len(seleccionados) >= limite:
            break

    print(
        f"RETRIEVAL PDF: consulta={consulta!r} "
        f"companias={sorted(companias_detectadas)} "
        f"archivos_procesados={cantidad_archivos} "
        f"archivos_seleccionados={len(archivos_permitidos)} "
        f"fragmentos={len(seleccionados)}"
    )

    return seleccionados



# ==========================================================
# LOGIN
# ==========================================================

@app.route(
    "/",
    methods=["GET", "POST"]
)
def login():

    if "usuario" in session:

        return redirect(
            url_for("documentos")
        )

    error = None

    if request.method == "POST":

        usuario = request.form.get(
            "usuario",
            ""
        ).strip()

        password = request.form.get(
            "password",
            ""
        )

        registro = obtener_usuario(usuario)
        if registro:
            try: valido = check_password_hash(registro["password"], password)
            except (ValueError, TypeError): valido = False
            if valido:
                session.clear()
                session["usuario"] = registro["usuario"]
                session["rol"] = registro["rol"]
                session["forzar_chat_nuevo"] = True
                return redirect(url_for("documentos"))
        error = "Usuario o contraseña incorrectos."

    return render_template(
        "login.html",
        error=error
    )


# ==========================================================
# INICIO
# ==========================================================

@app.route("/documentos")
@requiere_login
def documentos():
    # Consumir el flag una sola vez: solo fuerza un chat nuevo al iniciar sesión.
    forzar_chat_nuevo = session.pop("forzar_chat_nuevo", False)

    return render_template(
        "documentos.html",
        carpetas=obtener_companias(),
        usuario=session["usuario"],
        usuario_rol=session.get("rol", "usuario"),
        forzar_chat_nuevo=forzar_chat_nuevo
    )


# ==========================================================
# VER CARPETA
# ==========================================================

@app.route(
    "/carpeta/<path:carpeta>"
)
@requiere_login
def ver_carpeta(carpeta):
    # P0.5 — Anti path-traversal: el path resuelto debe quedar bajo DOCUMENTOS_DIR.
    try:
        carpeta_path = (DOCUMENTOS_DIR / carpeta).resolve()
        if not carpeta_path.is_relative_to(DOCUMENTOS_DIR.resolve()):
            return ("Acceso denegado", 403)
    except (OSError, ValueError):
        return ("Compañía no encontrada", 404)

    if not carpeta_path.exists() or not carpeta_path.is_dir():
        return ("Compañía no encontrada", 404)

    archivos = []
    for archivo in carpeta_path.rglob("*"):
        if archivo.is_file():
            try:
                rel = archivo.relative_to(carpeta_path)
            except ValueError:
                continue
            archivos.append({
                "nombre": archivo.name,
                "ruta": str(rel),
                "extension": archivo.suffix.lower(),
                "tamaño": round(archivo.stat().st_size / 1024, 1),
            })

    return render_template(
        "carpeta.html",
        carpeta=carpeta,
        nombre=nombre_compania(carpeta),
        archivos=archivos,
        usuario=session["usuario"],
        carpetas=obtener_companias()
    )


# ==========================================================
# ARCHIVO
# ==========================================================

@app.route(
    "/archivo/<path:carpeta>/<path:archivo>"
)
@requiere_login
def archivo(carpeta, archivo):
    # P0.5 — Anti path-traversal.
    try:
        carpeta_path = (DOCUMENTOS_DIR / carpeta).resolve()
        archivo_path = (carpeta_path / archivo).resolve()
        base = DOCUMENTOS_DIR.resolve()
        if not carpeta_path.is_relative_to(base) or not archivo_path.is_relative_to(base):
            return ("Acceso denegado", 403)
    except (OSError, ValueError):
        return ("Archivo no encontrado", 404)

    if not archivo_path.exists() or not archivo_path.is_file():
        return ("Archivo no encontrado", 404)

    # send_from_directory necesita el directorio base y el nombre relativo.
    rel = archivo_path.relative_to(carpeta_path)
    return send_from_directory(str(carpeta_path), str(rel))


# ==========================================================
# BUSCADOR
# ==========================================================

@app.route("/buscar")
@requiere_login
def buscar():

    return render_template(
        "buscar.html",
        usuario=session["usuario"],
        carpetas=obtener_companias()
    )


# ==========================================================
# BUSCADOR DE ARCHIVOS
# ==========================================================

@app.route(
    "/api/buscar"
)
@requiere_login
def api_buscar():

    consulta = request.args.get(
        "q",
        ""
    ).strip().lower()

    resultados = []

    if not consulta:

        return jsonify(
            resultados
        )

    for archivo in DOCUMENTOS_DIR.rglob("*"):

        if not archivo.is_file():

            continue

        if consulta not in archivo.name.lower():

            continue

        try:

            relativa = archivo.relative_to(
                DOCUMENTOS_DIR
            )

            partes = relativa.parts

            compania = (
                partes[0]
                if partes
                else ""
            )

            resultados.append({

                "nombre":
                    archivo.name,

                "compania":
                    nombre_compania(
                        compania
                    ),

                "ruta":
                    str(
                        relativa
                    ),

                "extension":
                    archivo.suffix.lower(),

                "tamaño":
                    round(
                        archivo.stat().st_size / 1024,
                        1
                    )

            })

        except Exception:

            pass

    return jsonify(
        resultados[:200]
    )


# ==========================================================
# BÚSQUEDA REAL EN CONTENIDO
# ==========================================================

@app.route(
    "/api/consultar_documentos",
    methods=["POST"]
)
@requiere_login
def consultar_documentos():

    data = request.get_json(
        silent=True
    )

    if not data:

        return jsonify({
            "ok": False,
            "resultados": []
        })

    consulta = data.get(
        "consulta",
        ""
    ).strip()

    resultados = buscar_en_documentos(
        consulta
    )

    respuesta = []

    for resultado in resultados[:5]:

        texto = resultado["texto"]

        palabras = re.findall(
            r"\w+",
            consulta.lower()
        )

        posiciones = []

        for palabra in palabras:

            posicion = texto.lower().find(
                palabra
            )

            if posicion >= 0:

                posiciones.append(
                    posicion
                )

        if posiciones:

            posicion = min(
                posiciones
            )

        else:

            posicion = 0

        inicio = max(
            0,
            posicion - 250
        )

        fin = min(
            len(texto),
            posicion + 750
        )

        fragmento = (
            texto[inicio:fin]
            .replace(
                "\n",
                " "
            )
        )

        respuesta.append({

            "archivo":
                resultado["archivo"],

            "compania":
                resultado["compania"],

            "coincidencias":
                resultado["coincidencias"],

            "fragmento":
                fragmento

        })

    return jsonify({

        "ok": True,

        "resultados":
            respuesta

    })


# ==========================================================
# ==========================================================
# EXCEL Y WORD INTERNOS
# ==========================================================

def _fila_vacia(fila):
    return not any(str(valor or "").strip() for valor in fila)

def _normalizar_matriz_excel(filas):
    """Normaliza la matriz sin eliminar filas intencionalmente creadas por el usuario."""
    if not isinstance(filas, list):
        raise ValueError("La matriz no es válida.")
    normalizadas = []
    for fila in filas[:500]:
        if not isinstance(fila, list):
            continue
        normalizadas.append(["" if valor is None else str(valor) for valor in fila[:30]])
    return normalizadas

def _limpiar_filas_excel(filas, conservar_vacias=False):
    """
    Limpieza explícita de filas vacías.
    Por defecto conserva filas vacías para que una fila nueva pueda existir
    y guardarse antes de que el usuario empiece a escribir.
    """
    normalizadas = _normalizar_matriz_excel(filas)
    if not normalizadas:
        return []
    if conservar_vacias:
        return normalizadas
    encabezado = normalizadas[0]
    cuerpo = [fila for fila in normalizadas[1:] if not _fila_vacia(fila)]
    return [encabezado] + cuerpo

def _limpiar_columnas_excel(filas):
    """Elimina sólo columnas completamente vacías cuando el usuario lo solicita."""
    filas = _normalizar_matriz_excel(filas)
    if not filas:
        return []
    max_cols = max((len(f) for f in filas), default=0)
    if not max_cols:
        return filas
    columnas_vivas = []
    for c in range(max_cols):
        if any(str(f[c] if c < len(f) else "").strip() for f in filas):
            columnas_vivas.append(c)
    if not columnas_vivas:
        return [[""]]
    return [[(fila[c] if c < len(fila) else "") for c in columnas_vivas] for fila in filas]


def _r2_excel_configurado():
    """Indica si R2 tiene las credenciales necesarias para persistir el Excel."""
    return all(
        os.getenv(nombre)
        for nombre in (
            "R2_ENDPOINT_URL",
            "R2_ACCESS_KEY_ID",
            "R2_SECRET_ACCESS_KEY",
            "R2_BUCKET_NAME",
        )
    )


def _crear_excel_inicial():
    wb = Workbook()
    ws = wb.active
    ws.title = "Datos"
    ws.append(["Dato", "Valor", "Observaciones"])
    wb.save(EXCEL_FILE)


def asegurar_excel_interno(libro_id="1"):
    """
    Mantiene una copia local de trabajo del Excel seleccionado, pero utiliza
    R2 como almacenamiento persistente cuando está configurado.
    """
    libro_id = str(libro_id or "1")
    if libro_id not in LIBROS_EXCEL:
        raise ValueError("Libro de Excel no válido.")
    libro = LIBROS_EXCEL[libro_id]
    archivo = BASE_DIR / libro["archivo"]
    r2_key = libro["r2_key"]

    if _r2_excel_configurado():
        try:
            descargado = descargar_excel_interno(archivo, r2_key)
            if descargado:
                return

            if not archivo.exists():
                wb = Workbook()
                ws = wb.active
                ws.title = "Datos"
                ws.append(["Dato", "Valor", "Observaciones"])
                wb.save(archivo)

            subir_excel_interno(archivo, r2_key)
            return
        except Exception as error:
            if archivo.exists():
                print("ADVERTENCIA EXCEL R2:", error)
                print(f"Se utilizará temporalmente la copia local del libro {libro_id}.")
                return
            raise RuntimeError(
                f"No se pudo recuperar el libro Excel {libro_id} desde Cloudflare R2 "
                "y tampoco existe una copia local."
            ) from error

    if not archivo.exists():
        wb = Workbook()
        ws = wb.active
        ws.title = "Datos"
        ws.append(["Dato", "Valor", "Observaciones"])
        wb.save(archivo)



def leer_excel_interno(libro_id="1"):
    libro_id = str(libro_id or "1")
    if libro_id not in LIBROS_EXCEL:
        raise ValueError("Libro de Excel no válido.")
    archivo = BASE_DIR / LIBROS_EXCEL[libro_id]["archivo"]
    asegurar_excel_interno(libro_id)
    wb = load_workbook(archivo, data_only=False)
    ws = wb.active
    filas = [["" if value is None else str(value) for value in row] for row in ws.iter_rows(values_only=True)]
    filas = _limpiar_filas_excel(filas, conservar_vacias=True)
    columnas = max([len(f) for f in filas], default=1)
    columnas = max(1, min(columnas, 30))
    filas = [f[:columnas] + [""] * (columnas - len(f)) for f in filas]
    return {"hoja": ws.title, "filas": filas, "columnas": columnas}


def guardar_matriz_excel(filas, nombre_hoja="Datos", libro_id="1"):
    libro_id = str(libro_id or "1")
    if libro_id not in LIBROS_EXCEL:
        raise ValueError("Libro de Excel no válido.")
    libro = LIBROS_EXCEL[libro_id]
    archivo = BASE_DIR / libro["archivo"]
    filas = _limpiar_filas_excel(filas, conservar_vacias=True)
    max_cols = max([len(f) for f in filas], default=1)
    max_cols = min(max_cols, 30)
    wb = Workbook()
    ws = wb.active
    ws.title = (nombre_hoja or "Datos")[:31]
    # Materializar TODAS las celdas de la matriz, incluso las vacías.
    # Esto permite que filas/columnas nuevas completamente vacías sobrevivan
    # al guardado y no dependan de si contienen datos.
    for r, fila in enumerate(filas, start=1):
        for c in range(1, max_cols + 1):
            valor = fila[c - 1] if c - 1 < len(fila) else ""
            ws.cell(row=r, column=c, value="" if valor is None else str(valor))
    for c in range(1, max_cols + 1):
        letra = get_column_letter(c)
        valores = [str(ws.cell(r, c).value or "") for r in range(1, min(ws.max_row, 30) + 1)]
        ancho = min(max([len(v) for v in valores] + [10]) + 2, 32)
        ws.column_dimensions[letra].width = ancho
    wb.save(archivo)

    # R2 es la persistencia permanente. Si la sincronización falla, elevamos
    # el error para que la API no informe falsamente que el guardado fue
    # exitoso y quede registrado en los logs de Render.
    if _r2_excel_configurado():
        try:
            subir_excel_interno(archivo, libro["r2_key"])
        except Exception as error:
            print(f"ERROR SINCRONIZANDO LIBRO {libro_id} CON R2:", error)
            raise

    # Un único punto de invalidación cubre limpiar/importar/agregar/editar: todas
    # las escrituras de la grilla pasan por guardar_matriz_excel. Sólo aplica al
    # libro principal que consulta la IA.
    if libro_id == "1":
        try:
            from servicios_ia import invalidar_cache_excel_interno
            invalidar_cache_excel_interno()
        except Exception as error:
            print("AVISO INVALIDANDO CACHE EXCEL IA:", error)

def asegurar_word_interno():
    if not WORD_FILE.exists():
        doc = Document()
        doc.add_paragraph("")
        doc.save(WORD_FILE)


def leer_word_interno():
    if _documento_interno_usar_pg():
        try:
            contenido = pg_obtener_documento_interno()
            return contenido if contenido is not None else ""
        except Exception as error:
            print("ERROR leer_word_interno PG:", error)
            return ""
    asegurar_word_interno()
    doc = Document(WORD_FILE)
    return "\n\n".join(p.text for p in doc.paragraphs)


def guardar_word_interno(contenido):
    if _documento_interno_usar_pg():
        pg_guardar_documento_interno(str(contenido or ""))
        return
    doc = Document()
    for linea in str(contenido or "").splitlines():
        doc.add_paragraph(linea)
    doc.save(WORD_FILE)


def _generar_docx_documento_interno():
    """Arma un .docx en memoria/disco con el contenido actual (Neon o
    local), para exportarlo. Devuelve la ruta del archivo generado."""
    if _documento_interno_usar_pg():
        contenido = leer_word_interno()
        doc = Document()
        for linea in contenido.splitlines():
            doc.add_paragraph(linea)
        doc.save(WORD_FILE)
    else:
        asegurar_word_interno()
    return WORD_FILE


@app.route("/notas")
@requiere_login
def notas():
    # Se conserva la URL para no romper marcadores antiguos; ahora muestra Excel + Word.
    return render_template("notas.html", usuario=session["usuario"], carpetas=obtener_companias())


@app.route("/api/excel", methods=["GET"])
@requiere_login
def api_excel():
    libro_id = request.args.get("libro_id", "1")
    try:
        return jsonify({"ok": True, **leer_excel_interno(libro_id)})
    except Exception as error:
        print("ERROR LEYENDO EXCEL INTERNO:", error)
        return jsonify({"ok": False, "error": "No se pudo leer la planilla."}), 500


@app.route("/api/excel", methods=["POST"])
@requiere_login
def api_excel_guardar():
    data = request.get_json(silent=True) or {}
    libro_id = data.get("libro_id", request.args.get("libro_id", "1"))
    try:
        guardar_matriz_excel(data.get("filas", []), data.get("hoja", "Datos"), libro_id)
        return jsonify({"ok": True, **leer_excel_interno(libro_id)})
    except Exception as error:
        print("ERROR GUARDANDO EXCEL INTERNO:", error)
        return jsonify({"ok": False, "error": "No se pudo guardar la planilla."}), 500


@app.route("/api/excel/limpiar", methods=["POST"])
@requiere_login
def api_excel_limpiar():
    data = request.get_json(silent=True) or {}
    libro_id = data.get("libro_id", request.args.get("libro_id", "1"))
    try:
        datos = leer_excel_interno(libro_id)
        filas_limpias = _limpiar_filas_excel(datos["filas"], conservar_vacias=False)
        guardar_matriz_excel(filas_limpias, datos["hoja"], libro_id)
        return jsonify({"ok": True, **leer_excel_interno(libro_id)})
    except Exception as error:
        print("ERROR LIMPIANDO EXCEL:", error)
        return jsonify({"ok": False, "error": "No se pudieron eliminar las filas vacías."}), 500


@app.route("/api/excel/limpiar-columnas", methods=["POST"])
@requiere_login
def api_excel_limpiar_columnas():
    data = request.get_json(silent=True) or {}
    libro_id = data.get("libro_id", request.args.get("libro_id", "1"))
    try:
        datos = leer_excel_interno(libro_id)
        filas_limpias = _limpiar_columnas_excel(datos["filas"])
        guardar_matriz_excel(filas_limpias, datos["hoja"], libro_id)
        return jsonify({"ok": True, **leer_excel_interno(libro_id)})
    except Exception:
        return jsonify({"ok": False, "error": "No se pudieron eliminar las columnas vacías."}), 500


@app.route("/api/excel/importar", methods=["POST"])
@requiere_login
def api_excel_importar():
    archivo = request.files.get("archivo")
    if not archivo or not archivo.filename:
        return jsonify({"ok": False, "error": "No se recibió ningún archivo."}), 400
    nombre = secure_filename(archivo.filename)
    if not nombre.lower().endswith((".xlsx", ".xlsm")):
        return jsonify({"ok": False, "error": "Usá un archivo Excel .xlsx o .xlsm."}), 400
    libro_id = request.form.get("libro_id", request.args.get("libro_id", "1"))
    if str(libro_id) not in LIBROS_EXCEL:
        return jsonify({"ok": False, "error": "Libro de Excel no válido."}), 400
    archivo_libro = BASE_DIR / LIBROS_EXCEL[str(libro_id)]["archivo"]
    temporal = archivo_libro.with_suffix(".upload.xlsx")
    try:
        archivo.save(temporal)
        wb = load_workbook(temporal, data_only=False)
        if not wb.sheetnames:
            raise ValueError("El Excel no contiene hojas.")
        ws = wb[wb.sheetnames[0]]
        filas = [["" if value is None else str(value) for value in row] for row in ws.iter_rows(values_only=True)]
        guardar_matriz_excel(filas, ws.title, libro_id)
        return jsonify({"ok": True, **leer_excel_interno(libro_id)})
    except Exception as error:
        print("ERROR IMPORTANDO EXCEL:", error)
        return jsonify({"ok": False, "error": "No se pudo importar el Excel."}), 400
    finally:
        try: temporal.unlink(missing_ok=True)
        except Exception: pass


@app.route("/excel/exportar")
@requiere_login
def excel_exportar():
    libro_id = request.args.get("libro_id", "1")
    if str(libro_id) not in LIBROS_EXCEL:
        return jsonify({"ok": False, "error": "Libro de Excel no válido."}), 400
    archivo = BASE_DIR / LIBROS_EXCEL[str(libro_id)]["archivo"]
    asegurar_excel_interno(libro_id)
    return send_from_directory(
        archivo.parent,
        archivo.name,
        as_attachment=True,
        download_name="OficinaIA.xlsx" if str(libro_id) == "1" else archivo.name,
    )


@app.route("/api/word", methods=["GET"])
@requiere_login
def api_word():
    try:
        return jsonify({"ok": True, "contenido": leer_word_interno()})
    except Exception as error:
        print("ERROR LEYENDO WORD:", error)
        return jsonify({"ok": False, "error": "No se pudo leer el documento."}), 500


@app.route("/api/word", methods=["POST"])
@requiere_login
def api_word_guardar():
    data = request.get_json(silent=True) or {}
    try:
        guardar_word_interno(data.get("contenido", ""))
        return jsonify({"ok": True})
    except Exception as error:
        print("ERROR GUARDANDO WORD:", error)
        return jsonify({"ok": False, "error": "No se pudo guardar el documento."}), 500


@app.route("/word/exportar")
@requiere_login
def word_exportar():
    _generar_docx_documento_interno()
    return send_from_directory(WORD_FILE.parent, WORD_FILE.name, as_attachment=True, download_name="OficinaIA.docx")




# ==========================================================
# METADATOS — FICHAS DE TEXTO LIBRE
# Preferencia: Neon/Postgres (persistente). Fallback: SQLite.
# ==========================================================

def _metadatos_usar_pg():
    """True si DATABASE_URL está configurada (producción Render/Neon)."""
    return bool(os.getenv("DATABASE_URL"))


@app.route("/api/metadatos", methods=["GET"])
@requiere_login
def listar_metadatos():
    if _metadatos_usar_pg():
        try:
            from database_pg import listar_metadatos as listar_pg
            filas = listar_pg()
            # La UI solo necesita id, titulo, actualizado_en en el listado.
            resumen = [
                {
                    "id": f["id"],
                    "titulo": f.get("titulo"),
                    "actualizado_en": f.get("actualizado_en"),
                }
                for f in filas
            ]
            return jsonify({"ok": True, "metadatos": resumen})
        except Exception as error:
            print("ERROR listar_metadatos PG:", error)
            return jsonify({"ok": False, "error": "No se pudieron listar los metadatos."}), 500

    with closing(conectar_db()) as db:
        rows = db.execute(
            "SELECT id, titulo, actualizado_en FROM metadatos ORDER BY actualizado_en DESC, id DESC"
        ).fetchall()
        return jsonify({"ok": True, "metadatos": [dict(row) for row in rows]})


@app.route("/api/metadatos/<int:metadato_id>", methods=["GET"])
@requiere_login
def obtener_metadato(metadato_id):
    if _metadatos_usar_pg():
        try:
            from database_pg import obtener_metadato as obtener_pg
            fila = obtener_pg(metadato_id)
            if not fila:
                return jsonify({"ok": False, "error": "Ficha no encontrada."}), 404
            return jsonify({"ok": True, "metadato": fila})
        except Exception as error:
            print("ERROR obtener_metadato PG:", error)
            return jsonify({"ok": False, "error": "No se pudo leer la ficha."}), 500

    with closing(conectar_db()) as db:
        row = db.execute(
            "SELECT id, titulo, contenido, creado_en, actualizado_en, usuario "
            "FROM metadatos WHERE id=?",
            (metadato_id,),
        ).fetchone()
        if not row:
            return jsonify({"ok": False, "error": "Ficha no encontrada."}), 404
        return jsonify({"ok": True, "metadato": dict(row)})


@app.route("/api/metadatos", methods=["POST"])
@requiere_login
def crear_metadato():
    data = request.get_json(silent=True) or {}
    titulo = str(data.get("titulo", "")).strip()
    contenido = str(data.get("contenido", "") or "")
    if not titulo:
        return jsonify({"ok": False, "error": "El título es obligatorio."}), 400
    if len(titulo) > 200:
        titulo = titulo[:200]

    if _metadatos_usar_pg():
        try:
            from database_pg import crear_metadato as crear_pg
            fila = crear_pg(session["usuario"], titulo, contenido)
            return jsonify({"ok": True, "metadato": fila})
        except Exception as error:
            print("ERROR crear_metadato PG:", error)
            return jsonify({"ok": False, "error": "No se pudo guardar la ficha."}), 500

    with closing(conectar_db()) as db:
        cur = db.execute(
            "INSERT INTO metadatos (usuario,titulo,contenido) VALUES (?,?,?)",
            (session["usuario"], titulo, contenido),
        )
        db.commit()
        row = db.execute(
            "SELECT id,titulo,contenido,creado_en,actualizado_en,usuario "
            "FROM metadatos WHERE id=?",
            (cur.lastrowid,),
        ).fetchone()
        return jsonify({"ok": True, "metadato": dict(row)})


@app.route("/api/metadatos/<int:metadato_id>", methods=["PUT"])
@requiere_login
def editar_metadato(metadato_id):
    data = request.get_json(silent=True) or {}
    titulo = str(data.get("titulo", "")).strip()
    contenido = str(data.get("contenido", "") or "")
    if not titulo:
        return jsonify({"ok": False, "error": "El título es obligatorio."}), 400
    if len(titulo) > 200:
        titulo = titulo[:200]

    if _metadatos_usar_pg():
        try:
            from database_pg import actualizar_metadato as actualizar_pg
            fila = actualizar_pg(metadato_id, titulo, contenido)
            if not fila:
                return jsonify({"ok": False, "error": "Ficha no encontrada."}), 404
            return jsonify({"ok": True, "metadato": fila})
        except Exception as error:
            print("ERROR editar_metadato PG:", error)
            return jsonify({"ok": False, "error": "No se pudo actualizar la ficha."}), 500

    with closing(conectar_db()) as db:
        row = db.execute("SELECT id FROM metadatos WHERE id=?", (metadato_id,)).fetchone()
        if not row:
            return jsonify({"ok": False, "error": "Ficha no encontrada."}), 404
        db.execute(
            "UPDATE metadatos SET titulo=?, contenido=?, actualizado_en=CURRENT_TIMESTAMP WHERE id=?",
            (titulo, contenido, metadato_id),
        )
        db.commit()
        row = db.execute(
            "SELECT id,titulo,contenido,creado_en,actualizado_en,usuario "
            "FROM metadatos WHERE id=?",
            (metadato_id,),
        ).fetchone()
        return jsonify({"ok": True, "metadato": dict(row)})


@app.route("/api/metadatos/<int:metadato_id>", methods=["DELETE"])
@requiere_login
def eliminar_metadato(metadato_id):
    if _metadatos_usar_pg():
        try:
            from database_pg import eliminar_metadato as eliminar_pg
            ok = eliminar_pg(metadato_id)
            if not ok:
                return jsonify({"ok": False, "error": "Ficha no encontrada."}), 404
            return jsonify({"ok": True})
        except Exception as error:
            print("ERROR eliminar_metadato PG:", error)
            return jsonify({"ok": False, "error": "No se pudo eliminar la ficha."}), 500

    with closing(conectar_db()) as db:
        row = db.execute("SELECT id FROM metadatos WHERE id=?", (metadato_id,)).fetchone()
        if not row:
            return jsonify({"ok": False, "error": "Ficha no encontrada."}), 404
        db.execute("DELETE FROM metadatos WHERE id=?", (metadato_id,))
        db.commit()
    return jsonify({"ok": True})

# ==========================================================
# CONVERSACIONES PERSISTENTES
# ==========================================================

def _crear_conversacion(usuario, titulo="Nueva conversación"):
    titulo = (titulo or "Nueva conversación")[:100] or "Nueva conversación"
    if _chats_usar_pg():
        try:
            return pg_crear_conversacion(usuario, titulo)
        except Exception as error:
            print("ERROR _crear_conversacion PG:", error)
            raise
    with closing(conectar_db()) as db:
        cur = db.execute(
            "INSERT INTO conversaciones (usuario,titulo) VALUES (?,?)",
            (usuario, titulo)
        )
        db.commit()
        return cur.lastrowid


def _validar_chat(chat_id, usuario):
    """True si la conversación existe y pertenece al usuario."""
    if _chats_usar_pg():
        try:
            return pg_validar_chat(chat_id, usuario)
        except Exception as error:
            print("ERROR _validar_chat PG:", error)
            return False
    with closing(conectar_db()) as db:
        row = db.execute(
            "SELECT id FROM conversaciones WHERE id=? AND usuario=?",
            (chat_id, usuario)
        ).fetchone()
        return row is not None


def _guardar_mensaje(chat_id, rol, contenido):
    """Inserta un mensaje y actualiza el timestamp de la conversación.
    P1.0b — Si Postgres falla, se propaga el error (no silenciar).
    """
    if _chats_usar_pg():
        try:
            pg_agregar_mensaje(chat_id, rol, contenido)
            return
        except Exception as error:
            print("ERROR _guardar_mensaje PG:", error)
            raise  # no tragar: el chat debe enterarse
    with closing(conectar_db()) as db:
        db.execute(
            "INSERT INTO mensajes (conversacion_id,rol,contenido) VALUES (?,?,?)",
            (chat_id, rol, contenido)
        )
        db.execute(
            "UPDATE conversaciones SET actualizado_en=CURRENT_TIMESTAMP WHERE id=?",
            (chat_id,)
        )
        db.commit()


def _historial_desde_db(chat_id, usuario, limite=10):
    """P1.3 — Historial para Gemini desde DB (no confiar en el cliente).
    Devuelve los últimos `limite` mensajes previos al turno actual.
    """
    if not chat_id:
        return []
    if _chats_usar_pg():
        try:
            from database_pg import listar_mensajes_historial
            return listar_mensajes_historial(chat_id, usuario, limite=limite)
        except Exception as error:
            print("ERROR _historial_desde_db PG:", error)
            return []
    with closing(conectar_db()) as db:
        if not db.execute(
            "SELECT id FROM conversaciones WHERE id=? AND usuario=?",
            (chat_id, usuario),
        ).fetchone():
            return []
        rows = db.execute(
            """
            SELECT rol, contenido FROM mensajes
            WHERE conversacion_id=?
            ORDER BY id DESC LIMIT ?
            """,
            (chat_id, limite),
        ).fetchall()
        out = [
            {"rol": r["rol"], "contenido": r["contenido"]}
            for r in reversed(list(rows))
            if r["rol"] in ("user", "assistant") and str(r["contenido"] or "").strip()
        ]
        return out


def _detectar_tipo_chat(mensaje):
    """Tipo operativo a partir del mensaje (comandos slash y señales claras)."""
    t = (mensaje or "").strip().lower()
    if t.startswith("/flota") or t.startswith("flota "):
        return "flota"
    if t.startswith("/coti"):
        return "coti"
    if t.startswith("/guardar") or t.startswith("/alta"):
        return "alta"
    if t.startswith("/envios") or t.startswith("/envíos"):
        return "envios"
    if "whatsapp" in t[:80]:
        return "whatsapp"
    return ""


def _asignar_tipo_chat(chat_id, usuario, mensaje):
    """Persiste tipo solo si el chat aún no tiene uno (P2.6)."""
    tipo = _detectar_tipo_chat(mensaje)
    if not tipo or not chat_id:
        return
    if _chats_usar_pg():
        try:
            from database_pg import actualizar_tipo_chat
            actualizar_tipo_chat(chat_id, usuario, tipo)
        except Exception as error:
            print("ERROR _asignar_tipo_chat PG:", error)
        return
    with closing(conectar_db()) as db:
        db.execute(
            """
            UPDATE conversaciones SET tipo=?
            WHERE id=? AND usuario=? AND (tipo IS NULL OR tipo='')
            """,
            (tipo, chat_id, usuario),
        )
        db.commit()


def _generar_titulo_chat(mensaje):
    """Título corto y útil a partir de la primera consulta. Determinístico, sin IA."""
    texto = " ".join(str(mensaje or "").split())
    if not texto:
        return "Nueva conversación"

    if texto.startswith("/"):
        partes = texto.split(None, 1)
        cmd = partes[0]
        resto = partes[1].strip() if len(partes) > 1 else ""
        if resto:
            corto = resto[:42] + ("…" if len(resto) > 42 else "")
            return f"{cmd} — {corto}"[:90]
        return cmd[:90]

    t = texto
    t = re.sub(r"^[¿?¡!\s]+", "", t)
    t = re.sub(r"[¿?¡!]+$", "", t).strip()
    # Quitar muletillas de pregunta frecuentes
    t = re.sub(
        r"^(?:cu[aá]ntos?|cu[aá]ntas?|qu[eé]|c[oó]mo|cu[aá]ndo|d[oó]nde|por\s+qu[eé]|para\s+qu[eé]|"
        r"me\s+pod[eé]s|pod[eé]s|necesit[oa]|quiero|haceme|armame|decime|mostrame|"
        r"busc[aá]|listame|ten[eé]s|dame|pasame|informame)\b[\s,:]*",
        "",
        t,
        flags=re.IGNORECASE,
    ).strip()
    t = re.sub(r"^(?:de\s+la|de\s+los|de\s+las|del|de|el|la|los|las|un|una)\s+", "", t, flags=re.IGNORECASE).strip()

    if not t:
        t = texto

    # Si aparece una compañía conocida, priorizarla al frente
    companias = (
        "ATM", "Sancor", "San Cor", "Mercantil Andina", "Mercantil", "Rivadavia",
        "Federación Patronal", "Federacion Patronal", "La Segunda", "Sura", "Allianz",
        "Mapfre", "Zurich", "Experta", "Provincia", "Triunfo", "Nación", "Nacion",
        "HDI", "Chubb", "SMG", "Galeno", "Prevención", "Prevencion",
    )
    encontrada = None
    lower = t.lower()
    for cia in sorted(companias, key=len, reverse=True):
        if cia.lower() in lower:
            encontrada = cia
            break
    if encontrada:
        resto = re.sub(re.escape(encontrada), " ", t, count=1, flags=re.IGNORECASE)
        # Limpiar verbos/conectores sobrantes del enunciado interrogativo
        resto = re.sub(
            r"\b(tiene|tienen|hay|son|es|para|sobre|con|del|de\s+la|de\s+los|de\s+las|de)\b",
            " ",
            resto,
            flags=re.IGNORECASE,
        )
        resto = re.sub(r"\s+", " ", resto).strip(" -—,.:;")
        if resto:
            if resto and resto[0].islower():
                resto = resto[0].upper() + resto[1:]
            if len(resto) > 40:
                corte = resto[:40].rsplit(" ", 1)[0] or resto[:40]
                resto = corte + "…"
            t = f"{encontrada} — {resto}"
        else:
            t = encontrada
    else:
        if t and t[0].islower():
            t = t[0].upper() + t[1:]
        if len(t) > 56:
            corte = t[:53].rsplit(" ", 1)[0] or t[:53]
            t = corte + "…"

    return (t or "Nueva conversación")[:90]


def _obtener_titulo_chat(chat_id, usuario):
    if _chats_usar_pg():
        try:
            from database_pg import obtener_titulo_chat as pg_obtener_titulo
            return pg_obtener_titulo(chat_id, usuario)
        except Exception as error:
            print("ERROR _obtener_titulo_chat PG:", error)
            return None
    with closing(conectar_db()) as db:
        row = db.execute(
            "SELECT titulo FROM conversaciones WHERE id=? AND usuario=?",
            (chat_id, usuario),
        ).fetchone()
        return row["titulo"] if row else None


def _actualizar_titulo_chat(chat_id, usuario, titulo):
    titulo = (titulo or "").strip()[:100] or "Nueva conversación"
    if _chats_usar_pg():
        try:
            from database_pg import actualizar_titulo_chat as pg_actualizar_titulo
            return pg_actualizar_titulo(chat_id, usuario, titulo)
        except Exception as error:
            print("ERROR _actualizar_titulo_chat PG:", error)
            return False
    with closing(conectar_db()) as db:
        cur = db.execute(
            "UPDATE conversaciones SET titulo=?, actualizado_en=CURRENT_TIMESTAMP WHERE id=? AND usuario=?",
            (titulo, chat_id, usuario),
        )
        db.commit()
        return cur.rowcount > 0


def _auto_titulo_si_corresponde(chat_id, usuario, mensaje):
    """Si el chat sigue con el título por defecto, lo reemplaza con uno derivado del mensaje."""
    actual = _obtener_titulo_chat(chat_id, usuario)
    if actual is None:
        return None
    if str(actual).strip().lower() not in {"nueva conversación", "nueva conversacion", ""}:
        return actual
    nuevo = _generar_titulo_chat(mensaje)
    if nuevo and nuevo != actual:
        _actualizar_titulo_chat(chat_id, usuario, nuevo)
        return nuevo
    return actual


_flota_usar_pg = _usuarios_usar_pg


def _flota_obtener(chat_id):
    """Devuelve el estado de la flota activa de esta conversación, o None.
    P1.4 — Si PG falla, se propaga (no evaporar flota en silencio).
    """
    if _flota_usar_pg():
        try:
            return pg_obtener_flota_activa(chat_id)
        except Exception as error:
            print("ERROR _flota_obtener PG:", error)
            raise
    with closing(conectar_db()) as db:
        fila = db.execute(
            "SELECT conversacion_id, estado, libro_id, datos_generales, vehiculos "
            "FROM flotas_activas WHERE conversacion_id=?",
            (chat_id,),
        ).fetchone()
        if not fila:
            return None
        return {
            "conversacion_id": fila[0],
            "estado": fila[1],
            "libro_id": fila[2],
            "datos_generales": json.loads(fila[3] or "{}"),
            "vehiculos": json.loads(fila[4] or "[]"),
        }


def _flota_guardar(chat_id, estado, libro_id, datos_generales, vehiculos):
    """Persiste (crea o actualiza) el estado de trabajo de una flota.
    P1.4 — Si PG falla, se propaga el error al chat.
    """
    if _flota_usar_pg():
        try:
            pg_guardar_flota_activa(chat_id, estado, libro_id, datos_generales, vehiculos)
            return
        except Exception as error:
            print("ERROR _flota_guardar PG:", error)
            raise
    with closing(conectar_db()) as db:
        db.execute(
            """
            INSERT INTO flotas_activas (conversacion_id, estado, libro_id, datos_generales, vehiculos, actualizado_en)
            VALUES (?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
            ON CONFLICT(conversacion_id) DO UPDATE SET
                estado=excluded.estado,
                libro_id=excluded.libro_id,
                datos_generales=excluded.datos_generales,
                vehiculos=excluded.vehiculos,
                actualizado_en=CURRENT_TIMESTAMP
            """,
            (
                chat_id,
                estado,
                libro_id,
                json.dumps(datos_generales or {}, ensure_ascii=False),
                json.dumps(vehiculos or [], ensure_ascii=False),
            ),
        )
        db.commit()


def _flota_borrar(chat_id):
    """Cierra/limpia el contexto de flota activa (al finalizar la tarea)."""
    if _flota_usar_pg():
        try:
            pg_borrar_flota_activa(chat_id)
            return
        except Exception as error:
            print("ERROR _flota_borrar PG:", error)
            raise
    with closing(conectar_db()) as db:
        db.execute("DELETE FROM flotas_activas WHERE conversacion_id=?", (chat_id,))
        db.commit()


@app.route("/api/chats", methods=["GET"])
@requiere_login
def listar_chats():
    if _chats_usar_pg():
        try:
            chats = pg_listar_chats(session["usuario"])
            return jsonify({"ok": True, "chats": chats})
        except Exception as error:
            print("ERROR listar_chats PG:", error)
            return jsonify({"ok": False, "error": "No se pudieron cargar las conversaciones."}), 500
    with closing(conectar_db()) as db:
        rows = db.execute(
            "SELECT id,titulo,COALESCE(tipo,'') AS tipo,creado_en,actualizado_en FROM conversaciones WHERE usuario=? ORDER BY actualizado_en DESC, id DESC",
            (session["usuario"],)
        ).fetchall()
        return jsonify({"ok": True, "chats":[dict(r) for r in rows]})

@app.route("/api/chats", methods=["POST"])
@requiere_login
def crear_chat():
    data=request.get_json(silent=True) or {}
    titulo=str(data.get("titulo","Nueva conversación")).strip()[:100] or "Nueva conversación"
    try:
        cid=_crear_conversacion(session["usuario"], titulo)
    except Exception:
        return jsonify({"ok": False, "error": "No se pudo crear la conversación."}), 500
    return jsonify({"ok":True,"id":cid,"titulo":titulo})

@app.route("/api/chats/<int:chat_id>", methods=["GET"])
@requiere_login
def obtener_chat(chat_id):
    if _chats_usar_pg():
        try:
            chat, mensajes = pg_obtener_chat_con_mensajes(chat_id, session["usuario"])
        except Exception as error:
            print("ERROR obtener_chat PG:", error)
            return jsonify({"ok": False, "error": "No se pudo cargar la conversación."}), 500
        if not chat:
            return jsonify({"ok": False, "error": "Conversación no encontrada."}), 404
        return jsonify({"ok": True, "chat": chat, "mensajes": mensajes})
    with closing(conectar_db()) as db:
        chat=db.execute("SELECT id,titulo FROM conversaciones WHERE id=? AND usuario=?",(chat_id,session["usuario"])).fetchone()
        if not chat: return jsonify({"ok":False,"error":"Conversación no encontrada."}),404
        mensajes=db.execute("SELECT id,rol,contenido,creado_en FROM mensajes WHERE conversacion_id=? ORDER BY id",(chat_id,)).fetchall()
        return jsonify({"ok":True,"chat":dict(chat),"mensajes":[dict(x) for x in mensajes]})

@app.route("/api/chats/<int:chat_id>", methods=["DELETE"])
@requiere_login
def eliminar_chat(chat_id):
    if _chats_usar_pg():
        try:
            borrado = pg_eliminar_chat(chat_id, session["usuario"])
        except Exception as error:
            print("ERROR eliminar_chat PG:", error)
            return jsonify({"ok": False, "error": "No se pudo eliminar la conversación."}), 500
        if not borrado:
            return jsonify({"ok": False, "error": "Conversación no encontrada."}), 404
        return jsonify({"ok": True})
    with closing(conectar_db()) as db:
        row=db.execute("SELECT id FROM conversaciones WHERE id=? AND usuario=?",(chat_id,session["usuario"])).fetchone()
        if not row: return jsonify({"ok":False,"error":"Conversación no encontrada."}),404
        db.execute("DELETE FROM mensajes WHERE conversacion_id=?",(chat_id,))
        db.execute("DELETE FROM conversaciones WHERE id=?",(chat_id,))
        db.commit()
    return jsonify({"ok":True})


@app.route("/api/chats/<int:chat_id>", methods=["PATCH", "PUT"])
@requiere_login
def renombrar_chat(chat_id):
    data = request.get_json(silent=True) or {}
    titulo = str(data.get("titulo", "")).strip()[:100]
    if not titulo:
        return jsonify({"ok": False, "error": "El título no puede estar vacío."}), 400
    try:
        ok = _actualizar_titulo_chat(chat_id, session["usuario"], titulo)
    except Exception as error:
        print("ERROR renombrar_chat:", error)
        return jsonify({"ok": False, "error": "No se pudo renombrar la conversación."}), 500
    if not ok:
        return jsonify({"ok": False, "error": "Conversación no encontrada."}), 404
    return jsonify({"ok": True, "id": chat_id, "titulo": titulo})


# CHAT
# ==========================================================

def _normalizar_encabezado(valor):
    texto = unicodedata.normalize("NFKD", str(valor or ""))
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    return re.sub(r"[^a-z0-9]+", "", texto.lower())


def _construir_fila_excel(campos_fila, indices, cantidad_columnas, libro_id):
    """Mapea un dict de campos (asegurado, patente, marca_modelo, etc.) a una
    fila del Excel real, usando los encabezados/alias existentes.

    Extraída de la propuesta manual (`/api/excel/agregar-fila`) para poder
    reutilizarse también desde el guardado autónomo de `/flota`: ambos casos
    necesitan exactamente el mismo mapeo de columnas.
    """
    normalizar = _normalizar_encabezado
    fila_nueva = [""] * cantidad_columnas

    alias = {
        normalizar("dominio"): "patente",
        normalizar("chapa"): "patente",
        normalizar("marca/modelo"): "marca_modelo",
        normalizar("marca_modelo"): "marca_modelo",
        normalizar("marca - modelo"): "marca_modelo",
        normalizar("vehiculo"): "vehiculo",
        normalizar("descripcion del vehiculo"): "vehiculo",
        normalizar("anio"): "año",
        normalizar("uso del vehiculo"): "uso",
        normalizar("suma asegurada"): "suma",
        normalizar("suma_asegurada"): "suma",
        normalizar("asegurado"): "asegurado",
        normalizar("domicilio"): "domicilio",
        normalizar("localidad"): "localidad",
        normalizar("cp"): "codigo postal",
        normalizar("codigo postal"): "codigo postal",
        normalizar("cia"): "compania",
        normalizar("compania"): "compania",
        normalizar("compañia"): "compania",
        normalizar("aseguradora"): "compania",
        normalizar("telefono"): "telefono",
        normalizar("teléfono"): "telefono",
    }
    campos_canonicos = {}
    for clave, valor in campos_fila.items():
        clave_normalizada = normalizar(clave)
        canonico = alias.get(clave_normalizada, str(clave or "").strip())
        if clave_normalizada in {
            normalizar("patente"),
            normalizar("dominio"),
            normalizar("chapa"),
        }:
            canonico = "PATENTE"
        campos_canonicos[normalizar(canonico)] = str(valor or "").strip()

    marca_modelo = str(campos_fila.get("marca_modelo") or "").strip()
    marca = str(campos_fila.get("marca") or "").strip()
    modelo = str(campos_fila.get("modelo") or "").strip()
    vehiculo = str(campos_fila.get("vehiculo") or "").strip()
    if marca_modelo:
        vehiculo = marca_modelo
    elif marca and modelo:
        vehiculo = f"{marca} {modelo}".strip()
    if vehiculo:
        campos_canonicos.setdefault(normalizar("vehiculo"), vehiculo)
        campos_canonicos.setdefault(normalizar("marca/modelo"), vehiculo)
        campos_canonicos.setdefault(normalizar("marca_modelo"), vehiculo)

    for campo, valor in campos_fila.items():
        clave = str(campo or "").strip()
        canonico = alias.get(normalizar(clave), clave)
        if normalizar(canonico) not in campos_canonicos:
            campos_canonicos[normalizar(canonico)] = str(valor or "").strip()

    if libro_id == "1":
        indice_asegurado = indices.get(normalizar("ASEGURADO"))
        indice_numero = indices.get(normalizar("NUMERO"))
        indice_patente = indices.get(normalizar("PATENTE"))
        asegurado = campos_canonicos.get(normalizar("ASEGURADO"), "")
        numero = campos_canonicos.get(normalizar("NUMERO"), "")
        patente = campos_canonicos.get(normalizar("PATENTE"), "")

        if not asegurado:
            raise ValueError(
                "Antes de guardar, el registro necesita al menos el nombre del ASEGURADO."
            )
        if not numero and not patente:
            raise ValueError(
                "Antes de guardar, indicá al menos NUMERO (DNI/póliza) o PATENTE."
            )
        if indice_asegurado is None:
            raise ValueError("El Excel no tiene la columna ASEGURADO.")
        if indice_numero is None and indice_patente is None:
            raise ValueError(
                "El Excel no tiene NUMERO ni PATENTE para identificar el registro."
            )

    for campo, valor in campos_canonicos.items():
        indice = indices.get(campo)
        if indice is not None:
            fila_nueva[indice] = valor

    if libro_id == "2":
        for encabezado, indice in indices.items():
            if encabezado in {
                normalizar("patente"),
                normalizar("dominio"),
                normalizar("chapa"),
            }:
                fila_nueva[indice] = campos_canonicos.get(
                    normalizar("patente"), fila_nueva[indice]
                )
            elif encabezado in {
                normalizar("marca"),
                normalizar("marca del vehiculo"),
            }:
                fila_nueva[indice] = str(
                    campos_fila.get("marca") or fila_nueva[indice]
                ).strip()
            elif encabezado in {
                normalizar("modelo"),
                normalizar("modelo del vehiculo"),
            }:
                fila_nueva[indice] = str(
                    campos_fila.get("modelo") or fila_nueva[indice]
                ).strip()
            elif encabezado in {
                normalizar("marca/modelo"),
                normalizar("marca_modelo"),
                normalizar("marca - modelo"),
                normalizar("vehiculo"),
                normalizar("descripcion del vehiculo"),
            }:
                fila_nueva[indice] = vehiculo or fila_nueva[indice]
            elif encabezado in {normalizar("asegurado"), normalizar("nombre asegurado")}:
                fila_nueva[indice] = campos_canonicos.get(normalizar("asegurado"), fila_nueva[indice])
            elif encabezado in {normalizar("domicilio"), normalizar("direccion")}:
                fila_nueva[indice] = campos_canonicos.get(normalizar("domicilio"), fila_nueva[indice])
            elif encabezado in {normalizar("localidad"), normalizar("ciudad")}:
                fila_nueva[indice] = campos_canonicos.get(normalizar("localidad"), fila_nueva[indice])
            elif encabezado in {normalizar("cp"), normalizar("codigo postal")}:
                fila_nueva[indice] = campos_canonicos.get(normalizar("cp"), fila_nueva[indice])
            elif encabezado in {normalizar("año"), normalizar("anio")}:
                fila_nueva[indice] = campos_canonicos.get(
                    normalizar("año"), fila_nueva[indice]
                )
            elif encabezado in {
                normalizar("uso"),
                normalizar("uso del vehiculo"),
            }:
                fila_nueva[indice] = campos_canonicos.get(
                    normalizar("uso"), fila_nueva[indice]
                )
            elif encabezado in {
                normalizar("suma"),
                normalizar("suma asegurada"),
            }:
                fila_nueva[indice] = campos_canonicos.get(
                    normalizar("suma"), fila_nueva[indice]
                )
            elif encabezado == normalizar("cobertura"):
                fila_nueva[indice] = campos_canonicos.get(
                    normalizar("cobertura"), fila_nueva[indice]
                )
            elif encabezado in {normalizar("motor")}:
                fila_nueva[indice] = campos_canonicos.get(normalizar("motor"), fila_nueva[indice])
            elif encabezado in {normalizar("chasis")}:
                fila_nueva[indice] = campos_canonicos.get(normalizar("chasis"), fila_nueva[indice])

    if not any(str(valor).strip() for valor in fila_nueva):
        raise ValueError(
            "Ninguno de los campos propuestos coincide con las columnas existentes del Excel."
        )
    return fila_nueva


@app.route("/api/excel/agregar-fila", methods=["POST"])
@requiere_login
def api_excel_agregar_fila():
    data = request.get_json(silent=True) or {}
    tipo_propuesta = str(data.get("tipo_propuesta") or "").strip().lower()
    filas_propuesta = data.get("filas")
    campos = data.get("campos")
    libro_id = str(data.get("libro_id") or session.get("guardar_asegurado_libro_id") or "1")

    es_flota = tipo_propuesta == "flota"
    if es_flota:
        if not isinstance(filas_propuesta, list) or not filas_propuesta:
            return jsonify({"ok": False, "error": "No se recibieron vehículos para agregar."}), 400
        if not all(isinstance(fila, dict) and fila for fila in filas_propuesta):
            return jsonify({"ok": False, "error": "La propuesta de flota contiene vehículos inválidos."}), 400
    elif not isinstance(campos, dict) or not campos:
        return jsonify({"ok": False, "error": "No se recibieron campos para agregar."}), 400

    if libro_id not in LIBROS_EXCEL:
        return jsonify({"ok": False, "error": "Libro de Excel no válido."}), 400

    try:
        datos = leer_excel_interno(libro_id)
        filas = list(datos.get("filas") or [])
        hoja_actual = datos.get("hoja", "Datos")
        if not filas:
            return jsonify({"ok": False, "error": "El Excel interno no tiene encabezados."}), 400

        encabezados = filas[0]
        cantidad_columnas = max(len(encabezados), 1)
        indices = {
            _normalizar_encabezado(encabezado): i
            for i, encabezado in enumerate(encabezados)
            if _normalizar_encabezado(encabezado)
        }

        if es_flota:
            filas_nuevas = [
                _construir_fila_excel(fila, indices, cantidad_columnas, libro_id)
                for fila in filas_propuesta
            ]
            guardar_matriz_excel(
                filas + filas_nuevas, hoja_actual, libro_id=libro_id
            )
        else:
            fila_nueva = _construir_fila_excel(campos, indices, cantidad_columnas, libro_id)
            guardar_matriz_excel(
                filas + [fila_nueva], hoja_actual, libro_id=libro_id
            )

        session.pop("guardar_asegurado_libro_id", None)
        texto_envios_ya = None
        if not es_flota and libro_id == "1":
            try:
                texto_envios_ya = _armar_texto_envios_ya(campos)
            except Exception as error:
                print("ERROR ARMANDO TEXTO ENVIOS YA:", error)
        return jsonify({
            "ok": True,
            "libro_id": libro_id,
            "filas_agregadas": len(filas_propuesta) if es_flota else 1,
            "texto_envios_ya": texto_envios_ya,
            **leer_excel_interno(libro_id),
        })
    except ValueError as error:
        return jsonify({"ok": False, "error": str(error)}), 400
    except Exception as error:
        print("ERROR AGREGANDO FILA DESDE CHAT:", error)
        return jsonify({"ok": False, "error": "No se pudo agregar el registro al Excel."}), 500


def _dividir_marca_modelo_flota(marca_modelo):
    """Igual que _dividir_marca_modelo, pero para flotas de camiones/
    acoplados: acá "primera palabra = marca" falla con marcas de dos
    palabras (M. BENZ, SOLA Y BRUSA). Probamos contra una lista de marcas
    conocidas del rubro transporte antes de caer al split simple."""
    texto = str(marca_modelo or "").strip()
    if not texto:
        return "", ""
    texto_norm = texto.upper()
    for marca in _MARCAS_TRANSPORTE_CONOCIDAS:
        if texto_norm == marca or texto_norm.startswith(marca + " "):
            modelo = texto[len(marca):].strip()
            return marca, modelo
    partes = texto.split(" ", 1)
    return partes[0], (partes[1].strip() if len(partes) > 1 else "")


_MARCAS_TRANSPORTE_CONOCIDAS = sorted(
    [
        "M. BENZ", "MERCEDES BENZ", "SCANIA", "FORD", "RANDON", "SOLA Y BRUSA",
        "ULTRANS", "FIAT", "IVECO", "VOLKSWAGEN", "VOLVO", "AGRALE", "DAF",
        "HYUNDAI", "MAN", "TOYOTA", "CHEVROLET",
    ],
    key=len,
    reverse=True,
)

# Enumerados fijos del formato de flota de La Segunda: es una tabla sin
# etiquetas por campo, pero con columnas en orden constante, así que se
# puede anclar con estos valores conocidos en vez de depender de Gemini.
_LASEGUNDA_TIPOS = [
    "SEMI-REMOLQUE", "ACOPLADO", "CAMION", "CAMIÓN",
    "P GI H/1000Kg", "P GII +1000Kg", "P GII H/1000Kg",
    "AUT/FAM/M.VAN", "JEEP/CAM.FAM",
]
_LASEGUNDA_CARROCERIAS = [
    "FURGON DE FABRICA", "FURGÓN DE FABRICA", "GRANEL", "ABIERTA", "CERRADA",
    "CARGAS GENERALES", "CON BARANDAS", "SEDAN", "RURAL", "DOBLE CABINA",
]
_LASEGUNDA_USOS = ["COMER.H/4 TT LD", "CARG PELIG L.D.", "COMERCIAL LD", "PARTICULAR"]


def _alt_regex(lista):
    return "|".join(re.escape(x) for x in sorted(lista, key=len, reverse=True))


_PATRON_LASEGUNDA = re.compile(
    r"(?P<tipo>" + _alt_regex(_LASEGUNDA_TIPOS) + r")\s+"
    r"(?P<carroceria>" + _alt_regex(_LASEGUNDA_CARROCERIAS) + r")\s+"
    r"(?P<marca_modelo>.+?)\s+"
    r"(?P<anio>(?:19|20)\d{2})\s+"
    r"(?P<uso>" + _alt_regex(_LASEGUNDA_USOS) + r")\s+"
    r"(?P<patente>[A-Z0-9]{6,7})\s+"
    r"(?P<motor>\S+)\s+"
    r"(?P<chasis>\S+)\s+"
    r"\$(?P<limite>[\d,]+\.\d{2})\s+"
    r"(?P<plan>\d{2})\s+"
    r"\$(?P<suma>[\d,]+\.\d{2})\s+"
    r"(?:SI|--)\s+(?:SI|--)\s+(?:SI|--)\s+(?:SI|--)\s+(?:SI|--)\s+(?:SI|--)",
    re.IGNORECASE,
)

# Variante de respaldo: algunas filas (típicamente IVECO/FIAT) traen motor
# y chasis pegados sin espacio en el texto original de la póliza (bug del
# parser de origen de La Segunda). En vez de perder toda la fila, se
# captura el bloque motor+chasis como un solo token y se marca "sospechoso"
# para que se revise/separe a mano en vez de asumir un corte automático.
_PATRON_LASEGUNDA_MOTOR_CHASIS_PEGADO = re.compile(
    r"(?P<tipo>" + _alt_regex(_LASEGUNDA_TIPOS) + r")\s+"
    r"(?P<carroceria>" + _alt_regex(_LASEGUNDA_CARROCERIAS) + r")\s+"
    r"(?P<marca_modelo>.+?)\s+"
    r"(?P<anio>(?:19|20)\d{2})\s+"
    r"(?P<uso>" + _alt_regex(_LASEGUNDA_USOS) + r")\s+"
    r"(?P<patente>[A-Z0-9]{6,7})\s+"
    r"(?P<motor_chasis>\S{15,})\s+"
    r"\$(?P<limite>[\d,]+\.\d{2})\s+"
    r"(?P<plan>\d{2})\s+"
    r"\$(?P<suma>[\d,]+\.\d{2})\s+"
    r"(?:SI|--)\s+(?:SI|--)\s+(?:SI|--)\s+(?:SI|--)\s+(?:SI|--)\s+(?:SI|--)",
    re.IGNORECASE,
)


def _parece_formato_lasegunda(texto):
    """Heurística barata para decidir si vale la pena probar este parser:
    el formato de La Segunda no tiene etiquetas de campo, pero sí trae al
    menos 2 de los enumerados fijos (tipo/carrocería/uso) + varios importes
    en pesos con el patrón típico de la planilla."""
    texto_norm = texto.upper()
    hits = sum(1 for t in _LASEGUNDA_TIPOS if t in texto_norm)
    hits += sum(1 for u in _LASEGUNDA_USOS if u in texto_norm)
    tiene_importes = len(re.findall(r"\$[\d,]+\.\d{2}", texto)) >= 2
    return hits >= 2 and tiene_importes


def _lasegunda_conteo_bloques(texto):
    """Cuenta cuántas filas de vehículo debería haber en el texto, contando
    ocurrencias de los TIPO fijos (SEMI-REMOLQUE, ACOPLADO, CAMION...) que
    marcan el arranque de cada fila en el formato de La Segunda. Sirve para
    detectar si _parsear_flota_lasegunda se salteó alguna (p.ej. porque
    motor y chasis vinieron pegados sin espacio, como pasó con el Fiat) sin
    tener que revisar el Excel fila por fila."""
    texto_norm = re.sub(r"\s+", " ", str(texto or "")).upper()
    patron_tipo = re.compile(r"\b(?:" + _alt_regex(_LASEGUNDA_TIPOS) + r")\b")
    return len(patron_tipo.findall(texto_norm))


# Marca el FINAL de cada fila de La Segunda: los 6 flags fijos SI/-- de
# coberturas puntuales, seguidos opcionalmente del importe de costo
# mensual que a veces viene pegado justo después en el mismo renglón.
# Se usa como límite de fila en vez del TIPO/CARROCERÍA del principio
# porque los vehículos PARTICULARES no traen esas dos columnas: si se
# ancla el corte solo por TIPO, el texto de un auto particular queda
# pegado a la fila del camión anterior y termina mezclado con ella (el
# caso del Fiat Strada + Ford Bronco fusionados en una sola celda).
_PATRON_TERMINADOR_FILA_LASEGUNDA = re.compile(
    r"(?:SI|--)(?:\s+(?:SI|--)){5}(?:\s+\$[\d,]+\.\d+)?",
    re.IGNORECASE,
)


def _dividir_filas_lasegunda(texto_plano):
    """Divide el texto (ya aplanado a una sola línea) en un bloque por
    fila de vehículo, cortando después de cada terminador de fila. Cada
    bloque resultante se procesa después de forma AISLADA, así un error
    de lectura en una fila (campos pegados, formato distinto) no puede
    arrastrar texto hacia la fila vecina."""
    terminadores = list(_PATRON_TERMINADOR_FILA_LASEGUNDA.finditer(texto_plano))
    if not terminadores:
        return [texto_plano] if texto_plano.strip() else []
    bloques = []
    inicio = 0
    for term in terminadores:
        bloque = texto_plano[inicio:term.end()].strip()
        # Un bloque que es sólo guiones/espacios (el separador
        # "----------" del formato de La Segunda cuando queda aislado
        # entre dos filas) no es un vehículo: se descarta para no
        # generar un vehículo fantasma "sin_parsear" vacío.
        if bloque and re.sub(r"[-\s]+", "", bloque):
            bloques.append(bloque)
        inicio = term.end()
    resto = texto_plano[inicio:].strip()
    if resto and re.sub(r"[-\s]+", "", resto):
        bloques.append(resto)
    return bloques


def _parsear_flota_lasegunda(texto):
    """Parser determinístico para el formato tabular de La Segunda (sin
    etiquetas por campo, columnas fijas). Devuelve una lista de vehículos
    en el mismo formato que usa el resto de /flota, o [] si no matchea
    nada (en ese caso el llamador sigue con el flujo normal).

    Primero se corta el texto en bloques por fila (ver
    `_dividir_filas_lasegunda`) y recién después se aplica el regex de
    campos a cada bloque por separado. Una fila que no matchea el patrón
    estricto de camión (típicamente un vehículo PARTICULAR, sin
    TIPO/CARROCERÍA) no se pierde ni se mezcla con la vecina: queda como
    fila "sin parsear" con el texto crudo, para completar a mano."""
    texto_plano = re.sub(r"\s+", " ", str(texto or "")).strip()
    vehiculos = []
    for bloque in _dividir_filas_lasegunda(texto_plano):
        m = _PATRON_LASEGUNDA.search(bloque)
        if m:
            marca, modelo = _dividir_marca_modelo_flota(m.group("marca_modelo"))
            vehiculos.append({
                "patente": m.group("patente").upper(),
                "marca_modelo": m.group("marca_modelo").strip(),
                "marca": marca,
                "modelo": modelo,
                "año": m.group("anio"),
                "motor": m.group("motor"),
                "chasis": m.group("chasis"),
                "uso": m.group("uso"),
                "suma_asegurada": m.group("suma"),
                # La Segunda no trae un texto de cobertura tipo "TODO
                # RIESGO": tiene un código de Plan + 6 flags SI/-- de
                # coberturas puntuales. Dejamos cobertura vacía a
                # propósito en vez de inventar una equivalencia; se
                # completa a mano si hace falta.
                "cobertura": "",
            })
            continue
        m2 = _PATRON_LASEGUNDA_MOTOR_CHASIS_PEGADO.search(bloque)
        if m2:
            marca, modelo = _dividir_marca_modelo_flota(m2.group("marca_modelo"))
            vehiculos.append({
                "patente": m2.group("patente").upper(),
                "marca_modelo": m2.group("marca_modelo").strip(),
                "marca": marca,
                "modelo": modelo,
                "año": m2.group("anio"),
                "motor": "",
                # Motor y chasis venían pegados sin espacio en el texto
                # original: se deja el bloque completo en chasis (no se
                # adivina dónde cortar) y se marca sospechoso para que
                # se revise/separe a mano antes de pegar en Excel.
                "chasis": m2.group("motor_chasis"),
                "uso": m2.group("uso"),
                "suma_asegurada": m2.group("suma"),
                "cobertura": "",
                "sospechoso": True,
                "motivo_sospecha": "Motor y chasis pegados sin espacio en el texto original — separar a mano.",
            })
        else:
            vehiculos.append({
                "patente": "",
                "marca_modelo": bloque,
                "marca": "",
                "modelo": "",
                "año": "",
                "motor": "",
                "chasis": "",
                "uso": "",
                "suma_asegurada": "",
                "cobertura": "",
                "sin_parsear": True,
            })
    return vehiculos


def interpretar_flota_a_json(texto):
    """Interpreta una o varias descripciones de vehículos de una póliza.

    Primero intenta extraer de forma determinista los campos que vienen con
    etiquetas explícitas en el frente de póliza. Gemini queda como respaldo
    para textos que no respeten ese formato.
    """
    texto = str(texto or "").replace("\r", "")

    campos = (
        "patente", "marca_modelo", "marca", "modelo", "año", "motor",
        "chasis", "uso", "suma_asegurada", "cobertura",
        "asegurado", "domicilio", "localidad", "cp"
    )

    def limpiar_valor(valor):
        return re.sub(r"\s+", " ", str(valor or "")).strip(" \t\n:;,-")

    def campo_etiquetado(bloque, etiqueta, etiquetas_siguientes):
        patron = rf"{re.escape(etiqueta)}\s*:\s*(.*?)(?=\s+(?:{'|'.join(re.escape(x) for x in etiquetas_siguientes)})\s*:|$)"
        m = re.search(patron, bloque, flags=re.IGNORECASE | re.DOTALL)
        return limpiar_valor(m.group(1)) if m else ""

    # FORMATO LA SEGUNDA: tabla sin etiquetas por campo, columnas fijas
    # (ver _parsear_flota_lasegunda). Se intenta ANTES del formato con
    # etiquetas, porque si el texto matchea este patrón tabular no va a
    # tener "DESCRIPCIÓN DEL VEHÍCULO ASEGURADO:" para lo otro.
    if _parece_formato_lasegunda(texto):
        vehiculos_lasegunda = _parsear_flota_lasegunda(texto)
        if vehiculos_lasegunda:
            resultado = {"vehiculos": vehiculos_lasegunda}
            sin_parsear = [v for v in vehiculos_lasegunda if v.get("sin_parsear")]
            sospechosos = [v for v in vehiculos_lasegunda if v.get("sospechoso")]
            avisos = []
            if sin_parsear:
                avisos.append(
                    f"{len(sin_parsear)} fila(s) no matchearon el patrón de camión/acoplado "
                    "conocido y quedaron con el texto original en MARCA/MODELO — completá "
                    "esas a mano."
                )
            if sospechosos:
                avisos.append(
                    f"{len(sospechosos)} fila(s) tenían motor y chasis pegados sin espacio en "
                    "el texto original: quedaron juntos en la columna CHASIS — separalos a "
                    "mano antes de pegar."
                )
            total_ok = len(vehiculos_lasegunda) - len(sin_parsear) - len(sospechosos)
            if avisos:
                resultado["aviso_conteo"] = (
                    f"Ojo: de {len(vehiculos_lasegunda)} vehículo(s) detectados, {total_ok} se "
                    "cargaron completos. " + " ".join(avisos)
                )
            return resultado

    # Cada aparición de "DESCRIPCIÓN DEL <VEHÍCULO/AUTOMOTOR/RODADO/UNIDAD>
    # ASEGURADO" marca un vehículo. Distintas compañías usan palabras
    # distintas para lo mismo en el frente de póliza, así que el marcador
    # acepta las variantes más comunes en vez de exigir "VEHÍCULO" literal.
    marcadores = list(re.finditer(
        r"DESCRIPCI[ÓO]N\s+DEL\s+"
        r"(?:VEH[ÍI]CULO|AUTOMOTOR|AUTOM[ÓO]VIL|RODADO|UNIDAD|AUTO|COCHE)\s+"
        r"ASEGURAD[OA]\s*:",
        texto,
        flags=re.IGNORECASE,
    ))

    if marcadores:
        vehiculos = []
        etiquetas = [
            "TIPO", "MARCA/MODELO", "AÑO", "PATENTE", "MOTOR", "CHASIS",
            "AUTO/JEEP/SUV PARTICULARES Y FAMILIARES (1-1-1)",
            "USO DEL VEHÍCULO", "USO DEL VEHICULO", "SUMA ASEGURADA", "COBERTURA",
        ]
        for i, marcador in enumerate(marcadores):
            fin = marcadores[i + 1].start() if i + 1 < len(marcadores) else len(texto)
            bloque = texto[marcador.end():fin]
            # Normalizamos sinónimos de "VEHÍCULO" dentro de la etiqueta
            # "USO DEL ..." para no tener que duplicar cada regex de más
            # abajo por cada variante (automotor, rodado, unidad, etc.).
            bloque = re.sub(
                r"USO\s+DEL\s+(?:AUTOMOTOR|AUTOM[ÓO]VIL|RODADO|UNIDAD|AUTO|COCHE)",
                "USO DEL VEHICULO",
                bloque,
                flags=re.IGNORECASE,
            )
            etiquetas_campos = [
                "TIPO", "MARCA/MODELO", "AÑO", "PATENTE", "MOTOR", "CHASIS",
                "USO DEL VEHÍCULO", "USO DEL VEHICULO", "SUMA ASEGURADA", "COBERTURA"
            ]
            marca_modelo = campo_etiquetado(bloque, "MARCA/MODELO", [
                "AÑO", "PATENTE", "MOTOR", "CHASIS", "USO DEL VEHÍCULO",
                "USO DEL VEHICULO", "SUMA ASEGURADA", "COBERTURA"
            ])
            patente = campo_etiquetado(bloque, "PATENTE", [
                "MOTOR", "CHASIS", "USO DEL VEHÍCULO", "USO DEL VEHICULO",
                "SUMA ASEGURADA", "COBERTURA"
            ])
            anio = campo_etiquetado(bloque, "AÑO", [
                "PATENTE", "MOTOR", "CHASIS", "USO DEL VEHÍCULO",
                "USO DEL VEHICULO", "SUMA ASEGURADA", "COBERTURA"
            ])
            motor = campo_etiquetado(bloque, "MOTOR", [
                "CHASIS", "USO DEL VEHÍCULO", "USO DEL VEHICULO",
                "SUMA ASEGURADA", "COBERTURA"
            ])
            m_chasis = re.search(
                r"CHASIS\s*:\s*(.*?)(?=\s+AUTO/JEEP/SUV\s+PARTICULARES\s+Y\s+FAMILIARES\s+\(1-1-1\)|\s+USO\s+DEL\s+VEH[ÍI]CULO\s*:|\s+SUMA\s+ASEGURADA\s*:|\s+COBERTURA\s*:|$)",
                bloque,
                flags=re.IGNORECASE | re.DOTALL,
            )
            chasis = limpiar_valor(m_chasis.group(1)) if m_chasis else ""
            uso = campo_etiquetado(bloque, "USO DEL VEHÍCULO", ["SUMA ASEGURADA", "COBERTURA"])
            if not uso:
                uso = campo_etiquetado(bloque, "USO DEL VEHICULO", ["SUMA ASEGURADA", "COBERTURA"])
            suma = campo_etiquetado(bloque, "SUMA ASEGURADA", ["COBERTURA"])
            cobertura = campo_etiquetado(bloque, "COBERTURA", [])
            # Los frentes pueden traer pie de página después de COBERTURA.
            cobertura = re.split(
                r"\s+-\s+(?:Advertencia\b|\d{4,}\s+Tel\.?|Tel\.?\s*:|Provincia\b|Condición\b|ASEGURADO\b|PRODUCTOR\b)",
                cobertura,
                maxsplit=1,
                flags=re.IGNORECASE,
            )[0].strip()


            # MARCA/MODELO se conserva EXACTAMENTE como figura en la póliza.
            # No se separa en marca y modelo para no alterar el dato original.
            vehiculo = {
                "patente": patente,
                "marca_modelo": marca_modelo,
                "marca": "",
                "modelo": "",
                "año": anio,
                "motor": motor,
                "chasis": chasis,
                "uso": uso,
                "suma_asegurada": suma,
                "cobertura": cobertura,
            }
            if any(vehiculo[k] for k in ("patente", "marca_modelo", "año", "motor", "chasis")):
                vehiculos.append(vehiculo)

        # Datos generales del asegurado: se aplican a todos los vehículos.
        # El CP se obtiene del bloque de cabecera si está presente como (NNNN)
        # o junto a una localidad/código postal explícito.
        cabecera = texto[:marcadores[0].start()]
        cp = ""
        m_cp = re.search(r"(?:\(|\b)(\d{4})\)?\b", cabecera)
        if m_cp:
            cp = m_cp.group(1)
        m_cp2 = re.search(r"(?:C[ÓO]D(?:IGO)?\s*POSTAL|CP)\s*:?\s*(\d{4})", cabecera, re.IGNORECASE)
        if m_cp2:
            cp = m_cp2.group(1)

        # En muchos frentes el nombre del asegurado aparece en la cabecera
        # antes del domicilio. Intentamos extraerlo sin tocar los datos de los
        # vehículos. Si no hay suficiente estructura, dejamos el campo vacío
        # antes que inventarlo.
        asegurado = ""
        domicilio = ""
        localidad = ""
        cabecera_limpia = re.sub(r"^/flota\s*", "", cabecera, flags=re.IGNORECASE).strip()
        cabecera_limpia = re.sub(r"\s+", " ", cabecera_limpia)

        # Caso habitual: NOMBRE + CALLE + ALTURA + LOCALIDAD + DNI/IVA + (CP).
        # Para no confundir nombre y calle, buscamos un número de altura y
        # usamos el segmento anterior como candidato. Si hay una etiqueta
        # explícita, ésta tiene prioridad.
        m_aseg_et = re.search(r"(?:ASEGURADO|TOMADOR)\s*:\s*([^:]{2,100})", cabecera_limpia, re.IGNORECASE)
        if m_aseg_et and limpiar_valor(m_aseg_et.group(1)):
            asegurado = limpiar_valor(m_aseg_et.group(1))

        if not asegurado:
            # Casos frecuentes de domicilios cuyo nombre de calle es fácilmente
            # identificable en el texto corrido (LA RIOJA, AVENIDA, CALLE, etc.).
            m_aseg_calle = re.search(
                r"^(.+?)\s+(?=(?:LA\s+RIOJA|AV(?:ENIDA)?|CALLE|RUTA|BARRIO)\b[^0-9]{0,50}\s+\d{1,6}\b)",
                cabecera_limpia,
                re.IGNORECASE,
            )
            if m_aseg_calle:
                asegurado = limpiar_valor(m_aseg_calle.group(1))

        if not asegurado:
            # El formato de ejemplo del frente usa tres palabras para el
            # nombre antes de la calle. Preferimos esa estructura cuando está
            # seguida por una calle + altura.
            m_aseg = re.search(
                r"^([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ'’-]+(?:\s+[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ'’-]+){1,4})\s+(?=[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ .'-]{2,40}\s+\d{1,6}\b)",
                cabecera_limpia,
                re.IGNORECASE,
            )
            if m_aseg:
                asegurado = limpiar_valor(m_aseg.group(1))

        # Extraer domicilio/localidad si la cabecera tiene una calle + altura.
        if asegurado:
            resto = cabecera_limpia[len(asegurado):].strip()
        else:
            resto = cabecera_limpia
        m_dir = re.search(
            r"^([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ .'-]{2,50}?)\s+(\d{1,6})\s+([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ .'-]{2,40}?)(?=\s+\d{6,11}\b|\s+CONSUMIDOR\b|\s+RESPONSABLE\b|\s+\(|$)",
            resto,
            re.IGNORECASE,
        )
        if m_dir:
            domicilio = limpiar_valor(f"{m_dir.group(1)} {m_dir.group(2)}")
            localidad = limpiar_valor(m_dir.group(3))


        for v in vehiculos:
            v["asegurado"] = asegurado
            v["domicilio"] = domicilio
            v["localidad"] = localidad
            v["cp"] = cp

        return {"vehiculos": vehiculos}

    # Fallback Gemini para formatos no estructurados.
    from servicios_ia import obtener_cliente_gemini, MODELOS_GEMINI
    from google.genai import types

    cliente = obtener_cliente_gemini()
    if cliente is None:
        raise RuntimeError("La IA todavía no está configurada. Falta GEMINI_API_KEY.")

    instruccion = """
Vas a recibir el texto crudo de un frente de póliza de flota de una compañía de
seguros argentina, pegado sin formato. Tu tarea es dividirlo en un vehículo por
fila e identificar sus datos.

NO asumas que conocés el formato exacto de esta compañía. Cada aseguradora
ordena las columnas distinto y usa sus propias etiquetas de tipo/carrocería/uso.
En vez de buscar palabras específicas de una compañía, ancla la división de
filas en estas tres señales, que SIEMPRE están presentes sin importar la
compañía:

1. PATENTE: una secuencia de 6-7 caracteres alfanuméricos en mayúscula con
   formato argentino (3 letras + 3 números, o 2 letras + 3 números + 2 letras).
   Aparece EXACTAMENTE UNA VEZ por vehículo. Es tu ancla principal: cada
   patente que encontrás marca un vehículo distinto.
2. AÑO: un número de 4 dígitos entre 1990 y el año actual.
3. MONTOS EN PESOS: valores tipo $X.XXX.XXX,XX o $X,XXX,XXX.XX. Cada vehículo
   trae al menos dos (límite de cobertura y suma asegurada).

CÓMO DIVIDIR LAS FILAS:
- Localizá todas las patentes del texto, en orden de aparición.
- Para cada patente, el vehículo es el tramo de texto que va desde el final de
  la fila anterior (o el inicio del texto, si es la primera patente) hasta el
  final de los datos asociados a esa patente (generalmente después del último
  monto o de los últimos flags SI/NO/-- de esa fila, antes de que empiece el
  texto de tipo/carrocería del vehículo siguiente).
- Si dos patentes aparecen muy cerca sin datos numéricos entre medio, revisá si
  en verdad es la misma fila partida en dos renglones (unilas) o dos vehículos
  distintos con muy poca descripción.

Devolvé ÚNICAMENTE JSON válido, sin markdown, comentarios ni texto fuera del
JSON. La estructura obligatoria es:
{
  "vehiculos": [
    {
      "patente": "",
      "marca_modelo": "",
      "año": "",
      "motor": "",
      "chasis": "",
      "uso": "",
      "suma_asegurada": "",
      "cobertura": "",
      "asegurado": "",
      "domicilio": "",
      "localidad": "",
      "cp": "",
      "sospechoso": false,
      "motivo_sospecha": ""
    }
  ]
}

REGLA CRÍTICA: MARCA/MODELO (o el texto libre de tipo/carrocería/marca/modelo
cuando no hay etiquetas separadas) debe copiarse EXACTAMENTE como aparece en la
póliza, sin separar, resumir, corregir, traducir ni reinterpretar. Por ejemplo,
"PEUGEOT PARTNER PATA. 1.6 VTC PLUS L10/17" debe quedar exactamente así.

CASOS QUE TENÉS QUE MARCAR EN VEZ DE ADIVINAR:
- Si motor y chasis aparecen pegados sin espacio (una sola cadena muy larga,
  más de 15 caracteres sin separación), NO intentes cortar arbitrariamente
  dónde termina uno y empieza el otro. Poné el bloque completo en "chasis",
  dejá "motor" vacío, "sospechoso": true y "motivo_sospecha": "motor y chasis
  pegados, revisar a mano".
- Si un tramo de texto no tiene una patente reconocible cerca, no lo conviertas
  en un vehículo: probablemente es texto de cabecera (datos del asegurado) o
  pie de página.

No mezcles información entre vehículos. No inventes datos: si un campo no
aparece con certeza, dejalo vacío en vez de adivinarlo. La cantidad de
vehículos del resultado debe corresponder a la cantidad real de patentes
detectadas — preferí dejar más texto libre en una fila antes que perder un
vehículo o mezclar dos.
"""
    ultimo_error = None
    for modelo in MODELOS_GEMINI:
        try:
            config = types.GenerateContentConfig(
                temperature=0,
                max_output_tokens=4000,
                response_mime_type="application/json",
                system_instruction=instruccion.strip(),
            )
            respuesta = cliente.models.generate_content(
                model=modelo,
                contents=texto.strip(),
                config=config,
            )
            bruto = str(getattr(respuesta, "text", "") or "").strip()
            if not bruto:
                raise ValueError("Gemini no devolvió JSON.")
            datos = json.loads(bruto)
            vehiculos = datos.get("vehiculos")
            if not isinstance(vehiculos, list):
                raise ValueError("Gemini no devolvió la lista de vehículos.")
            salida = []
            sospechosos = 0
            for v in vehiculos:
                if not isinstance(v, dict):
                    continue
                fila = {k: limpiar_valor(v.get(k, "")) for k in campos}
                if v.get("sospechoso"):
                    fila["sospechoso"] = True
                    fila["motivo_sospecha"] = limpiar_valor(v.get("motivo_sospecha", ""))
                    sospechosos += 1
                salida.append(fila)
            resultado = {"vehiculos": salida}
            if sospechosos:
                palabra = "vehículo" if sospechosos == 1 else "vehículos"
                resultado["aviso_conteo"] = (
                    f"Ojo: {sospechosos} {palabra} de los {len(salida)} que encontré "
                    "conviene revisarlos antes de pasarlos al Excel, porque algunos datos "
                    "no se leyeron del todo bien."
                )
            return resultado
        except Exception as error:
            ultimo_error = error
            print("ERROR GEMINI /FLOTA", modelo, ":", error)
    raise RuntimeError(f"No pude interpretar el frente de póliza como JSON: {ultimo_error}")


# ==========================================================
# /FLOTA — CONTEXTO PERSISTENTE, FUSIÓN, DEDUP Y AUTOGUARDADO
# ==========================================================
#
# Estas funciones implementan el comportamiento pedido para /flota:
# la flota vive en `flotas_activas` (una fila por conversación) y cada
# mensaje nuevo ENRIQUECE ese estado en vez de arrancar de cero. Los
# vehículos se identifican por patente/chasis (o por ITEM si no hay
# ninguno de los dos) para no crear duplicados, se guardan en el Excel
# real apenas tienen algo mínimamente identificable, y las correcciones
# posteriores ("el 7 es C3") pisan la fila ya guardada en vez de agregar
# una nueva.

CAMPOS_VEHICULO = (
    "patente", "marca_modelo", "marca", "modelo", "año", "motor",
    "chasis", "uso", "suma_asegurada", "cobertura",
)

_ETIQUETAS_CAMPO_NATURAL = {
    "patente": ("patente", "dominio", "chapa"),
    "chasis": ("chasis",),
    "motor": ("motor",),
    "suma_asegurada": ("suma asegurada", "suma"),
    "año": ("año", "anio", "modelo año", "año modelo"),
    "uso": ("uso",),
    "cobertura": ("cobertura",),
}


def _vacio(valor):
    return not str(valor or "").strip()


def _normalizar_identificador(valor):
    texto = unicodedata.normalize("NFKD", str(valor or ""))
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    return re.sub(r"[^A-Z0-9]", "", texto.upper())


def _vehiculo_nuevo_vacio(item):
    vehiculo = {campo: "" for campo in CAMPOS_VEHICULO}
    vehiculo["item"] = item
    vehiculo["fila_excel"] = None
    return vehiculo


def _mismo_vehiculo(existente, nuevo):
    """True si `nuevo` describe el mismo vehículo físico que `existente`,
    usando patente o chasis como identificador confiable (Sección 22)."""
    pat_a = _normalizar_identificador(existente.get("patente"))
    pat_b = _normalizar_identificador(nuevo.get("patente"))
    if pat_a and pat_b:
        return pat_a == pat_b
    cha_a = _normalizar_identificador(existente.get("chasis"))
    cha_b = _normalizar_identificador(nuevo.get("chasis"))
    if cha_a and cha_b:
        return cha_a == cha_b
    return False


def _fusionar_campos_vehiculo(existente, nuevo):
    """Completa campos vacíos con datos nuevos. Nunca pisa un dato ya
    presente con uno vacío; si llega un valor distinto para un campo ya
    completo, lo actualiza (se asume que es una corrección del usuario,
    Sección 20/21) salvo que sea idéntico."""
    cambio = False
    for campo in CAMPOS_VEHICULO:
        valor_nuevo = str(nuevo.get(campo, "") or "").strip()
        if not valor_nuevo:
            continue
        if valor_nuevo != str(existente.get(campo, "") or "").strip():
            existente[campo] = valor_nuevo
            cambio = True
    return cambio


def _vehiculo_guardable(vehiculo):
    """Mínimo para que un registro ya tenga sentido guardado en el Excel
    (Sección 15/20): algo que lo identifique (patente o chasis) o al menos
    la descripción del vehículo."""
    return bool(
        str(vehiculo.get("patente") or "").strip()
        or str(vehiculo.get("chasis") or "").strip()
        or str(vehiculo.get("marca_modelo") or "").strip()
    )


_CAMPOS_RELEVANTES_PARA_COMPLETITUD = (
    "patente", "marca_modelo", "año", "motor", "chasis", "uso", "suma_asegurada", "cobertura",
)


def _campos_pendientes_vehiculo(vehiculo):
    # "marca" y "modelo" quedan afuera del reporte de faltantes: el extractor
    # siempre trabaja con "marca_modelo" combinado (Sección 30) y nunca llena
    # esos dos por separado, así que reportarlos como "faltantes" sería ruido.
    return [c for c in _CAMPOS_RELEVANTES_PARA_COMPLETITUD if _vacio(vehiculo.get(c))]


# Patente argentina vieja (AAA000) o Mercosur (AA000AA). Sirve para marcar
# como sospechosa una patente que no matchea ninguno de los dos formatos —
# típicamente señal de que el parser agarró un pedazo de otro campo.
_PATENTE_VIEJA = re.compile(r"^[A-Z]{3}\d{3}$")
_PATENTE_MERCOSUR = re.compile(r"^[A-Z]{2}\d{3}[A-Z]{2}$")

# Motor/chasis reales rondan entre 8 y 20 caracteres. Por encima de este
# largo, lo más probable es que dos campos hayan quedado pegados sin
# espacio en el texto original (el caso del Fiat de la flota de La
# Segunda: motor+chasis llegaron como un solo bloque de 30+ caracteres).
_LARGO_MOTOR_CHASIS_SOSPECHOSO = 30


def _patente_formato_valido(patente):
    p = re.sub(r"\s+", "", str(patente or "").strip().upper())
    if not p:
        return True  # patente vacía ya se reporta aparte como campo pendiente
    return bool(_PATENTE_VIEJA.match(p) or _PATENTE_MERCOSUR.match(p))


def _vehiculo_avisos(vehiculo):
    """Chequeos baratos (regex, sin Gemini) para detectar filas que
    probablemente tengan un problema de LECTURA (no de dato faltante):
    patente con formato raro, o motor/chasis con longitud fuera de lo
    normal. No corrige nada solo; da una pista de qué revisar antes de
    pegar en el Excel."""
    avisos = []
    if vehiculo.get("sin_parsear"):
        avisos.append("fila no reconocida por el parser (revisar formato/columnas)")
    if not _patente_formato_valido(vehiculo.get("patente")):
        avisos.append("patente con formato raro")
    for campo, etiqueta in (("motor", "motor"), ("chasis", "chasis")):
        if len(str(vehiculo.get(campo) or "")) > _LARGO_MOTOR_CHASIS_SOSPECHOSO:
            avisos.append(f"{etiqueta} con longitud fuera de lo normal (¿campos pegados?)")
    return avisos


def _fusionar_flota(estado_flota, datos_generales_nuevos, vehiculos_nuevos):
    """Aplica la Sección 21/22/23: completa datos generales, y para cada
    vehículo nuevo busca coincidencia por patente/chasis antes de decidir
    si actualiza uno existente o agrega uno nuevo con el próximo ITEM."""
    datos_generales = estado_flota.setdefault("datos_generales", {})
    for campo, valor in (datos_generales_nuevos or {}).items():
        valor = str(valor or "").strip()
        if valor and _vacio(datos_generales.get(campo)):
            datos_generales[campo] = valor

    vehiculos = estado_flota.setdefault("vehiculos", [])
    tocados = set()
    for nuevo in vehiculos_nuevos or []:
        nuevo = {campo: str(nuevo.get(campo, "") or "").strip() for campo in CAMPOS_VEHICULO}
        if not any(nuevo.values()):
            continue
        coincidencia = next((v for v in vehiculos if _mismo_vehiculo(v, nuevo)), None)
        if coincidencia is not None:
            _fusionar_campos_vehiculo(coincidencia, nuevo)
            tocados.add(coincidencia["item"])
        else:
            item = (max((v["item"] for v in vehiculos), default=0)) + 1
            registro = _vehiculo_nuevo_vacio(item)
            _fusionar_campos_vehiculo(registro, nuevo)
            vehiculos.append(registro)
            tocados.add(item)
    return tocados


_PATRON_ACTUALIZACION_ETIQUETADA = re.compile(
    r"\b(?:la\s+)?(patente|dominio|chapa|chasis|motor|suma\s+asegurada|suma|"
    r"cobertura|a[ñn]o|uso)\s+(?:del|de(?:l)?\s+veh[íi]culo)\s+(\d{1,3})\s+"
    r"(?:es|son|queda|qued[oó])\s*:?\s*(.+?)(?:[.;]|$|\by\b)",
    re.IGNORECASE,
)

_PATRON_ACTUALIZACION_POR_ITEM = re.compile(
    r"\b(?:el|al|vehiculo|veh[íi]culo|unidad)\s*(?:n[úu]mero)?\s*(\d{1,3})\b\s*"
    r"(?:es|son|tiene|tambi[ée]n|queda|qued[oó])?\s*[:\-]?\s*"
    r"(.*?)(?=(?:\s*(?:,|\by\b)\s*(?:el|al|vehiculo|veh[íi]culo|unidad)\s*\d)|[.;]|$)",
    re.IGNORECASE,
)


def _campo_por_etiqueta(etiqueta):
    etiqueta = etiqueta.lower().strip()
    for campo, alias in _ETIQUETAS_CAMPO_NATURAL.items():
        if etiqueta in alias:
            return campo
    return None


def _adivinar_campo_por_valor(valor):
    valor = valor.strip()
    if re.fullmatch(r"[A-Z]{1,2}\s?-?\s?\d{1,3}", valor, re.IGNORECASE):
        return "cobertura"
    if re.fullmatch(r"[A-Z]{2,3}\s?\d{3}\s?[A-Z]{0,2}", valor, re.IGNORECASE):
        return "patente"
    if re.fullmatch(r"\d{4}", valor):
        return "año"
    return None


def _extraer_campo_explicito(resto):
    """Si el texto empieza nombrando el campo ('cobertura C3', 'patente
    AB123CD'), lo separa del valor en vez de guardar la etiqueta pegada al
    dato (evita guardar 'cobertura C3' como si fuera el valor)."""
    m = re.match(
        r"^(patente|dominio|chapa|chasis|motor|suma\s+asegurada|suma|cobertura|a[ñn]o|uso)\s*[:\-]?\s*(.+)$",
        resto,
        re.IGNORECASE,
    )
    if not m:
        return None, resto
    campo = _campo_por_etiqueta(m.group(1))
    return campo, m.group(2).strip()


def _detectar_actualizaciones_naturales(mensaje):
    """Lee frases como 'el 7 tiene C3', 'la patente del 8 es AB123CD' o
    'el 4 es C3 y el 18 también' y devuelve una lista de
    {item, campo, valor}. No inventa: si no puede determinar el campo con
    confianza, descarta esa coincidencia en vez de adivinar mal (Sección 16)."""
    actualizaciones = []
    ultimo_valor = None
    ultimo_campo = None

    for m in _PATRON_ACTUALIZACION_ETIQUETADA.finditer(mensaje):
        etiqueta, item, valor = m.group(1), int(m.group(2)), m.group(3).strip(" .")
        campo = _campo_por_etiqueta(etiqueta)
        if campo and valor:
            actualizaciones.append({"item": item, "campo": campo, "valor": valor})

    if actualizaciones:
        return actualizaciones

    for m in _PATRON_ACTUALIZACION_POR_ITEM.finditer(mensaje):
        item = int(m.group(1))
        resto = (m.group(2) or "").strip(" .:-")
        if not resto or resto.lower() in {"tambien", "también"}:
            if ultimo_valor and ultimo_campo:
                actualizaciones.append({"item": item, "campo": ultimo_campo, "valor": ultimo_valor})
            continue
        campo_explicito, valor_sin_etiqueta = _extraer_campo_explicito(resto)
        if campo_explicito:
            campo, valor = campo_explicito, valor_sin_etiqueta
        else:
            campo, valor = (_adivinar_campo_por_valor(resto) or "cobertura"), resto
        actualizaciones.append({"item": item, "campo": campo, "valor": valor})
        ultimo_valor, ultimo_campo = valor, campo

    return actualizaciones


def _aplicar_actualizaciones_naturales(estado_flota, mensaje):
    actualizaciones = _detectar_actualizaciones_naturales(mensaje)
    if not actualizaciones:
        return set()
    vehiculos = estado_flota.setdefault("vehiculos", [])
    por_item = {v["item"]: v for v in vehiculos}
    tocados = set()
    for cambio in actualizaciones:
        vehiculo = por_item.get(cambio["item"])
        if vehiculo is None:
            continue
        vehiculo[cambio["campo"]] = cambio["valor"]
        tocados.add(cambio["item"])
    return tocados


def _campos_flota_a_datos_generales(campos_flota):
    """interpretar_flota_a_json ya adjunta asegurado/domicilio/localidad/cp
    a cada vehículo; acá los desprendemos para que vivan una sola vez a
    nivel flota (Sección 8), no repetidos por vehículo."""
    vehiculos = campos_flota.get("vehiculos") or []
    datos_generales = {}
    if vehiculos:
        primero = vehiculos[0]
        for campo in ("asegurado", "domicilio", "localidad", "cp"):
            if primero.get(campo):
                datos_generales[campo] = primero[campo]
    return datos_generales


def _dividir_marca_modelo(marca_modelo):
    """El extractor guarda MARCA/MODELO combinado tal como figura en la
    póliza (Sección 30) para no alterar el dato original. La planilla de
    flotas, en cambio, tiene columnas separadas MARCA y MODELO. Acá lo
    partimos SOLO para volcarlo al bloque tabulado: la primera palabra se
    asume marca, el resto modelo. Es una heurística simple a propósito —
    cualquier caso raro se corrige a mano en el bloque antes de pegar."""
    texto = str(marca_modelo or "").strip()
    if not texto:
        return "", ""
    partes = texto.split(" ", 1)
    marca = partes[0]
    modelo = partes[1].strip() if len(partes) > 1 else ""
    return marca, modelo


# Orden EXACTO de columnas de la fila 16 de la planilla de flotas
# (excel/flotas), sólo las columnas que la IA completa. COSTO MENSUAL,
# PREMIO ANUAL y la SUMA ASEGURADA de "COMPETENCIA" son de cotización
# manual y no se tocan.
_COLUMNAS_TSV_FLOTA = (
    "ITEM", "MARCA", "MODELO", "AÑO", "COBERTURA SOLICITADA", "USO",
    "MOTOR", "CHASIS", "ACCESORIO", "PATENTE", "SUMA ASEGURADA",
)


def _normalizar_suma_asegurada(valor):
    """Dos frentes de póliza de compañías distintas casi nunca escriben la
    suma asegurada igual: una la pone como "$ 15.000.000,00", otra como
    "ARS 15000000", otra sin símbolo. El extractor guarda el valor tal cual
    vino (para no perder el dato original), pero para el bloque que se pega
    en Excel conviene sacarle cualquier símbolo de moneda y espacio interno:
    así entra como número real sin importar de qué compañía salió, en vez
    de depender de que cada frente use el mismo formato que La Segunda."""
    texto = str(valor or "").strip()
    if not texto:
        return ""
    texto = re.sub(r"(?i)^\s*(ars|usd|u\$s|\$)\s*", "", texto)
    return re.sub(r"\s+", "", texto)


def _armar_bloque_tsv_flota(vehiculos):
    """Arma el bloque de texto separado por TABULADORES (uno por vehículo),
    en el mismo orden de columnas que la fila 16 de la planilla de flotas.
    Este bloque es lo que el usuario copia y pega directo en Excel a partir
    de la fila del ITEM correspondiente — Excel interpreta cada \\t como un
    salto de columna, así que las celdas quedan alineadas solas."""
    lineas = []
    for indice, vehiculo in enumerate(vehiculos, start=1):
        if not _vehiculo_guardable(vehiculo):
            continue
        # Si el extractor de origen ya calculó marca/modelo separados (caso
        # La Segunda, con marcas de dos palabras como "M. BENZ"), se
        # respetan tal cual. Sólo se recalcula con la heurística genérica
        # cuando no vinieron ya resueltos (formato con etiquetas).
        if vehiculo.get("marca") or vehiculo.get("modelo"):
            marca, modelo = vehiculo.get("marca", ""), vehiculo.get("modelo", "")
        else:
            marca, modelo = _dividir_marca_modelo_flota(vehiculo.get("marca_modelo"))
        fila = {
            "ITEM": str(vehiculo.get("item") or indice),
            "MARCA": marca,
            "MODELO": modelo,
            "AÑO": vehiculo.get("año", ""),
            "COBERTURA SOLICITADA": vehiculo.get("cobertura", ""),
            "USO": vehiculo.get("uso", ""),
            "MOTOR": vehiculo.get("motor", ""),
            "CHASIS": vehiculo.get("chasis", ""),
            "ACCESORIO": "",
            "PATENTE": vehiculo.get("patente", ""),
            "SUMA ASEGURADA": _normalizar_suma_asegurada(vehiculo.get("suma_asegurada", "")),
        }
        valores = [str(fila[col] or "").strip() for col in _COLUMNAS_TSV_FLOTA]
        lineas.append("\t".join(valores))
    return "\n".join(lineas)


def _listar_numeros(items, conector="y"):
    """Arma "16, 30, 32 y 36" a partir de una lista de números/strings."""
    items = [str(i) for i in items]
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    return ", ".join(items[:-1]) + f" {conector} " + items[-1]


def _resumen_estado_flota(estado_flota, items_tocados_ahora, con_error, es_primera_vez, aviso_conteo=None):
    """Arma el mensaje que Sofia le muestra al usuario después de procesar
    una póliza de flota. /flota no escribe el archivo real, sólo arma el
    bloque tabulado; acá se cuenta en lenguaje simple qué se encontró, qué
    conviene revisar y qué falta, sin tecnicismos ni jerga interna."""
    vehiculos = estado_flota.get("vehiculos") or []
    total = len(vehiculos)
    tocados_ahora = len(items_tocados_ahora) if items_tocados_ahora else 0
    pendientes = [v for v in vehiculos if _campos_pendientes_vehiculo(v) and _vehiculo_guardable(v)]
    sospechosos = [v for v in vehiculos if _vehiculo_guardable(v) and _vehiculo_avisos(v)]

    partes = []

    if tocados_ahora:
        if es_primera_vez:
            palabra = "vehículo" if total == 1 else "vehículos"
            partes.append(f"La póliza se cargó y encontré {total} {palabra}.")
        else:
            palabra = "vehículo" if tocados_ahora == 1 else "vehículos"
            partes.append(f"Sumé {tocados_ahora} {palabra} más. Ahora hay {total} en total.")
    elif total and es_primera_vez:
        partes.append(f"Encontré {total} vehículos, pero todavía no tienen datos suficientes para armar el bloque.")
    elif not total:
        partes.append(
            "Empecé a cargar la flota. Pasame los datos de la póliza y los vehículos, "
            "todos juntos o de a poco, como te resulte más cómodo."
        )

    if con_error:
        singular = len(con_error) == 1
        articulo, palabra = ("El", "vehículo") if singular else ("Los", "vehículos")
        numeros = _listar_numeros([i for i, _ in con_error])
        verbo = "tuvo" if singular else "tuvieron"
        partes.append(f"{articulo} {palabra} {numeros} {verbo} un problema puntual, pero no afectó al resto.")

    if sospechosos:
        singular = len(sospechosos) == 1
        articulo, palabra = ("El", "vehículo") if singular else ("Los", "vehículos")
        numeros = _listar_numeros([v["item"] for v in sospechosos])
        verbo = "conviene revisarlo" if singular else "conviene revisarlos"
        partes.append(
            f"{articulo} {palabra} {numeros} {verbo} porque algunos datos parecen estar mal separados."
        )

    if pendientes:
        campos_faltantes = set()
        for v in pendientes:
            campos_faltantes.update(_campos_pendientes_vehiculo(v))
        if len(pendientes) <= 3:
            singular = len(pendientes) == 1
            articulo, palabra = ("Al", "vehículo") if singular else ("A los", "vehículos")
            numeros = _listar_numeros([v["item"] for v in pendientes])
            verbo = "le falta" if singular else "les falta"
            partes.append(f"{articulo} {palabra} {numeros} todavía {verbo} algún dato. Pasámelo cuando lo tengas y actualizo el bloque.")
        else:
            ejemplo = sorted(campos_faltantes)[0] if campos_faltantes else "algún dato"
            partes.append(
                f"Además, hay {len(pendientes)} vehículos a los que todavía les falta "
                f"algún dato (por ejemplo, {ejemplo}). Pasámelos cuando los tengas y actualizo el bloque."
            )

    if not partes:
        partes.append(
            "Flota iniciada. Pasame los datos generales de la póliza y/o los vehículos "
            "(podés mandarlos todos juntos, en tandas, o uno por uno)."
        )

    return " ".join(partes)


def _flota_procesar_turno(chat_id, mensaje, contexto_pdf_adjunto):
    """Punto central del flujo /flota persistente. Devuelve
    (respuesta, True, bloque_tsv) si el mensaje fue absorbido por la tarea
    de flota, o (None, False, None) si no tiene nada que ver y debe seguir
    el flujo normal (Gemini / otros comandos).

    IMPORTANTE (rediseño): /flota YA NO escribe directo en excel/flotas.xlsx.
    Sólo interpreta el/los frente(s) de póliza pegados, acumula el contexto
    de la flota en `flotas_activas` (para poder seguir sumando vehículos en
    mensajes sucesivos) y devuelve un bloque de texto separado por
    TABULADORES con TODOS los vehículos detectados hasta el momento, listo
    para copiar y pegar manualmente en el Excel. El usuario revisa el
    bloque, corrige lo que haga falta y pega — así ningún error de lectura
    llega al Excel real sin que se vea antes."""

    es_comando_flota = bool(re.match(r"^/flota\b", mensaje, re.IGNORECASE))
    estado_flota = _flota_obtener(chat_id)
    es_primera_vez = estado_flota is None

    if not es_comando_flota and estado_flota is None:
        return None, False, None

    if not es_comando_flota and estado_flota is not None and estado_flota.get("estado") == "completada":
        return None, False, None

    if es_comando_flota:
        texto_flota = re.sub(r"^/flota\s*", "", mensaje, count=1, flags=re.IGNORECASE).strip()
    else:
        texto_flota = mensaje.strip()

    if re.match(r"^(termin(a|ar|amos|é)|listo|finaliza(r)?|cerrar\s+flota|flota\s+completa)\b", texto_flota, re.IGNORECASE) and estado_flota:
        vehiculos = estado_flota.get("vehiculos") or []
        bloque_final = _armar_bloque_tsv_flota(vehiculos)
        _flota_guardar(chat_id, "completada", estado_flota.get("libro_id", "2"), estado_flota.get("datos_generales", {}), vehiculos)
        return (
            f"Flota cerrada con {len(vehiculos)} vehículo(s). Te dejo el bloque completo abajo para "
            "copiar y pegar en el Excel. Si aparece más información después, escribí /flota de nuevo "
            "y la sumo.",
            True,
            bloque_final or None,
        )

    if estado_flota is None:
        estado_flota = {"estado": "nueva", "libro_id": "2", "datos_generales": {}, "vehiculos": []}

    fuente = texto_flota
    if contexto_pdf_adjunto:
        fuente = f"{texto_flota}\n\n{contexto_pdf_adjunto}".strip() if texto_flota else contexto_pdf_adjunto

    tocados = set()

    # Las correcciones/updates cortos en lenguaje natural ("el 7 es C3") se
    # intentan primero y son baratos (regex, sin llamar a nada externo). El
    # parser de "volcado de vehículos" (interpretar_flota_a_json, que puede
    # caer a Gemini si el texto no trae etiquetas explícitas) sólo se invoca
    # cuando el mensaje realmente parece traer datos de póliza/vehículos, no
    # en cada corrección puntual — para no gastar una llamada a Gemini de
    # más ni arriesgarse a que reinterprete mal una frase corta.
    if not es_comando_flota:
        tocados |= _aplicar_actualizaciones_naturales(estado_flota, mensaje)

    parece_volcado_vehiculos = (
        es_comando_flota
        or bool(contexto_pdf_adjunto)
        or re.search(r"DESCRIPCI[ÓO]N\s+DEL\s+VEH[ÍI]CULO", fuente, re.IGNORECASE)
        or len(fuente) > 200
    )

    aviso_conteo = None
    if fuente and not tocados and parece_volcado_vehiculos:
        try:
            campos_flota = interpretar_flota_a_json(fuente)
            vehiculos_nuevos = campos_flota.get("vehiculos") or []
        except Exception as error:
            print("ERROR PROCESANDO /FLOTA:", error)
            vehiculos_nuevos = []
            campos_flota = {}
        if vehiculos_nuevos:
            datos_generales_nuevos = _campos_flota_a_datos_generales(campos_flota)
            tocados |= _fusionar_flota(estado_flota, datos_generales_nuevos, vehiculos_nuevos)
        aviso_conteo = campos_flota.get("aviso_conteo")

    if not tocados and not fuente and es_comando_flota:
        # "/flota" pelado: si es la primera vez, arrancamos la tarea. Si ya
        # había una flota activa, sólo informamos el estado (Sección 40).
        estado_flota["estado"] = estado_flota.get("estado") or "nueva"
        _flota_guardar(chat_id, estado_flota["estado"], estado_flota.get("libro_id", "2"), estado_flota["datos_generales"], estado_flota["vehiculos"])
        if es_primera_vez:
            return (
                "Entendido, arranco una flota nueva. Pasame los datos generales de la póliza "
                "(asegurado, número de póliza, compañía, etc.) y los vehículos — en el orden y de "
                "a la cantidad que te resulte más cómoda, todos juntos o de a tandas. Cuando tenga "
                "algo, te devuelvo el bloque tabulado para copiar y pegar en el Excel.",
                True,
                None,
            )
        bloque = _armar_bloque_tsv_flota(estado_flota["vehiculos"])
        return _resumen_estado_flota(estado_flota, [], [], es_primera_vez=False), True, (bloque or None)

    if not tocados and not es_comando_flota:
        # No era ni un dato de flota ni una actualización reconocible:
        # dejamos pasar el mensaje al flujo normal (puede ser una pregunta
        # sin relación, Sección 27).
        return None, False, None

    # Ya NO se escribe directo en excel/flotas.xlsx (Sección rediseño): se
    # arma el bloque tabulado con TODOS los vehículos acumulados hasta acá
    # para que el usuario lo copie y pegue a mano, revisando antes de tocar
    # el Excel real.
    estado_flota["estado"] = "en_progreso" if estado_flota["vehiculos"] else "nueva"
    _flota_guardar(
        chat_id,
        estado_flota["estado"],
        estado_flota.get("libro_id", "2"),
        estado_flota["datos_generales"],
        estado_flota["vehiculos"],
    )

    bloque = _armar_bloque_tsv_flota(estado_flota["vehiculos"])
    respuesta = _resumen_estado_flota(estado_flota, tocados, [], es_primera_vez, aviso_conteo=aviso_conteo)
    return respuesta, True, (bloque or None)


# ==========================================================
# TANDA 4 — PÓLIZA → ASEGURADO (flujo nuevo, separado de /flota)
# ==========================================================
#
# Este flujo NO toca /flota: no comparte estado (flotas_activas), no
# reutiliza su deduplicación ni sus mensajes. Lo que se reutiliza de
# /flota es el aprendizaje de cómo interpretar un frente de póliza
# (mismo patrón de "Gemini con instrucción estricta, salida JSON, sin
# inventar campos vacíos") — no su código de flota en sí.
#
# El comando es "/alta" y procesa UNA póliza individual (no una flota)
# para dejar lista la propuesta de alta de asegurado. Es de un solo turno:
# no arrastra estado entre mensajes como sí hace /flota.

_COLUMNAS_ALTA_ASEGURADO = (
    "ASEGURADO", "NUMERO", "VEHICULO", "PATENTE", "ENVIOS YA", "COMPAÑIA",
    "MEDIO DE PAGO", "CODIGO POSTAL", "EMITIDO DÍA:", "IMPORTE APROX", "TELEFONO",
)

# Mensaje que arma solo el propio /api/chat cuando el usuario adjunta un PDF
# sin escribir nada (botón o drag&drop). Se usa como referencia para saber
# si, al detectar automáticamente una póliza (más abajo), conviene ignorar
# ese texto de relleno o respetar algo que el usuario sí haya escrito.
_MENSAJE_PDF_POR_DEFECTO = "Analizá el PDF que acabo de adjuntar y explicame de qué trata."

# Patente argentina: formato viejo (3 letras + 3 números) o mercosur
# (2 letras + 3 números + 2 letras).
_PATENTE_REGEX_ALTA = re.compile(r"\b[A-Z]{2}\d{3}[A-Z]{2}\b|\b[A-Z]{3}\d{3}\b")


def _pdf_parece_poliza_individual_para_alta(contexto_pdf_adjunto):
    """Tanda 8 — drag&drop sin comando: decide si un PDF recién adjuntado
    se parece al frente de UNA póliza individual (un asegurado, un
    vehículo), para poder ofrecer el alta automáticamente sin que el
    usuario tenga que escribir "/alta".

    Es deliberadamente conservadora: si hay señales de que en realidad es
    una flota (más de una patente, numeración de ítems), no dispara nada
    acá y el mensaje sigue el camino normal (o /flota, si corresponde)."""
    texto = str(contexto_pdf_adjunto or "")
    if not texto.strip():
        return False

    patentes = {m.group(0).upper() for m in _PATENTE_REGEX_ALTA.finditer(texto)}
    if len(patentes) >= 2:
        return False  # más de un vehículo → parece flota, no alta individual

    if len(re.findall(r"(?:^|\n)\s*(?:ITEM|ÍTEM|N[°º]\s*\d+)\b", texto, re.IGNORECASE)) >= 2:
        return False  # listado numerado → parece flota

    return bool(re.search(
        r"\b(P[ÓO]LIZA|ASEGURADO|PREMIO|PRIMA|COBERTURA|VIGENCIA|CERTIFICADO)\b",
        texto, re.IGNORECASE,
    ))

# Mismas compañías que ya tiene el sistema en los accesos directos
# (Sección /flota-adyacente): se reutilizan los nombres, no se inventan
# nuevas. Sirven como pista para detectar la compañía en el texto de la
# póliza cuando no viene con una etiqueta explícita "COMPAÑIA:".
_COMPANIAS_CONOCIDAS_ALTA = [nombre for nombre, _url in CIAS_LINKS]


def _detectar_compania_poliza(texto):
    """Sólo confirma la compañía si aparece de forma razonablemente clara
    en el propio texto de la póliza. No adivina a partir del nombre de
    archivo ni de suposiciones: si no hay evidencia, se deja vacío."""
    texto_norm = str(texto or "")
    m = re.search(r"COMPA[ÑN][IÍ]A\s*:?\s*([A-ZÁÉÍÓÚÑ .]{3,40})", texto_norm, re.IGNORECASE)
    if m:
        candidato = re.sub(r"\s+", " ", m.group(1)).strip(" .")
        if candidato:
            return candidato
    for nombre in _COMPANIAS_CONOCIDAS_ALTA:
        if re.search(rf"\b{re.escape(nombre)}\b", texto_norm, re.IGNORECASE):
            return nombre
    return ""


def interpretar_poliza_asegurado_a_json(texto):
    """Extrae de un frente de póliza individual los datos para proponer el
    alta de un asegurado. A diferencia de interpretar_flota_a_json, esto
    NO es una lista de vehículos de una flota: es un único asegurado con
    un único vehículo, con las columnas exactas que pide la Tanda 4.

    Reglas que Gemini debe respetar (idénticas a la instrucción, para no
    inventar datos):
    - "envios_ya" NUNCA se completa acá: queda para carga manual.
    - "medio_pago" sólo si hay evidencia clara en el texto; si no, vacío.
    - "codigo_postal" sólo si aparece explícito o es inequívoco; si no, vacío.
    - "emitido" es la FECHA DE EMISIÓN de la póliza (no la fecha actual).
    - "importe_aprox" sale del PREMIO (precio final), no de una suma propia.
    """
    texto = str(texto or "").replace("\r", "")
    if not texto.strip():
        return {}

    from servicios_ia import obtener_cliente_gemini, MODELOS_GEMINI
    from google.genai import types

    cliente = obtener_cliente_gemini()
    if cliente is None:
        raise RuntimeError("La IA todavía no está configurada. Falta GEMINI_API_KEY.")

    instruccion = """
Vas a recibir el texto crudo del frente de una póliza de seguro individual
(un solo asegurado, un solo vehículo) de una compañía argentina, pegado sin
formato. Tu tarea es identificar un conjunto fijo de datos para proponer el
alta de ese asegurado en una planilla.

Devolvé ÚNICAMENTE JSON válido, sin markdown ni texto fuera del JSON, con
esta estructura exacta:
{
  "asegurado": "",
  "telefono": "",
  "numero": "",
  "vehiculo": "",
  "patente": "",
  "compania": "",
  "medio_pago": "",
  "codigo_postal": "",
  "emitido": "",
  "premio": ""
}

REGLAS ESTRICTAS (no las rompas aunque el texto sea ambiguo):

1. "asegurado": el nombre de la persona o razón social asegurada, tal como
   figura en la póliza. Si no está claro, dejalo vacío.
2. "telefono": el teléfono de contacto del asegurado, SOLO si figura explícitamente en la póliza. Si no aparece, dejalo vacío. NUNCA uses el DNI ni el número de póliza como teléfono.
3. "numero": el número de póliza.
4. "vehiculo": marca/modelo o descripción del vehículo tal como figura.
5. "patente": la patente, si figura.
6. "compania": la compañía de seguros que emitió la póliza.
7. "medio_pago": SOLO si el texto permite identificarlo con evidencia clara
   (por ejemplo dice explícitamente "cuponera", "CBU", "tarjeta de crédito").
   NUNCA lo asumas a partir de qué compañía es. Si no hay evidencia clara,
   dejalo vacío. No adivines.
8. "codigo_postal": SOLO si aparece explícito en el texto o se puede
   identificar sin ambigüedad a partir de los propios datos de la póliza.
   Si no aparece, dejalo vacío. No lo inventes ni lo busques externamente.
9. "emitido": la FECHA DE EMISIÓN de la póliza (no la fecha de hoy, no la
   fecha de vigencia, no el vencimiento). Formato DD/MM/AAAA. Si no hay una
   fecha de emisión confiable, dejalo vacío.
10. "premio": el importe del PREMIO (el precio final del seguro), tal como
   figura. Si la póliza usa otra palabra que equivale claramente al premio
   final, usá ese valor. NUNCA sumes conceptos, ni calcules un precio
   nuevo, ni estimes en base a la cobertura. Si no hay un premio
   identificable con certeza, dejalo vacío.

No inventes ningún dato. Un campo que no aparece con certeza en el texto
queda vacío ("") — nunca "N/D", "no informado" ni similares.
"""
    ultimo_error = None
    for modelo in MODELOS_GEMINI:
        try:
            config = types.GenerateContentConfig(
                temperature=0,
                max_output_tokens=1500,
                response_mime_type="application/json",
                system_instruction=instruccion.strip(),
            )
            respuesta = cliente.models.generate_content(
                model=modelo,
                contents=texto.strip(),
                config=config,
            )
            bruto = str(getattr(respuesta, "text", "") or "").strip()
            if not bruto:
                raise ValueError("Gemini no devolvió JSON.")
            datos = json.loads(bruto)
            if not isinstance(datos, dict):
                raise ValueError("Gemini no devolvió un objeto JSON.")

            def limpio(clave):
                valor = re.sub(r"\s+", " ", str(datos.get(clave, "") or "")).strip()
                return valor

            return {
                "asegurado": limpio("asegurado"),
                "telefono": limpio("telefono"),
                "numero": limpio("numero"),
                "vehiculo": limpio("vehiculo"),
                "patente": limpio("patente"),
                "compania": normalizar_compania(limpio("compania") or _detectar_compania_poliza(texto)),
                "medio_pago": limpio("medio_pago"),
                "codigo_postal": limpio("codigo_postal"),
                "emitido": limpio("emitido"),
                "premio": limpio("premio"),
            }
        except Exception as error:
            ultimo_error = error
            print("ERROR GEMINI /ALTA:", modelo, ":", error)
    raise RuntimeError(f"No pude interpretar la póliza: {ultimo_error}")


def _propuesta_alta_a_columnas(propuesta):
    """Arma el dict con las columnas EXACTAS (mismo orden y mismos nombres)
    que pide la Tanda 4/5. ENVIOS YA queda siempre vacío: es manual."""
    return {
        "ASEGURADO": propuesta.get("asegurado", ""),
        "NUMERO": propuesta.get("numero", ""),
        "VEHICULO": propuesta.get("vehiculo", ""),
        "PATENTE": propuesta.get("patente", ""),
        "ENVIOS YA": "",
        "COMPAÑIA": propuesta.get("compania", ""),
        "MEDIO DE PAGO": propuesta.get("medio_pago", ""),
        "CODIGO POSTAL": propuesta.get("codigo_postal", ""),
        "EMITIDO DÍA:": propuesta.get("emitido", ""),
        "IMPORTE APROX": propuesta.get("premio", ""),
        "TELEFONO": propuesta.get("telefono", ""),
    }


def _resumen_alta_asegurado(columnas):
    """Mensaje en lenguaje simple (Tanda 1) que resume lo que Sofia
    encontró en la póliza, sin tecnicismos ni listar cada campo vacío."""
    asegurado = columnas.get("ASEGURADO") or ""
    compania = columnas.get("COMPAÑIA") or ""
    numero = columnas.get("NUMERO") or ""

    partes = []
    if asegurado and compania:
        partes.append(f"Leí la póliza: es de {asegurado}, en {compania}.")
    elif asegurado:
        partes.append(f"Leí la póliza: es de {asegurado}.")
    else:
        partes.append("Leí la póliza, pero no pude identificar con seguridad a nombre de quién está.")

    if numero:
        partes.append(f"El número de póliza es {numero}.")

    faltan = [
        etiqueta for etiqueta, clave in (
            ("la compañía", "COMPAÑIA"),
            ("el vehículo", "VEHICULO"),
            ("la patente", "PATENTE"),
        )
        if not columnas.get(clave)
    ]
    if faltan:
        partes.append(f"No encontré {', '.join(faltan)}; si lo tenés, pasámelo y lo agrego.")

    if not columnas.get("MEDIO DE PAGO"):
        partes.append("El medio de pago no queda claro en la póliza, así que lo dejé vacío para que lo completes vos.")
    if not columnas.get("CODIGO POSTAL"):
        partes.append("Tampoco encontré el código postal.")
    if not columnas.get("IMPORTE APROX"):
        partes.append("No encontré un premio identificable, así que el importe quedó vacío.")

    partes.append("Envíos Ya lo dejo vacío siempre, para que lo cargues a mano.")

    regla_pago = calcular_regla_pago(columnas.get("COMPAÑIA"), columnas.get("MEDIO DE PAGO"))
    if regla_pago:
        partes.append(regla_pago["descripcion"])

    return " ".join(partes)


def _procesar_alta_asegurado(mensaje, contexto_pdf_adjunto, automatico=False):
    """Punto de entrada del comando /alta. Devuelve (respuesta, True,
    propuesta_dict) si el mensaje fue absorbido por este flujo, o
    (None, False, None) si no corresponde y debe seguir el flujo normal.

    Este flujo es de UN SOLO TURNO (a diferencia de /flota): no arrastra
    contexto entre mensajes.

    `automatico=True` (Tanda 8) indica que "/alta" no lo escribió el
    usuario: se agregó solo porque el PDF adjunto se detectó como una
    póliza individual. Sólo cambia la primera frase de la respuesta, para
    que quede claro que Sofia lo reconoció sola."""
    if not re.match(r"^/alta\b", mensaje, re.IGNORECASE):
        return None, False, None

    texto_comando = re.sub(r"^/alta\s*", "", mensaje, count=1, flags=re.IGNORECASE).strip()
    fuente = texto_comando
    if contexto_pdf_adjunto:
        fuente = f"{texto_comando}\n\n{contexto_pdf_adjunto}".strip() if texto_comando else contexto_pdf_adjunto

    if not fuente:
        return (
            "Para dar de alta un asegurado desde una póliza, adjuntame el PDF "
            "(con el clip o arrastrándolo al chat) y escribí /alta, o pegame "
            "el texto del frente de póliza después de /alta.",
            True,
            None,
        )

    try:
        propuesta = interpretar_poliza_asegurado_a_json(fuente)
    except Exception as error:
        print("ERROR PROCESANDO /ALTA:", error)
        return (
            "No pude leer bien esa póliza para armar el alta. Probá adjuntarla "
            "de nuevo o pegar el texto del frente.",
            True,
            None,
        )

    columnas = _propuesta_alta_a_columnas(propuesta)
    respuesta = _resumen_alta_asegurado(columnas)
    if automatico:
        respuesta = "Reconocí que me pasaste una póliza, así que te dejo directamente la propuesta de alta. " + respuesta
    return respuesta, True, columnas


# ==========================================================
# TANDA 5 — DOS SALIDAS: TABULADO O GUARDAR EN EXCEL
# ==========================================================
#
# A partir de la propuesta que arma la Tanda 4 (columnas), se ofrecen dos
# caminos SIN volver a interpretar la póliza:
#   - Tabulado: una fila lista para copiar/pegar, mismo orden y mismos
#     nombres de columna que pide la instrucción (los campos que faltan
#     quedan vacíos, nunca "N/D").
#   - Guardar en Excel: reutiliza el mecanismo ya existente de
#     /guardar asegurado (mismo formulario, misma validación, misma
#     escritura en el Excel de Asegurados) en vez de duplicar esa lógica.

def _armar_tabulado_alta(columnas):
    """Fila TSV en el orden EXACTO de _COLUMNAS_ALTA_ASEGURADO, lista para
    pegar en Excel. Un campo sin dato queda vacío, nunca "N/D" ni similar."""
    valores = [str((columnas or {}).get(col, "") or "") for col in _COLUMNAS_ALTA_ASEGURADO]
    return "\t".join(valores)


def _alta_a_campos_guardar_asegurado(columnas):
    """Convierte la propuesta de alta (columnas de la Tanda 4/5) a los
    mismos campos que ya usa /guardar asegurado para el Excel de
    Asegurados (Libro 1), para reutilizar ese formulario y esa función de
    guardado sin crear un camino paralelo.

    COMPAÑIA -> CIA y CODIGO POSTAL -> CP porque son los nombres reales de
    esas columnas en ese Excel. EMITIDO DÍA e IMPORTE APROX todavía no
    tienen columna en esa planilla, así que no viajan acá: quedan
    disponibles en el Tabulado."""
    columnas = columnas or {}
    return {
        "LIBRO_ID": "1",
        "ASEGURADO": columnas.get("ASEGURADO", ""),
        "NUMERO": columnas.get("NUMERO", ""),
        "VEHICULO": columnas.get("VEHICULO", ""),
        "PATENTE": columnas.get("PATENTE", ""),
        "ENVIOS YA": "",
        "CIA": normalizar_compania(columnas.get("COMPAÑIA", "")),
        "MEDIO DE PAGO": columnas.get("MEDIO DE PAGO", ""),
        "CP": columnas.get("CODIGO POSTAL", ""),
        "MAIL": "",
        "TELEFONO": columnas.get("TELEFONO", ""),
    }


# ==========================================================
# TANDA 6 — REGLAS DE NEGOCIO DE PAGOS
# ==========================================================
#
# Con COMPAÑIA + MEDIO DE PAGO (los mismos datos que ya deja armados la
# Tanda 4) se puede anticipar cómo funciona el pago: si el precio queda
# fijo por varios meses, si hay cupones físicos o si el importe cambia
# mes a mes. Son sólo las primeras reglas conocidas; el diccionario está
# pensado para sumar compañías y medios de pago nuevos sin tocar la
# lógica que los aplica.
#
# Si falta la compañía, falta el medio de pago, o la combinación todavía
# no tiene una regla cargada, no se inventa nada: no se aplica ninguna
# regla y listo.

def _normalizar_texto_pago(valor):
    """Deja el texto en mayúsculas y sin acentos, para poder comparar
    compañías y medios de pago sin importar cómo los haya escrito Gemini
    o el usuario (con o sin tildes, mayúsculas, espacios de más)."""
    texto = unicodedata.normalize("NFKD", str(valor or ""))
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", texto.upper()).strip()


def _normalizar_compania_pago(compania):
    """Usa la misma normalización maestra que el resto de OficinaIA."""
    return normalizar_compania(compania)


def _normalizar_medio_pago(medio_pago):
    """Agrupa las formas habituales de escribir cada medio de pago en una
    clave única. Si no reconoce ninguna, devuelve vacío: mejor no aplicar
    ninguna regla que aplicar una equivocada."""
    texto = _normalizar_texto_pago(medio_pago)
    if not texto:
        return ""
    if "CUPON" in texto:
        return "CUPONERA"
    if "CBU" in texto or "DEBITO" in texto or "TRANSFERENCIA" in texto:
        return "CBU"
    if "CREDITO" in texto or "TARJETA" in texto:
        return "CREDITO"
    return ""


# Reglas conocidas hasta ahora. Para agregar una compañía o medio de pago
# nuevo alcanza con sumar una entrada acá, sin tocar el resto del sistema.
_REGLAS_PAGO = {
    ("ATM", "CUPONERA"): {
        "precio_fijo": True,
        "meses_precio_fijo": 3,
        "cantidad_cupones": 3,
        "cupones_fisicos": True,
        "descripcion": (
            "Con ATM y cuponera el precio se mantiene fijo 3 meses. "
            "Hay 3 cupones, uno para cada mes, y se pueden pagar físicamente."
        ),
    },
    ("ATM", "CBU"): {
        "precio_fijo": False,
        "meses_precio_fijo": None,
        "cantidad_cupones": 0,
        "cupones_fisicos": False,
        "descripcion": "Con ATM y CBU el precio puede cambiar todos los meses. No hay cupones.",
    },
    ("ATM", "CREDITO"): {
        "precio_fijo": False,
        "meses_precio_fijo": None,
        "cantidad_cupones": 0,
        "cupones_fisicos": False,
        "descripcion": "Con ATM y crédito el precio puede cambiar todos los meses. No hay cupones.",
    },
    ("AGS", "CUPONERA"): {
        "precio_fijo": True,
        "meses_precio_fijo": 4,
        "cantidad_cupones": 4,
        "cupones_fisicos": True,
        "descripcion": (
            "Con AGS y cuponera el precio se mantiene fijo 4 meses. "
            "Hay 4 cupones, uno para cada mes."
        ),
    },
}


def calcular_regla_pago(compania, medio_pago):
    """Devuelve la regla de pago que corresponde a esta combinación de
    compañía + medio de pago, o None si falta alguno de los dos datos o si
    todavía no hay una regla cargada para esa combinación. Nunca supone
    una regla que no esté explícitamente definida arriba."""
    compania_norm = _normalizar_compania_pago(compania)
    medio_norm = _normalizar_medio_pago(medio_pago)
    if not compania_norm or not medio_norm:
        return None
    regla = _REGLAS_PAGO.get((compania_norm, medio_norm))
    if not regla:
        return None
    return dict(regla)


def _normalizar_patente(valor):
    """Deja sólo letras/números en mayúscula, para poder comparar patentes
    sin importar espacios, guiones o mayúsculas/minúsculas."""
    return re.sub(r"[^A-Z0-9]", "", str(valor or "").upper())


def _normalizar_telefono(valor):
    """Envíos Ya no acepta espacios ni guiones en el teléfono: se dejan
    sólo los dígitos."""
    return re.sub(r"\D", "", str(valor or ""))


def _buscar_asegurado_por_patente(patente_buscada, libro_id="1"):
    """Busca en el Excel interno (asegurados) la fila cuya columna PATENTE
    coincide con la patente pedida, comparando sin espacios/guiones/mayúsculas.
    Devuelve un dict {encabezado: valor} de esa fila, o None si no aparece."""
    patente_norm = _normalizar_patente(patente_buscada)
    if not patente_norm:
        return None
    datos = leer_excel_interno(libro_id)
    filas = datos.get("filas") or []
    if not filas:
        return None
    encabezados = filas[0]
    indices = {
        _normalizar_encabezado(encabezado): i
        for i, encabezado in enumerate(encabezados)
        if _normalizar_encabezado(encabezado)
    }
    indice_patente = indices.get(_normalizar_encabezado("PATENTE"))
    if indice_patente is None:
        return None
    for fila in filas[1:]:
        valor_patente = fila[indice_patente] if indice_patente < len(fila) else ""
        if _normalizar_patente(valor_patente) == patente_norm:
            return {
                encabezado: (fila[i] if i < len(fila) else "")
                for encabezado, i in (
                    (encabezados[i], i) for i in range(len(encabezados))
                )
            }
    return None


def _armar_texto_envios_ya(datos_asegurado):
    """Arma el bloque de texto listo para pegar en Envíos Ya a partir de un
    dict de campos del asegurado (aceptando tanto encabezados de Excel como
    las claves canónicas usadas en /guardar asegurado)."""
    def obtener(*claves):
        for clave in claves:
            for k, v in datos_asegurado.items():
                if _normalizar_encabezado(k) == _normalizar_encabezado(clave):
                    if str(v or "").strip():
                        return str(v).strip()
        return ""

    nombre = obtener("ASEGURADO", "nombre asegurado")
    telefono = _normalizar_telefono(obtener("TELEFONO"))
    vehiculo = obtener("VEHICULO", "marca_modelo", "marca/modelo")
    patente = obtener("PATENTE", "dominio", "chapa").upper()
    cia = obtener("CIA", "compañia", "compania")

    aviso_telefono = ""
    if telefono and len(telefono) != 10:
        aviso_telefono = f" (ojo: tiene {len(telefono)} dígitos, revisá que sea correcto)"

    return (
        f"NOMBRE Y APELLIDO: {nombre}\n"
        f"TELEFONO: {telefono}{aviso_telefono}\n"
        f"VEHICULO: {vehiculo}\n"
        f"PATENTE: {patente}\n"
        f"COMPAÑIA: {cia}"
    )


def _parsear_comando_envios_ya(mensaje):
    """Parsea /envios ya (patente) o /envios ya patente, sin depender de
    Gemini. Devuelve la patente pedida, o None si el mensaje no es este
    comando."""
    texto = str(mensaje or "").strip()
    m = re.match(r"^/envios\s+ya\b\s*(.*)$", texto, re.IGNORECASE)
    if not m:
        return None
    resto = m.group(1).strip()
    resto = re.sub(r"^\(([^)]*)\)$", r"\1", resto).strip()
    resto = resto.strip("'\" ")
    if not resto:
        return {"error": "Usá el formato /envios ya (patente), indicando la patente del vehículo."}
    return {"patente": resto}


def _parsear_comando_guardar_asegurado(mensaje):
    """

    Parsea el comando explícito /guardar asegurado sin depender de Gemini.

    Formato principal:
      /guardar asegurado (asegurado) (numero) (vehiculo) (patente) (cia)
      (medio de pago) (cp) (mail) (telefono opcional) [1|2]

    También acepta los mismos campos separados por comas. ENVIOS YA no forma
    parte del comando corto; puede completarse luego en la propuesta.
    """
    texto = str(mensaje or "").strip()
    patron = re.compile(r"^/guardar\s+asegurado\b", re.IGNORECASE)
    if not patron.match(texto):
        return None

    resto = patron.sub("", texto, count=1).strip()
    campos = (
        "ASEGURADO",
        "NUMERO",
        "VEHICULO",
        "PATENTE",
        "CIA",
        "MEDIO DE PAGO",
        "CP",
        "MAIL",
        "TELEFONO",
    )

    libro_id = "1"
    sufijo = re.search(r"(?:,\s*|\s+)([12])\s*$", resto)
    if sufijo:
        libro_id = sufijo.group(1)
        resto = resto[:sufijo.start()].rstrip(" ,")

    valores = re.findall(r"\(([^)]*)\)", resto)
    if valores:
        if len(valores) > len(campos):
            return {"error": "El comando tiene más campos de los esperados."}
        valores = [v.strip() for v in valores]
    elif "," in resto:
        valores = [v.strip() for v in resto.split(",")]
        if len(valores) > len(campos):
            return {"error": "El comando tiene más campos de los esperados."}
    else:
        return {
            "error": (
                "Usá el formato /guardar asegurado (asegurado) (numero) "
                "(vehiculo) (patente) (cia) (medio de pago) (cp) (mail) "
                "(telefono opcional)."
            )
        }

    propuesta = {
        campo: (valores[i] if i < len(valores) else "")
        for i, campo in enumerate(campos)
    }

    placeholders = {
        "ASEGURADO": "asegurado",
        "NUMERO": "numero",
        "VEHICULO": "vehiculo",
        "PATENTE": "patente",
        "CIA": "cia",
        "MEDIO DE PAGO": "medio de pago",
        "CP": "cp",
        "MAIL": "mail",
        "TELEFONO": "telefono",
    }
    for campo, placeholder in placeholders.items():
        if propuesta[campo].strip().lower() == placeholder:
            propuesta[campo] = ""

    propuesta["CIA"] = normalizar_compania(propuesta.get("CIA", ""))
    propuesta["ENVIOS YA"] = ""

    return {
        "propuesta": propuesta,
        "libro_id": libro_id,
        "valida": bool(propuesta["ASEGURADO"] and (
            propuesta["NUMERO"] or propuesta["PATENTE"]
        )),
    }



def _consulta_requiere_metadatos(pregunta):
    """Detecta consultas donde conviene precargar metadatos internos antes de Gemini."""
    texto = unicodedata.normalize("NFKD", str(pregunta or "")).encode("ascii", "ignore").decode("ascii").lower()
    texto = re.sub(r"\s+", " ", texto).strip()
    if not texto or texto.startswith("/"):
        return False

    # Excluir operaciones que ya tienen routing determinístico/estructurado.
    terminos_estructurados = (
        "asegurado", "asegurados", "patente", "patentes", "poliza", "póliza",
        "polizas", "pólizas", "planilla", "excel", "dni", "numero de poliza",
        "número de póliza", "cuantos registros", "cuántos registros",
        "cantidad de vehiculos", "cantidad de vehículos",
    )
    if any(t in texto for t in terminos_estructurados):
        return False

    # Una frase como "cuántos remolques tiene ATM" puede hablar del inventario
    # de pólizas/vehículos, no de la cantidad de servicios de grúa de una
    # cobertura. Los conteos de entidades registradas deben ir al Excel.
    es_conteo = bool(re.search(r"\b(cuantos|cuantas|cantidad|total)\b", texto))
    entidades_excel = (
        "asegurado", "asegurados", "poliza", "polizas", "vehiculo", "vehiculos",
        "remolque", "remolques", "trailer", "trailers", "acoplado", "acoplados",
        "moto", "motos", "auto", "autos", "camion", "camiones",
    )
    if es_conteo and any(t in texto for t in entidades_excel):
        return False

    terminos_documentales = (
        "cobertura", "cubre", "cubrir", "asistencia", "asistencias",
        "grua", "grúa", "remolque", "remolques", "auxilio", "traslado",
        "limite", "límite", "limites", "límites", "condicion", "condición",
        "condiciones", "procedimiento", "procedimientos", "compañia",
        "compañía", "compañias", "compañías", "prestacion", "prestación",
        "prestaciones", "servicio", "servicios", "evento", "kilometros",
        "kilómetros", "cerradura", "cerraduras", "granizo", "vidrio",
        "vidrios", "rueda", "ruedas", "robo", "incendio", "destruccion",
        "destrucción", "responsabilidad civil", "rc",
    )
    return any(t in texto for t in terminos_documentales)


def _formatear_metadatos_para_contexto(resultado):
    """Convierte el resultado existente de buscar_en_metadatos en contexto legible para Gemini."""
    if not isinstance(resultado, dict) or not resultado.get("cantidad"):
        return ""
    fichas = resultado.get("fichas") or []
    if not fichas:
        return ""

    partes = [
        "\n\n===== METADATOS INTERNOS PRIORITARIOS =====",
        "Estos datos fueron recuperados localmente por OficinaIA antes de llamar a Gemini.",
        "Usalos como fuente prioritaria para información operativa interna.",
        "",
    ]
    for ficha in fichas:
        titulo = str(ficha.get("titulo") or "Metadato").strip()
        contenido = str(ficha.get("contenido") or "").strip()
        if not contenido:
            continue
        partes.append(f"[Metadato: {titulo}]")
        partes.append(contenido)
        partes.append("")
    partes.append("===== FIN METADATOS INTERNOS PRIORITARIOS =====\n")
    return "\n".join(partes)


def _mensaje_error_chat(error):
    """Clasifica errores escapados del chat sin modificar su causa ni el logging."""
    texto = f"{type(error).__module__} {type(error).__name__} {error}".lower()

    indicadores_ia = (
        "gemini", "google.genai", "genai", "resource_exhausted",
        "quota", "api key", "apikey", "model", "generate_content",
    )
    if any(indicador in texto for indicador in indicadores_ia):
        return (
            "No pude procesar la consulta con el servicio de IA en este momento. "
            "Intentá nuevamente en unos segundos."
        )

    indicadores_metadatos = (
        "metadato", "metadata", "buscar_en_metadatos", "_cargar_metadatos",
    )
    if any(indicador in texto for indicador in indicadores_metadatos):
        return (
            "No pude completar la búsqueda de información en este momento. "
            "Intentá nuevamente."
        )

    return (
        "Ocurrió un problema al procesar la consulta. Intentá nuevamente. "
        "Si el problema continúa, avisá al administrador."
    )


def _envolver_chat_con_manejo_de_errores(func):
    """Evita que una excepción no clasificada del chat termine en un 500 genérico."""
    @wraps(func)
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as error:
            print("ERROR CHAT NO CONTROLADO:", error)
            return jsonify({"ok": False, "error": _mensaje_error_chat(error)}), 500
    return wrapper


@app.route(
    "/api/chat",
    methods=["POST"]
)
@requiere_login
@_envolver_chat_con_manejo_de_errores
def chat():

    # El chat acepta JSON para consultas normales y multipart/form-data
    # cuando el usuario adjunta un PDF. El PDF se procesa en memoria y no
    # se guarda como documento permanente.
    if request.is_json:
        data = request.get_json(silent=True) or {}
        mensaje = str(data.get("mensaje", "")).strip()
        chat_id = data.get("chat_id")
        historial = data.get("historial") or []
        archivo_pdf = None
    else:
        data = request.form
        mensaje = str(data.get("mensaje", "")).strip()
        chat_id = data.get("chat_id")
        historial_raw = data.get("historial", "[]")
        try:
            import json
            historial = json.loads(historial_raw)
        except Exception:
            historial = []
        archivo_pdf = request.files.get("pdf")

    if not data and not archivo_pdf:
        return jsonify({"respuesta": "No recibí ningún mensaje."})

    # El libro de un /guardar asegurado queda pendiente sólo para su confirmación
    # inmediata; cualquier nuevo mensaje de chat invalida ese destino para no
    # arrastrarlo accidentalmente a otra propuesta.
    session.pop("guardar_asegurado_libro_id", None)

    try:
        chat_id = int(chat_id) if chat_id else None
    except (TypeError, ValueError):
        chat_id = None

    if chat_id and not _validar_chat(chat_id, session["usuario"]):
        chat_id = None
    if not chat_id:
        titulo = _generar_titulo_chat(mensaje)
        chat_id = _crear_conversacion(session["usuario"], titulo)
    else:
        # Chat precargado como "Nueva conversación": titular con la primera consulta real.
        _auto_titulo_si_corresponde(chat_id, session["usuario"], mensaje)

    # P1.3 — Historial desde DB (fuente de verdad). El JSON del cliente es
    # legacy/opcional y solo se usa si la DB no tiene mensajes aún.
    historial_db = _historial_desde_db(chat_id, session["usuario"], limite=10)
    if historial_db:
        historial = historial_db
    else:
        if not isinstance(historial, list):
            historial = []
        historial = [
            x for x in historial
            if isinstance(x, dict)
            and x.get("rol") in {"user", "assistant"}
            and str(x.get("contenido", "")).strip()
        ][-10:]

    # P2.6 — tipo de chat persistido (flota/coti/alta/envios)
    _asignar_tipo_chat(chat_id, session["usuario"], mensaje)

    contexto_pdf_adjunto = ""
    nombre_pdf_adjunto = ""
    if archivo_pdf and archivo_pdf.filename:
        nombre_pdf_adjunto = secure_filename(archivo_pdf.filename) or "documento.pdf"
        if not nombre_pdf_adjunto.lower().endswith(".pdf"):
            return jsonify({"ok": False, "error": "El archivo adjunto debe ser un PDF."}), 400

        try:
            archivo_pdf.stream.seek(0, os.SEEK_END)
            tamaño = archivo_pdf.stream.tell()
            archivo_pdf.stream.seek(0)
            if tamaño > 20 * 1024 * 1024:
                return jsonify({"ok": False, "error": "El PDF es demasiado grande. El máximo permitido es 20 MB."}), 413

            datos_pdf = archivo_pdf.stream.read()
            archivo_pdf.stream.seek(0)

            if len(datos_pdf) > MAX_PDF_FILE_SIZE_BYTES:
                return jsonify({
                    "ok": False,
                    "error": "El PDF es demasiado grande para procesarlo en el chat."
                }), 413

            paginas = []
            total_chars = 0
            try:
                documento = fitz.open(stream=datos_pdf, filetype="pdf")
                try:
                    max_paginas = min(documento.page_count, MAX_PDF_PAGES_CHAT)
                    for numero in range(max_paginas):
                        if total_chars >= MAX_PDF_TEXT_CHARS_CHAT:
                            break
                        try:
                            pagina = documento.load_page(numero)
                            texto = pagina.get_text("text", sort=True) or ""
                            del pagina
                        except Exception as exc:
                            print(f"ERROR EXTRAYENDO PDF ADJUNTO PÁGINA {numero + 1}: {exc}")
                            continue

                        texto = re.sub(r"[ \t]+", " ", texto).strip()
                        if not texto:
                            continue

                        restante = MAX_PDF_TEXT_CHARS_CHAT - total_chars
                        texto = texto[:restante]
                        paginas.append(f"PÁGINA {numero + 1}\n{texto}")
                        total_chars += len(texto)
                finally:
                    documento.close()
            except Exception as exc:
                return jsonify({
                    "ok": False,
                    "error": f"No se pudo leer el PDF adjunto: {exc}"
                }), 400
            finally:
                del datos_pdf

            if not paginas:
                return jsonify({
                    "ok": False,
                    "error": "El PDF parece ser escaneado o no contiene texto seleccionable. En esta versión puedo leer PDFs con texto."
                }), 422

            contexto_pdf_adjunto = (
                "\n\n===== PDF ADJUNTADO EN EL CHAT =====\n"
                f"ARCHIVO: {nombre_pdf_adjunto}\n"
                f"PÁGINAS PROCESADAS: {max_paginas}\n\n"
                + "\n\n".join(paginas)
                + "\n===== FIN PDF ADJUNTADO =====\n"
            )
        except Exception as error:
            print("ERROR PDF ADJUNTO:", error)
            return jsonify({"ok": False, "error": "No pude leer ese PDF. Verificá que el archivo no esté dañado."}), 422

    if not mensaje and archivo_pdf:
        mensaje = _MENSAJE_PDF_POR_DEFECTO

    if not mensaje:

        return jsonify({
            "respuesta":
                "Escribime una consulta."
        })

    mensaje_guardado = mensaje
    if nombre_pdf_adjunto:
        mensaje_guardado = f"[PDF adjunto: {nombre_pdf_adjunto}]\n{mensaje}"
    _guardar_mensaje(chat_id, "user", mensaje_guardado)

    # ======================================================
    # COMANDO /COTI — RESOLUCIÓN LOCAL Y DETERMINÍSTICA
    # ======================================================
    # Se intercepta antes de cualquier llamada a Gemini. El catálogo y el
    # parser viven en coti.py para que puedan ampliarse sin tocar /api/chat.
    respuesta_coti = procesar_comando_coti(mensaje)
    if respuesta_coti is not None:
        _guardar_mensaje(chat_id, "assistant", str(respuesta_coti))

        return jsonify({
            "respuesta": respuesta_coti,
            "chat_id": chat_id,
            "archivo_adjunto": nombre_pdf_adjunto or None,
            "propuesta_excel": None,
            "propuesta_metadato": None,
        })

    # ======================================================
    # COMANDO /FLOTA — TAREA CONVERSACIONAL PERSISTENTE
    # ======================================================
    # A diferencia de /coti y /guardar asegurado, /flota no se resuelve en
    # un solo turno: el contexto (datos generales + vehículos, guardados o
    # pendientes) vive en `flotas_activas` atado a chat_id, así que TODOS
    # los mensajes de esta conversación pasan por acá primero — no sólo los
    # que empiezan con "/flota" — para poder reconocer continuaciones
    # ("vehículos 11-20"), correcciones ("el 7 es C3") y el cierre de la
    # tarea sin que el usuario tenga que repetir el comando cada vez.
    respuesta_flota, atendido_por_flota, tabulado_flota = _flota_procesar_turno(
        chat_id, mensaje, contexto_pdf_adjunto
    )
    if atendido_por_flota:
        _guardar_mensaje(chat_id, "assistant", str(respuesta_flota))
        return jsonify({
            "respuesta": respuesta_flota,
            "chat_id": chat_id,
            "archivo_adjunto": nombre_pdf_adjunto or None,
            "propuesta_excel": None,
            "propuesta_metadato": None,
            "tabulado_flota": tabulado_flota,
        })

    # ======================================================
    # COMANDO /ALTA — PÓLIZA INDIVIDUAL → PROPUESTA DE ASEGURADO
    # ======================================================
    # Tanda 4: flujo nuevo y separado de /flota. Un solo turno, sin estado
    # persistente por chat_id. La propuesta queda en sesión para que, más
    # adelante (Tanda 5), el usuario pueda elegir "tabular" o "guardar en
    # Excel" sin tener que repetir la extracción.
    # Tanda 8 — detección automática: si el usuario tira el PDF de una
    # póliza individual al chat (con el clip o arrastrándolo) SIN escribir
    # "/alta", Sofia reconoce igual que es una póliza de un solo asegurado
    # y arma la propuesta de alta directamente, sin obligarlo a escribir
    # el comando. Sólo se activa cuando nada más ya se hizo cargo del
    # mensaje (ni /coti, ni /flota) y el contenido del PDF no parece una
    # flota (ver _pdf_parece_poliza_individual_para_alta).
    es_alta_explicito = bool(re.match(r"^/alta\b", mensaje, re.IGNORECASE))
    mensaje_para_alta = mensaje
    alta_automatica = False
    if (
        not es_alta_explicito
        and contexto_pdf_adjunto
        and _pdf_parece_poliza_individual_para_alta(contexto_pdf_adjunto)
    ):
        texto_extra = "" if mensaje.strip() == _MENSAJE_PDF_POR_DEFECTO else mensaje.strip()
        mensaje_para_alta = ("/alta " + texto_extra).strip()
        alta_automatica = True

    respuesta_alta, atendido_por_alta, propuesta_alta = _procesar_alta_asegurado(
        mensaje_para_alta, contexto_pdf_adjunto, automatico=alta_automatica
    )
    if atendido_por_alta:
        _guardar_mensaje(chat_id, "assistant", str(respuesta_alta))
        tabulado_alta = _armar_tabulado_alta(propuesta_alta) if propuesta_alta else None
        campos_guardar_alta = (
            _alta_a_campos_guardar_asegurado(propuesta_alta) if propuesta_alta else None
        )
        if propuesta_alta:
            session["alta_asegurado_propuesta"] = propuesta_alta
        return jsonify({
            "respuesta": respuesta_alta,
            "chat_id": chat_id,
            "archivo_adjunto": nombre_pdf_adjunto or None,
            "propuesta_excel": None,
            "propuesta_metadato": None,
            "propuesta_alta_asegurado": propuesta_alta,
            "tabulado_alta_asegurado": tabulado_alta,
            "campos_guardar_alta_asegurado": campos_guardar_alta,
        })

    # ======================================================
    # COMANDO /ENVIOS YA — BUSCAR ASEGURADO POR PATENTE
    # ======================================================
    propuesta_envios_ya = _parsear_comando_envios_ya(mensaje)
    if propuesta_envios_ya is not None:
        if propuesta_envios_ya.get("error"):
            respuesta = propuesta_envios_ya["error"]
            texto_envios_ya = None
        else:
            fila_asegurado = _buscar_asegurado_por_patente(propuesta_envios_ya["patente"])
            if fila_asegurado is None:
                respuesta = (
                    f"No encontré ningún asegurado con la patente "
                    f"{propuesta_envios_ya['patente'].upper()} en el Excel. Revisá que esté "
                    "bien escrita o que el asegurado ya esté guardado."
                )
                texto_envios_ya = None
            else:
                texto_envios_ya = _armar_texto_envios_ya(fila_asegurado)
                respuesta = "Te dejo los datos listos para pegar en Envíos Ya:"

        _guardar_mensaje(chat_id, "assistant", str(respuesta))
        return jsonify({
            "respuesta": respuesta,
            "chat_id": chat_id,
            "archivo_adjunto": nombre_pdf_adjunto or None,
            "propuesta_excel": None,
            "propuesta_metadato": None,
            "texto_envios_ya": texto_envios_ya,
        })

    # El comando explícito se parsea de forma determinista en backend y no
    # se entrega a Gemini para que adivine posiciones o separadores.
    propuesta_comando = _parsear_comando_guardar_asegurado(mensaje)
    if propuesta_comando is not None:
        if propuesta_comando.get("error"):
            respuesta = propuesta_comando["error"]
            propuesta_excel = None
        else:
            libro_id = str(propuesta_comando.get("libro_id") or "1")
            libro = LIBROS_EXCEL[libro_id]
            session["guardar_asegurado_libro_id"] = libro_id
            respuesta = (
                f"Voy a guardar este asegurado en Excel {libro_id} ({libro['nombre']}):\n\n"
                f"ASEGURADO: {propuesta_comando['propuesta'].get('ASEGURADO', '')}\n"
                f"NUMERO: {propuesta_comando['propuesta'].get('NUMERO', '')}\n"
                f"VEHICULO: {propuesta_comando['propuesta'].get('VEHICULO', '')}\n"
                f"PATENTE: {propuesta_comando['propuesta'].get('PATENTE', '')}\n"
                f"CIA: {propuesta_comando['propuesta'].get('CIA', '')}\n"
                f"MEDIO DE PAGO: {propuesta_comando['propuesta'].get('MEDIO DE PAGO', '')}\n"
                f"CP: {propuesta_comando['propuesta'].get('CP', '')}\n"
                f"MAIL: {propuesta_comando['propuesta'].get('MAIL', '')}\n\n"
                "¿Confirmás?"
            )
            propuesta_excel = propuesta_comando.get("propuesta")
            propuesta_excel["LIBRO_ID"] = libro_id

        _guardar_mensaje(chat_id, "assistant", str(respuesta))

        return jsonify({
            "respuesta": respuesta,
            "chat_id": chat_id,
            "archivo_adjunto": nombre_pdf_adjunto or None,
            "propuesta_excel": propuesta_excel,
            "propuesta_metadato": None,
        })

    # ======================================================
    # CONTEXTO DIRECTO + PRE-ROUTING DE METADATOS
    # ======================================================
    # Los metadatos internos se recuperan localmente antes de Gemini
    # para que el modelo no tenga que decidir si necesita consultarlos.
    # La tool buscar_en_metadatos sigue disponible para búsquedas adicionales.
    contexto = contexto_pdf_adjunto
    error_pre_routing_metadatos = False

    if _consulta_requiere_metadatos(mensaje):
        try:
            resultado_metadatos = buscar_en_metadatos(mensaje)
            if isinstance(resultado_metadatos, dict) and resultado_metadatos.get("error"):
                error_pre_routing_metadatos = True
                print(
                    "ERROR PRE-ROUTING METADATOS: %s consulta=%r"
                    % (resultado_metadatos.get("error"), mensaje)
                )
            else:
                contexto_metadatos = _formatear_metadatos_para_contexto(resultado_metadatos)
                if contexto_metadatos:
                    contexto += contexto_metadatos
                    print(
                        "PRE-ROUTING METADATOS: resultados=%s consulta=%r"
                        % (resultado_metadatos.get("cantidad", 0), mensaje)
                    )
                else:
                    print("PRE-ROUTING METADATOS: sin resultados consulta=%r" % mensaje)
        except Exception as error:
            error_pre_routing_metadatos = True
            # No se elimina el detalle técnico del log; el usuario recibe un mensaje controlado.
            print("ERROR PRE-ROUTING METADATOS:", error)

    # ======================================================
    # GEMINI
    # ======================================================

    propuesta_excel = None
    propuesta_metadato = None
    try:

        from servicios_ia import (
            consultar_gemini
        )

        resultado_gemini = consultar_gemini(
            mensaje,
            contexto,
            historial=historial
        )
        if isinstance(resultado_gemini, tuple):
            respuesta = resultado_gemini[0]
            propuesta_excel = resultado_gemini[1] if len(resultado_gemini) > 1 else None
            propuesta_metadato = resultado_gemini[2] if len(resultado_gemini) > 2 else None
        else:
            respuesta, propuesta_excel = resultado_gemini, None

    except Exception as error:

        print(
            "ERROR CHAT GEMINI:",
            error
        )

        if error_pre_routing_metadatos:
            respuesta = (
                "No pude completar la búsqueda de información en este momento. "
                "Intentá nuevamente."
            )
        elif contexto:

            respuesta = (
                "Encontré información "
                "relacionada en la oficina.\n\n"
                + contexto[:8000]
            )

        else:

            respuesta = (
                "No encontré información suficiente sobre ese punto en las fuentes "
                "disponibles actualmente. Puede que todavía no tengamos esa información "
                "cargada en los metadatos. Si querés, podemos revisar la documentación "
                "correspondiente o cargar esa información para que OfIA pueda utilizarla "
                "en futuras consultas."
            )

    _guardar_mensaje(chat_id, "assistant", str(respuesta))

    return jsonify({
        "respuesta": respuesta,
        "chat_id": chat_id,
        "archivo_adjunto": nombre_pdf_adjunto or None,
        "propuesta_excel": propuesta_excel,
        "propuesta_metadato": propuesta_metadato,
    })


# ==========================================================
# CONFIGURACIÓN
# ==========================================================

@app.route("/manuales")
@requiere_login
def manuales():
    # Mantener /manuales como ruta principal para compatibilidad con favoritos y enlaces antiguos.
    return redirect(url_for("biblioteca"))


@app.route("/biblioteca")
@requiere_login
def biblioteca():
    try:
        filas = pg_listar_polizas()
        polizas = [
            {
                "archivo": f["r2_key"],
                "nombre": f["nombre"],
                "fecha": f["fecha_subida"].strftime("%d/%m/%Y %H:%M") if f.get("fecha_subida") else "",
                "tamaño": round((f.get("tamaño") or 0) / 1024, 1),
            }
            for f in filas
        ]
    except Exception as error:
        print("ERROR LISTANDO POLIZAS:", error)
        polizas = []
    return render_template("biblioteca.html", manuales=manuales_companias(), polizas=polizas,
                           usuario=session["usuario"], usuario_rol=session.get("rol","usuario"),
                           usuario_es_admin=usuario_es_admin())

@app.route("/api/polizas", methods=["POST"])
@requiere_admin
def subir_poliza():
    archivo=request.files.get("poliza")
    if not archivo or not archivo.filename:
        return jsonify(ok=False,error="Seleccioná un archivo PDF."),400
    nombre=secure_filename(Path(archivo.filename).name)
    if not nombre.lower().endswith(".pdf"):
        return jsonify(ok=False,error="El archivo debe ser un PDF."),400

    try:
        archivo.stream.seek(0, os.SEEK_END)
        tamaño = archivo.stream.tell()
        archivo.stream.seek(0)
    except Exception:
        return jsonify(ok=False,error="No se pudo leer el archivo."),400

    if tamaño <= 0:
        return jsonify(ok=False,error="El PDF está vacío."),400
    if tamaño > MAX_PDF_FILE_SIZE_BYTES:
        return jsonify(ok=False,error=f"El PDF es demasiado grande. El máximo permitido es {MAX_PDF_FILE_SIZE_BYTES // (1024*1024)} MB."),413

    try:
        cabecera = archivo.stream.read(5)
        archivo.stream.seek(0)
        if cabecera != b"%PDF-":
            return jsonify(ok=False,error="El archivo no parece ser un PDF válido."),400
        datos_validacion = archivo.stream.read()
        archivo.stream.seek(0)
        documento = fitz.open(stream=datos_validacion, filetype="pdf")
        documento.close()
        del datos_validacion
    except Exception:
        return jsonify(ok=False,error="No se pudo leer el PDF. Verificá que no esté dañado."),400

    import uuid
    r2_key = f"polizas/{uuid.uuid4().hex}__{nombre}"

    try:
        r2_subir_pdf(archivo.stream, r2_key, tamaño)
    except Exception as error:
        print("ERROR SUBIENDO POLIZA A R2:", error)
        return jsonify(ok=False,error="No se pudo guardar la póliza en Cloudflare R2."),502

    try:
        registrar_poliza(nombre, r2_key, tamaño)
    except Exception as error:
        print("ERROR REGISTRANDO POLIZA EN NEON:", error)
        try:
            r2_eliminar_pdf(r2_key)
        except Exception as rollback_error:
            print("ERROR ROLLBACK R2 POLIZA:", rollback_error)
        return jsonify(ok=False,error="La póliza se subió a R2 pero no pudo registrarse en PostgreSQL. La operación no se completó."),502

    return jsonify(ok=True, archivo=r2_key, nombre=nombre, tamaño=tamaño)

@app.route("/api/polizas/<path:nombre>", methods=["DELETE"])
@requiere_admin
def eliminar_poliza(nombre):
    r2_key = str(nombre or "").strip()
    if not r2_key.startswith("polizas/") or not r2_key.lower().endswith(".pdf"):
        return jsonify(ok=False,error="Póliza no encontrada."),404

    existente = obtener_poliza_por_r2_key(r2_key)
    if not existente:
        return jsonify(ok=False,error="Póliza no encontrada."),404

    try:
        eliminado = eliminar_poliza_pg(r2_key)
        if not eliminado:
            return jsonify(ok=False,error="Póliza no encontrada."),404
    except Exception as error:
        print("ERROR ELIMINANDO POLIZA DE NEON:", error)
        return jsonify(ok=False,error="No se pudo actualizar PostgreSQL. La póliza no fue eliminada."),502

    try:
        r2_eliminar_pdf(r2_key)
    except Exception as error:
        print("ERROR ELIMINANDO POLIZA DE R2:", error)
        try:
            registrar_poliza(existente["nombre"], existente["r2_key"], existente["tamaño"])
        except Exception as rollback_error:
            print("ERROR RESTAURANDO POLIZA EN NEON:", rollback_error)
        return jsonify(ok=False,error="No se pudo eliminar el PDF de Cloudflare R2. La póliza se mantuvo registrada."),502

    return jsonify(ok=True)

@app.route("/polizas/<path:nombre>")
@requiere_login
def ver_poliza(nombre):
    r2_key = str(nombre or "").strip()
    if not r2_key.startswith("polizas/") or not r2_key.lower().endswith(".pdf"):
        return ("Póliza no encontrada",404)

    existente = obtener_poliza_por_r2_key(r2_key)
    if not existente:
        return ("Póliza no encontrada",404)

    try:
        objeto = obtener_objeto_stream(r2_key)
        body = objeto["Body"]
        content_length = objeto.get("ContentLength")

        @stream_with_context
        def generar():
            try:
                while True:
                    bloque = body.read(1024 * 1024)
                    if not bloque:
                        break
                    yield bloque
            finally:
                body.close()

        headers = {
            "Content-Type": "application/pdf",
            "Content-Disposition": f'inline; filename="{existente["nombre"]}"',
            "Cache-Control": "private, max-age=300",
        }
        if content_length is not None:
            headers["Content-Length"] = str(content_length)

        return Response(generar(), headers=headers)
    except Exception as error:
        print("ERROR SIRVIENDO POLIZA R2:", error)
        return ("No se pudo abrir la póliza.", 502)

@app.route("/configuracion")
@requiere_login
def configuracion():
    config = cargar_configuracion()
    usuarios=[]
    if usuario_es_admin():
        if _usuarios_usar_pg():
            try:
                usuarios = pg_listar_usuarios()
            except Exception as error:
                print("ERROR listar usuarios PG:", error)
                usuarios = []
        else:
            with closing(conectar_db()) as db:
                usuarios=db.execute("SELECT id,usuario,email,rol,protegido FROM usuarios ORDER BY usuario COLLATE NOCASE").fetchall()
    return render_template("configuracion.html",config=config,usuario=session["usuario"],carpetas=obtener_companias(),usuarios=usuarios)

@app.route("/api/configuracion", methods=["POST"])
@requiere_admin
def guardar_configuracion():
    data=request.get_json(silent=True) or {}
    config=cargar_configuracion()
    nombre=str(data.get("nombre_oficina",config["nombre_oficina"])).strip()
    if not nombre:
        return jsonify(ok=False,error="El nombre de la oficina no puede estar vacío."),400

    colores = {
        "color_principal": data.get("color_principal", config["color_principal"]),
        "color_acento": data.get("color_acento", config["color_acento"]),
        "color_fondo": data.get("color_fondo", config["color_fondo"]),
        "color_sidebar": data.get("color_sidebar", config["color_sidebar"]),
        "color_botones": data.get("color_botones", config["color_botones"]),
    }
    for clave, valor in colores.items():
        valor = str(valor).strip()
        if not re.fullmatch(r"#[0-9a-fA-F]{6}", valor):
            return jsonify(ok=False,error=f"El color {clave} no es válido."),400
        colores[clave] = valor.upper()

    herramientas_recibidas = data.get("herramientas", config.get("herramientas", []))
    if not isinstance(herramientas_recibidas, list):
        return jsonify(ok=False,error="La configuración de herramientas no es válida."),400

    herramientas = []
    ids = set()
    for i, item in enumerate(herramientas_recibidas):
        if not isinstance(item, dict):
            return jsonify(ok=False,error="Hay una herramienta inválida."),400
        nombre_h = str(item.get("nombre", "") or "").strip()
        url_h = str(item.get("url", "") or "").strip()
        if not nombre_h or not url_h:
            return jsonify(ok=False,error="Cada herramienta necesita nombre y URL."),400
        if len(nombre_h) > 60 or len(url_h) > 1000:
            return jsonify(ok=False,error="Nombre o URL de herramienta demasiado largo."),400
        if not re.match(r"^https?://", url_h, re.IGNORECASE):
            return jsonify(ok=False,error=f"La URL de {nombre_h} debe comenzar con http:// o https://"),400
        ident = re.sub(r"[^a-z0-9_-]+", "-", str(item.get("id", "") or "").lower()).strip("-")
        if not ident:
            ident = f"herramienta-{i+1}"
        base = ident; n = 2
        while ident in ids:
            ident = f"{base}-{n}"; n += 1
        ids.add(ident)
        herramientas.append({
            "id": ident,
            "nombre": nombre_h,
            "url": url_h,
            "visible": bool(item.get("visible", True)),
        })

    config["nombre_oficina"]=nombre
    config["notificaciones"]=bool(data.get("notificaciones",config["notificaciones"]))
    config["herramientas"]=herramientas
    config.pop("herramientas_visibles", None)
    config.pop("herramientas_urls", None)
    config["excel_visible"]=bool(data.get("excel_visible", config.get("excel_visible", True)))
    config.update(colores)
    try:
        if _config_usar_pg():
            pg_guardar_configuracion(config)
        else:
            CONFIG_FILE.write_text(json.dumps(config,ensure_ascii=False,indent=2),encoding="utf-8")
        return jsonify(ok=True, config=config)
    except Exception as error:
        print("ERROR guardar_configuracion:", error)
        return jsonify(ok=False,error="No se pudo guardar la configuración."),500


@app.route("/api/usuarios", methods=["POST"])
@requiere_admin
def crear_usuario():
    data=request.get_json(silent=True) or {}
    usuario=str(data.get("usuario","")).strip(); password=str(data.get("password","")); email=str(data.get("email","")).strip(); rol=str(data.get("rol","usuario")).strip().lower()
    if not usuario: return jsonify(ok=False,error="El usuario es obligatorio."),400
    if not password: return jsonify(ok=False,error="La contraseña es obligatoria."),400
    if rol not in ROLES_VALIDOS: return jsonify(ok=False,error="Rol inválido."),400
    if not validar_email(email): return jsonify(ok=False,error="El correo electrónico no es válido."),400
    if _usuarios_usar_pg():
        try:
            if pg_usuario_existe(usuario):
                return jsonify(ok=False,error="Ese usuario ya existe."),409
            fila_nueva = pg_crear_usuario(usuario, generate_password_hash(password), email, rol)
            nuevo_id = fila_nueva["id"] if isinstance(fila_nueva, dict) else fila_nueva
        except Exception as error:
            print("ERROR crear_usuario PG:", error)
            return jsonify(ok=False,error="No se pudo crear el usuario."),500
        return jsonify(ok=True, mensaje="Usuario creado correctamente.", id=nuevo_id)
    with closing(conectar_db()) as db:
        if db.execute("SELECT 1 FROM usuarios WHERE lower(usuario)=lower(?)",(usuario,)).fetchone(): return jsonify(ok=False,error="Ese usuario ya existe."),409
        cur = db.execute("INSERT INTO usuarios (usuario,password,email,rol,protegido) VALUES (?,?,?,?,0)",(usuario,generate_password_hash(password),email,rol)); db.commit()
        nuevo_id = cur.lastrowid
    return jsonify(ok=True, mensaje="Usuario creado correctamente.", id=nuevo_id)

@app.route("/api/usuarios/<int:usuario_id>", methods=["PUT"])
@requiere_admin
def editar_usuario(usuario_id):
    registro=obtener_usuario_por_id(usuario_id)
    if not registro: return jsonify(ok=False,error="Usuario no encontrado."),404
    data=request.get_json(silent=True) or {}; email=str(data.get("email","")).strip(); rol=str(data.get("rol",registro["rol"])).strip().lower(); password=str(data.get("password",""))
    if not validar_email(email): return jsonify(ok=False,error="El correo electrónico no es válido."),400
    if rol not in ROLES_VALIDOS: return jsonify(ok=False,error="Rol inválido."),400
    if registro["protegido"]: rol="admin"
    if _usuarios_usar_pg():
        try:
            pg_actualizar_usuario(usuario_id, email, rol, generate_password_hash(password) if password else None)
        except Exception as error:
            print("ERROR editar_usuario PG:", error)
            return jsonify(ok=False,error="No se pudo actualizar el usuario."),500
        return jsonify(ok=True,mensaje="Usuario actualizado correctamente.")
    with closing(conectar_db()) as db:
        db.execute("UPDATE usuarios SET email=?,rol=? WHERE id=?",(email,rol,usuario_id))
        if password: db.execute("UPDATE usuarios SET password=? WHERE id=?",(generate_password_hash(password),usuario_id))
        db.commit()
    return jsonify(ok=True,mensaje="Usuario actualizado correctamente.")

@app.route("/api/usuarios/<int:usuario_id>", methods=["DELETE"])
@requiere_admin
def eliminar_usuario(usuario_id):
    registro=obtener_usuario_por_id(usuario_id)
    if not registro: return jsonify(ok=False,error="Usuario no encontrado."),404
    if registro["protegido"]: return jsonify(ok=False,error="El administrador principal está protegido."),403
    if registro["usuario"]==session.get("usuario"): return jsonify(ok=False,error="No podés eliminar tu propia cuenta."),400
    if _usuarios_usar_pg():
        try:
            pg_eliminar_usuario(usuario_id)
        except Exception as error:
            print("ERROR eliminar_usuario PG:", error)
            return jsonify(ok=False,error="No se pudo eliminar el usuario."),500
        return jsonify(ok=True,mensaje="Usuario eliminado correctamente.")
    with closing(conectar_db()) as db: db.execute("DELETE FROM usuarios WHERE id=?",(usuario_id,)); db.commit()
    return jsonify(ok=True,mensaje="Usuario eliminado correctamente.")


@app.route("/api/manuales/<slug>", methods=["POST"])
@requiere_admin
def subir_manual(slug):
    """
    Agrega manuales de forma acumulativa.

    - Acepta uno o varios archivos en el campo "manual".
    - Cada PDF obtiene una clave R2 única, por lo que nunca reemplaza otro
      manual salvo que el frontend envíe explícitamente "replace".
    - Un fallo de un archivo no elimina los manuales que ya estaban guardados
      ni los que se hayan completado antes en esta misma solicitud.
    """
    compania = next(
        (c for c in MANUALES_COMPANIAS if slug_manual_compania(c) == slug),
        None,
    )
    if not compania:
        return jsonify(ok=False, error="Compañía no válida."), 404

    archivos = [a for a in request.files.getlist("manual") if a and a.filename]
    if not archivos:
        # Compatibilidad con clientes antiguos que enviaban request.files.get().
        archivo_unico = request.files.get("manual")
        if archivo_unico and archivo_unico.filename:
            archivos = [archivo_unico]

    if not archivos:
        return jsonify(ok=False, error="Seleccioná al menos un archivo PDF."), 400

    reemplazar = str(request.form.get("replace", "")).strip()
    if reemplazar and len(archivos) != 1:
        return jsonify(
            ok=False,
            error="El reemplazo de un manual debe hacerse con un solo PDF."
        ), 400

    # El límite de Flask protege cada petición. Además validamos cada archivo
    # individualmente para devolver un error claro y evitar cargas parciales
    # por formato/tamaño antes de escribir en R2.
    archivos_preparados = []
    for archivo in archivos:
        nombre_seguro = secure_filename(Path(archivo.filename).name)
        if not nombre_seguro or Path(nombre_seguro).suffix.lower() != ".pdf":
            return jsonify(
                ok=False,
                error=f'El archivo "{archivo.filename}" no es un PDF válido.'
            ), 400

        try:
            archivo.stream.seek(0, os.SEEK_END)
            tamaño = archivo.stream.tell()
            archivo.stream.seek(0)
        except Exception:
            return jsonify(
                ok=False,
                error=f'No se pudo leer el archivo "{archivo.filename}".'
            ), 400

        if tamaño <= 0:
            return jsonify(
                ok=False,
                error=f'El PDF "{archivo.filename}" está vacío.'
            ), 400

        if tamaño > 20 * 1024 * 1024:
            return jsonify(
                ok=False,
                error=(
                    f'El PDF "{archivo.filename}" es demasiado grande. '
                    "El máximo permitido es 20 MB por archivo."
                )
            ), 413

        archivos_preparados.append((archivo, nombre_seguro, tamaño))

    r2_key_anterior = None
    existente = None

    if reemplazar:
        prefijo = f"manuales/{slug}/"
        if not reemplazar.startswith(prefijo) or not reemplazar.lower().endswith(".pdf"):
            return jsonify(ok=False, error="El manual a reemplazar no es válido."), 400

        existente = obtener_manual_por_r2_key(reemplazar)
        if not existente:
            return jsonify(ok=False, error="El manual a reemplazar no existe."), 404
        r2_key_anterior = reemplazar

    import uuid

    resultados = []
    subidos_r2 = []

    for archivo, nombre_seguro, tamaño in archivos_preparados:
        r2_key_nuevo = (
            f"manuales/{slug}/"
            f"{uuid.uuid4().hex}__{nombre_seguro}"
        )

        # Validamos antes de persistir. Los PDFs siguen siendo privados en R2;
        # no se crea una copia permanente local.
        try:
            archivo.stream.seek(0)
            if archivo.stream.read(5) != b"%PDF-":
                raise ValueError(
                    f'El archivo "{archivo.filename}" no parece ser un PDF válido.'
                )
            archivo.stream.seek(0)
            ficha_sugerida = None
            try:
                datos_validacion = archivo.stream.read()
                archivo.stream.seek(0)
                if len(datos_validacion) > MAX_PDF_FILE_SIZE_BYTES:
                    raise ValueError(
                        f'El PDF "{archivo.filename}" supera el máximo de '
                        f'{MAX_PDF_FILE_SIZE_BYTES // (1024 * 1024)} MB.'
                    )
                documento = fitz.open(stream=datos_validacion, filetype="pdf")
                documento.close()
                # P1.7 — proponer ficha desde el mismo bytes (sin segunda lectura).
                try:
                    texto_pdf = extraer_texto_pdf_bytes(datos_validacion)
                    ficha_sugerida = _proponer_ficha_desde_manual(
                        texto_pdf, compania=compania, nombre_archivo=nombre_seguro
                    )
                except Exception as ficha_err:
                    print("ADVERTENCIA ficha sugerida manual:", ficha_err)
                    ficha_sugerida = None
                del datos_validacion
            except Exception as exc:
                raise ValueError(
                    f'No se pudo leer el PDF "{archivo.filename}". '
                    "Verificá que no esté dañado."
                ) from exc
            archivo.stream.seek(0)
        except ValueError as exc:
            return jsonify(
                ok=False,
                error=str(exc),
                cargados=resultados,
            ), 400
        except Exception as exc:
            print("ERROR VALIDANDO MANUAL:", exc)
            return jsonify(
                ok=False,
                error=f'No se pudo validar el PDF "{archivo.filename}".',
                cargados=resultados,
            ), 400

        try:
            r2_subir_pdf(archivo.stream, r2_key_nuevo, tamaño)
            subidos_r2.append(r2_key_nuevo)
        except Exception as error:
            print("ERROR SUBIENDO MANUAL A R2:", error)
            # Si esta carga no pudo completarse, eliminamos sólo los objetos
            # creados por esta petición que todavía no tienen registro.
            for key in subidos_r2:
                try:
                    if not any(r.get("archivo") == key for r in resultados):
                        r2_eliminar_pdf(key)
                except Exception as rollback_error:
                    print("ERROR ROLLBACK R2:", rollback_error)
            return jsonify(
                ok=False,
                error=f'No se pudo guardar "{archivo.filename}" en Cloudflare R2.',
                cargados=resultados,
            ), 502

        try:
            if r2_key_anterior:
                actualizar_manual(
                    r2_key_anterior,
                    nombre_seguro,
                    r2_key_nuevo,
                    tamaño,
                )
            else:
                registrar_manual(
                    nombre_seguro,
                    r2_key_nuevo,
                    tamaño,
                )
        except Exception as error:
            print("ERROR REGISTRANDO MANUAL EN NEON:", error)
            try:
                r2_eliminar_pdf(r2_key_nuevo)
            except Exception as rollback_error:
                print("ERROR ROLLBACK R2:", rollback_error)

            return jsonify(
                ok=False,
                error=(
                    f'"{archivo.filename}" se subió a R2, pero no pudo '
                    "registrarse en PostgreSQL. La operación no se completó."
                ),
                cargados=resultados,
            ), 502

        item_resultado = {
            "archivo": r2_key_nuevo,
            "nombre": nombre_seguro,
            "tamaño": tamaño,
        }
        if ficha_sugerida:
            item_resultado["ficha_sugerida"] = ficha_sugerida
        resultados.append(item_resultado)

        # Sólo un reemplazo explícito elimina el objeto anterior. Una carga
        # normal jamás toca los manuales existentes.
        if r2_key_anterior:
            try:
                r2_eliminar_pdf(r2_key_anterior)
            except Exception as error:
                print("ERROR ELIMINANDO EL PDF ANTERIOR DE R2:", error)
                try:
                    actualizar_manual(
                        r2_key_nuevo,
                        existente["nombre"],
                        r2_key_anterior,
                        existente["tamaño"],
                    )
                except Exception as rollback_db_error:
                    print("ERROR ROLLBACK NEON:", rollback_db_error)
                try:
                    r2_eliminar_pdf(r2_key_nuevo)
                except Exception as rollback_r2_error:
                    print("ERROR ROLLBACK R2:", rollback_r2_error)
                return jsonify(
                    ok=False,
                    error=(
                        "No se pudo completar el reemplazo porque el PDF "
                        "anterior no pudo eliminarse de R2."
                    ),
                    cargados=[],
                ), 502

            # Evita que el siguiente elemento de una hipotética petición múltiple
            # se interprete como otro reemplazo.
            r2_key_anterior = None

    return jsonify(
        ok=True,
        mensaje=(
            f"{len(resultados)} manual(es) de {compania} "
            "cargado(s) correctamente."
        ),
        archivos=resultados,
        cantidad=len(resultados),
    )


@app.route("/api/manuales/<slug>/<path:nombre_archivo>", methods=["DELETE"])
@requiere_admin
def eliminar_manual(slug, nombre_archivo):
    if slug not in {slug_manual_compania(c) for c in MANUALES_COMPANIAS}:
        return jsonify(ok=False, error="Compañía no válida."), 404

    r2_key = str(nombre_archivo or "").strip()
    prefijo = f"manuales/{slug}/"
    if not r2_key.startswith(prefijo) or not r2_key.lower().endswith(".pdf"):
        return jsonify(ok=False, error="Manual no válido para esa compañía."), 400

    existente = obtener_manual_por_r2_key(r2_key)
    if not existente:
        return jsonify(ok=False, error="Manual no encontrado."), 404

    # Primero quitamos el registro de Neon. Si R2 falla, restauramos el
    # registro para que la operación quede atómica desde el punto de vista
    # de la aplicación.
    try:
        eliminado = eliminar_manual_pg(r2_key)
        if not eliminado:
            return jsonify(ok=False, error="Manual no encontrado."), 404
    except Exception as error:
        print("ERROR ELIMINANDO MANUAL DE NEON:", error)
        return jsonify(
            ok=False,
            error="No se pudo actualizar PostgreSQL. El PDF no fue eliminado."
        ), 502

    try:
        r2_eliminar_pdf(r2_key)
    except Exception as error:
        print("ERROR ELIMINANDO MANUAL DE R2:", error)
        try:
            registrar_manual(
                existente["nombre"],
                existente["r2_key"],
                existente["tamaño"],
            )
        except Exception as rollback_error:
            print("ERROR RESTAURANDO MANUAL EN NEON:", rollback_error)
        return jsonify(
            ok=False,
            error="No se pudo eliminar el PDF de Cloudflare R2. El manual se mantuvo registrado."
        ), 502

    return jsonify(ok=True)


@app.route("/manuales/<slug>/<path:nombre_archivo>")
@requiere_login
def ver_manual(slug, nombre_archivo):
    if slug not in {slug_manual_compania(c) for c in MANUALES_COMPANIAS}:
        return ("Manual no encontrado", 404)

    r2_key = str(nombre_archivo or "").strip()
    prefijo = f"manuales/{slug}/"
    if not r2_key.startswith(prefijo) or not r2_key.lower().endswith(".pdf"):
        return ("Manual no encontrado", 404)

    existente = obtener_manual_por_r2_key(r2_key)
    if not existente:
        return ("Manual no encontrado", 404)

    try:
        objeto = obtener_objeto_stream(r2_key)
        body = objeto["Body"]
        content_length = objeto.get("ContentLength")

        @stream_with_context
        def generar():
            try:
                while True:
                    bloque = body.read(1024 * 1024)
                    if not bloque:
                        break
                    yield bloque
            finally:
                body.close()

        headers = {
            "Content-Type": "application/pdf",
            "Content-Disposition": f'inline; filename="{existente["nombre"]}"',
            "Cache-Control": "private, max-age=300",
        }
        if content_length is not None:
            headers["Content-Length"] = str(content_length)

        return Response(generar(), headers=headers)

    except Exception as error:
        print("ERROR SIRVIENDO MANUAL R2:", error)
        return ("No se pudo abrir el manual.", 502)


# ==========================================================
# LOGOUT
# ==========================================================

@app.route("/logout")
def logout():

    session.clear()

    return redirect(
        url_for("login")
    )


# ==========================================================
# CREAR ESTRUCTURA
# ==========================================================

def crear_estructura():
    MANUALES_DIR.mkdir(parents=True, exist_ok=True)
    try:
        inicializar_postgres()
        print('NEON POSTGRESQL: tabla manuales verificada.')
    except Exception as error:
        print('NEON POSTGRESQL: no se pudo verificar la tabla manuales:', error)
    POLIZAS_DIR.mkdir(parents=True, exist_ok=True)
    # Las compañías de documentos deben coincidir exactamente con las 9
    # compañías habilitadas en la sección Manuales.
    companias = [
        "atm",
        "mercantil_andina",
        "federacion_patronal",
        "san_cristobal",
        "rivadavia",
        "euroamerica",
        "agrosalta",
        "triunfo",
        "prof",
    ]

    for compania in companias:
        (DOCUMENTOS_DIR / compania).mkdir(parents=True, exist_ok=True)


# ==========================================================
# INICIAR
# ==========================================================

inicializar_base_datos()

try:
    inicializar_postgres()
    print('NEON POSTGRESQL: tabla manuales verificada.')
except Exception as error:
    print('NEON POSTGRESQL: no se pudo verificar la tabla manuales al iniciar:', error)




# ==========================================================
# OFICINAIA PLUS — Bandeja de pendientes + ficha desde texto
# ==========================================================

PLANTILLAS_METADATO = [
    {
        "id": "remolque",
        "titulo": "Remolque / asistencia — {COMPANIA} {PRODUCTO}",
        "contenido": (
            "COMPAÑÍA: {COMPANIA}\n"
            "PRODUCTO / PLAN: {PRODUCTO}\n\n"
            "Servicios de remolque al año:\n"
            "Kilómetros por servicio (ida + vuelta):\n"
            "Servicio especial (si aplica):\n"
            "Límites o exclusiones:\n"
            "Observaciones operativas:\n"
        ),
    },
    {
        "id": "cobertura",
        "titulo": "Cobertura — {COMPANIA} {PRODUCTO}",
        "contenido": (
            "COMPAÑÍA: {COMPANIA}\n"
            "PRODUCTO / PLAN: {PRODUCTO}\n\n"
            "Responsabilidad civil:\n"
            "Robo / hurto total:\n"
            "Incendio total:\n"
            "Daños parciales / total:\n"
            "Granizo / cristales:\n"
            "Franquicias:\n"
            "Exclusiones relevantes:\n"
        ),
    },
    {
        "id": "procedimiento",
        "titulo": "Procedimiento — {TEMA}",
        "contenido": (
            "TEMA: {TEMA}\n\n"
            "Cuándo aplica:\n"
            "Pasos a seguir:\n"
            "1.\n2.\n3.\n"
            "Documentación requerida:\n"
            "Contactos / teléfonos:\n"
            "Notas internas:\n"
        ),
    },
    {
        "id": "contacto",
        "titulo": "Contacto operativo — {COMPANIA}",
        "contenido": (
            "COMPAÑÍA: {COMPANIA}\n"
            "Área:\n"
            "Teléfono:\n"
            "Mail:\n"
            "Horario:\n"
            "Observaciones:\n"
        ),
    },
]


def _pendientes_usar_pg():
    """P0.2 — Con Neon los pendientes viven en Postgres (sobreviven redeploy)."""
    return bool(os.getenv("DATABASE_URL"))


def _conectar_db_pendientes():
    return conectar_db()


@app.route("/api/pendientes", methods=["GET"])
@requiere_login
def api_listar_pendientes():
    estado = (request.args.get("estado") or "pendiente").strip().lower()
    if estado not in pendientes_ops.ESTADOS_VALIDOS:
        estado = "pendiente"
    if _pendientes_usar_pg():
        try:
            from database_pg import listar_pendientes as pg_listar, contar_pendientes as pg_contar
            items = pg_listar(session["usuario"], estado=estado)
            total = pg_contar(session["usuario"], estado="pendiente")
            return jsonify({"ok": True, "pendientes": items, "total_pendientes": total})
        except Exception as error:
            print("ERROR api_listar_pendientes PG:", error)
            return jsonify({"ok": False, "error": "No se pudieron cargar los pendientes."}), 500
    with closing(_conectar_db_pendientes()) as db:
        pendientes_ops.asegurar_tabla(db)
        items = pendientes_ops.listar(db, session["usuario"], estado=estado)
        total = pendientes_ops.contar(db, session["usuario"], estado="pendiente")
    return jsonify({"ok": True, "pendientes": items, "total_pendientes": total})


@app.route("/api/pendientes", methods=["POST"])
@requiere_login
def api_crear_pendiente():
    data = request.get_json(silent=True) or {}
    tipo = str(data.get("tipo") or "generico")
    titulo = str(data.get("titulo") or "Pendiente").strip()
    payload = data.get("payload") if isinstance(data.get("payload"), dict) else {}
    if _pendientes_usar_pg():
        try:
            from database_pg import crear_pendiente as pg_crear, contar_pendientes as pg_contar
            pid = pg_crear(session["usuario"], tipo, titulo, payload)
            total = pg_contar(session["usuario"], estado="pendiente")
            return jsonify({"ok": True, "id": pid, "total_pendientes": total})
        except Exception as error:
            print("ERROR api_crear_pendiente PG:", error)
            return jsonify({"ok": False, "error": "No se pudo crear el pendiente."}), 500
    with closing(_conectar_db_pendientes()) as db:
        pendientes_ops.asegurar_tabla(db)
        pid = pendientes_ops.crear(db, session["usuario"], tipo, titulo, payload)
        total = pendientes_ops.contar(db, session["usuario"], estado="pendiente")
    return jsonify({"ok": True, "id": pid, "total_pendientes": total})


@app.route("/api/pendientes/<int:pendiente_id>", methods=["PATCH"])
@requiere_login
def api_actualizar_pendiente(pendiente_id):
    # Tanda 8: este PATCH ahora sirve tanto para el cambio rápido de estado
    # (marcar hecho/descartado, como antes) como para editar un pendiente
    # completo (título, tipo, contenido) desde la solapa. Sólo se tocan los
    # campos que vengan en el body; el resto queda igual.
    data = request.get_json(silent=True) or {}

    estado = data.get("estado")
    if estado is not None:
        estado = str(estado).strip().lower()
        if estado not in pendientes_ops.ESTADOS_VALIDOS:
            return jsonify({"ok": False, "error": "Estado no válido."}), 400

    titulo = data.get("titulo")
    if titulo is not None:
        titulo = str(titulo).strip()
        if not titulo:
            return jsonify({"ok": False, "error": "El título no puede quedar vacío."}), 400

    tipo = data.get("tipo")
    if tipo is not None:
        tipo = str(tipo).strip().lower()

    payload = data.get("payload") if isinstance(data.get("payload"), dict) else None

    if estado is None and titulo is None and tipo is None and payload is None:
        return jsonify({"ok": False, "error": "No hay cambios para guardar."}), 400

    if _pendientes_usar_pg():
        try:
            from database_pg import editar_pendiente as pg_editar, contar_pendientes as pg_contar
            ok = pg_editar(pendiente_id, session["usuario"], tipo=tipo, titulo=titulo, payload=payload, estado=estado)
            if not ok:
                return jsonify({"ok": False, "error": "Pendiente no encontrado."}), 404
            total = pg_contar(session["usuario"], estado="pendiente")
            return jsonify({"ok": True, "total_pendientes": total})
        except Exception as error:
            print("ERROR api_actualizar_pendiente PG:", error)
            return jsonify({"ok": False, "error": "No se pudo actualizar el pendiente."}), 500
    with closing(_conectar_db_pendientes()) as db:
        pendientes_ops.asegurar_tabla(db)
        ok = pendientes_ops.editar(db, pendiente_id, session["usuario"], tipo=tipo, titulo=titulo, payload=payload, estado=estado)
        if not ok:
            return jsonify({"ok": False, "error": "Pendiente no encontrado."}), 404
        total = pendientes_ops.contar(db, session["usuario"], estado="pendiente")
    return jsonify({"ok": True, "total_pendientes": total})


@app.route("/api/pendientes/<int:pendiente_id>", methods=["DELETE"])
@requiere_login
def api_eliminar_pendiente(pendiente_id):
    if _pendientes_usar_pg():
        try:
            from database_pg import eliminar_pendiente as pg_del, contar_pendientes as pg_contar
            ok = pg_del(pendiente_id, session["usuario"])
            if not ok:
                return jsonify({"ok": False, "error": "Pendiente no encontrado."}), 404
            total = pg_contar(session["usuario"], estado="pendiente")
            return jsonify({"ok": True, "total_pendientes": total})
        except Exception as error:
            print("ERROR api_eliminar_pendiente PG:", error)
            return jsonify({"ok": False, "error": "No se pudo eliminar el pendiente."}), 500
    with closing(_conectar_db_pendientes()) as db:
        pendientes_ops.asegurar_tabla(db)
        ok = pendientes_ops.eliminar(db, pendiente_id, session["usuario"])
        if not ok:
            return jsonify({"ok": False, "error": "Pendiente no encontrado."}), 404
        total = pendientes_ops.contar(db, session["usuario"], estado="pendiente")
    return jsonify({"ok": True, "total_pendientes": total})


@app.route("/api/plantillas-metadato", methods=["GET"])
@requiere_login
def api_plantillas_metadato():
    return jsonify({"ok": True, "plantillas": PLANTILLAS_METADATO})


@app.route("/api/ficha-desde-texto", methods=["POST"])
@requiere_login
def api_ficha_desde_texto():
    """Arma una ficha de metadato a partir de texto (respuesta de Sofia o extracto de PDF)."""
    data = request.get_json(silent=True) or {}
    texto = str(data.get("texto") or "").strip()
    titulo = str(data.get("titulo") or "").strip()
    if not texto:
        return jsonify({"ok": False, "error": "No hay texto para armar la ficha."}), 400
    if not titulo:
        # primeras palabras útiles
        linea = texto.split("\n", 1)[0].strip()
        titulo = (" ".join(linea.split())[:80] or "Ficha desde chat")
    # Normalizar un poco el contenido conservando información
    cuerpo = texto.strip()
    if len(cuerpo) > 12000:
        cuerpo = cuerpo[:12000] + "\n\n[Texto recortado por longitud]"
    return jsonify({
        "ok": True,
        "ficha": {
            "titulo": titulo[:200],
            "contenido": cuerpo,
        },
    })


@app.route("/api/validar-excel-fila", methods=["POST"])
@requiere_login
def api_validar_excel_fila():
    """Validación liviana de una fila de asegurados/flota antes de persistir.
    P1.6 — cableada desde el front; incluye aviso de patente duplicada.
    """
    data = request.get_json(silent=True) or {}
    libro_id = str(data.get("libro_id") or "1")
    campos = data.get("campos") if isinstance(data.get("campos"), dict) else {}
    avisos = []
    errores = []

    def norm(v):
        return str(v or "").strip()

    patente_limpia = ""
    if libro_id == "1":
        aseg = norm(campos.get("ASEGURADO") or campos.get("asegurado"))
        num = norm(campos.get("NUMERO") or campos.get("numero"))
        pat = norm(campos.get("PATENTE") or campos.get("patente"))
        cia = norm(campos.get("CIA") or campos.get("cia"))
        if not aseg:
            errores.append("Falta ASEGURADO.")
        if not num and not pat:
            errores.append("Completá NUMERO (DNI/póliza) o PATENTE.")
        if pat:
            limpio = re.sub(r"[^A-Za-z0-9]", "", pat).upper()
            if len(limpio) < 6 or len(limpio) > 8:
                avisos.append(f"La patente '{pat}' tiene un formato poco habitual.")
            campos["PATENTE"] = limpio or pat
            patente_limpia = limpio
        if not cia:
            avisos.append("CIA vacío: conviene completarlo para búsquedas futuras.")
    else:
        pat = norm(campos.get("patente") or campos.get("PATENTE"))
        if not pat:
            errores.append("Falta patente del vehículo.")
        else:
            limpio = re.sub(r"[^A-Za-z0-9]", "", pat).upper()
            if len(limpio) < 6 or len(limpio) > 8:
                avisos.append(f"La patente '{pat}' tiene un formato poco habitual.")
            campos["patente"] = limpio or pat
            patente_limpia = limpio

    # Anti-duplicado por patente (aviso, no bloqueo duro).
    if patente_limpia and len(errores) == 0:
        try:
            datos = leer_excel_interno(libro_id)
            filas = datos.get("filas") or []
            if filas:
                headers = [str(h or "").strip().upper() for h in filas[0]]
                idx_pat = None
                for i, h in enumerate(headers):
                    if h in ("PATENTE", "DOMINIO", "CHAPA"):
                        idx_pat = i
                        break
                if idx_pat is not None:
                    for row in filas[1:]:
                        if idx_pat >= len(row):
                            continue
                        existente = re.sub(r"[^A-Za-z0-9]", "", str(row[idx_pat] or "")).upper()
                        if existente and existente == patente_limpia:
                            avisos.append(
                                f"La patente {patente_limpia} ya figura en el Excel. "
                                "Revisá si es un duplicado antes de guardar."
                            )
                            break
        except Exception as error:
            print("ADVERTENCIA validar patente duplicada:", error)

    return jsonify({
        "ok": len(errores) == 0,
        "errores": errores,
        "avisos": avisos,
        "campos": campos,
    })


@app.route("/pendientes")
@requiere_login
def pagina_pendientes():
    return render_template("pendientes.html", usuario=session["usuario"])

@app.route("/internet")
@requiere_login
def pagina_internet():
    return render_template("internet.html", usuario=session["usuario"])



if __name__ == "__main__":

    crear_estructura()

    print("")
    print(
        "===================================="
    )

    print(
        "     OFICINA SEGUROS INICIADA"
    )

    print(
        "===================================="
    )

    print(
        "Servidor: http://127.0.0.1:5000"
    )

    print("")

    app.run(
        debug=True,
        host="127.0.0.1",
        port=5000
    )